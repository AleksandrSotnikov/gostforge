# ruff: noqa: RUF002, RUF003
"""Аннотация .docx: пометки нарушений в документе.

Два режима работы:

* ``style="comments"`` (по умолчанию) — настоящие OOXML-комментарии Word.
  Создаётся отдельная part ``word/comments.xml`` в docx-архиве, в
  ``document.xml`` ставятся пары ``<w:commentRangeStart/>`` /
  ``<w:commentRangeEnd/>`` и run с ``<w:commentReference/>``. Word и
  LibreOffice отображают такие комментарии как боковые выноски, на них
  можно отвечать и резолвить их.

* ``style="inline"`` — старый режим (Фаза 1): в начало проблемного параграфа
  вставляется TextRun с текстом ``[CODE: message]`` красным курсивом.
  Сохраняется как fallback, для совместимости и в случаях, когда нужно
  «вшить» пометки прямо в текст (например, при экспорте в форматы, где
  комментариев нет).

Маппинг Violation → параграф (общий для обоих режимов):

* Если ``location`` содержит ``paragraph[p-N]`` — ищем N-ый параграф
  ``<w:p>`` в ``<w:body>`` (нумерация с 1).
* Иначе — пометка считается document-level и привязывается к первому
  параграфу документа.
"""

from __future__ import annotations

import re
import shutil
import zipfile
from datetime import UTC, datetime
from pathlib import Path
from typing import Literal, cast

from docx import Document as DocxDocument
from docx.shared import RGBColor
from lxml import etree  # type: ignore[import-untyped]

from gostforge.parser import parse_docx
from gostforge.profile import Profile
from gostforge.validator import validate
from gostforge.validator.engine import Violation

# RGB красного цвета для inline-маркеров.
_MARKER_COLOR = RGBColor(0xC0, 0x00, 0x00)

# Локации, для которых не существует «настоящего» параграфа в документе —
# тогда маркер падает в первый параграф.
_DOCUMENT_LEVEL_PREFIXES = (
    "page_sections.",  # геометрия страницы, секции вёрстки
    "header.",
    "footer.",
    "metadata.",
    "bibliography.",
)

# Регулярки для извлечения идентификатора блока из location.
# Пример: "page_sections.main.children[3]" → не block-level.
# Пример: "paragraph.<id>" или "paragraph[<id>]" → block-level.
_BLOCK_ID_PATTERNS = (
    re.compile(r"paragraph[s]?[\.\[]([a-zA-Z0-9_\-]+)\]?"),
    re.compile(r"figure[s]?[\.\[]([a-zA-Z0-9_\-]+)\]?"),
    re.compile(r"table[s]?[\.\[]([a-zA-Z0-9_\-]+)\]?"),
    re.compile(r"block[s]?[\.\[]([a-zA-Z0-9_\-]+)\]?"),
)

# Регулярка для извлечения индекса параграфа из ID вида "p-12".
_PARAGRAPH_ID_INDEX = re.compile(r"^p-(\d+)$")

# OOXML namespaces.
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_PR_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
_CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
_NSMAP_W = {"w": _W_NS}

_COMMENTS_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
)
_COMMENTS_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)


def _extract_block_id(location: str) -> str | None:
    """Вытащить ID блока из location-строки, если он там есть."""
    for pattern in _BLOCK_ID_PATTERNS:
        match = pattern.search(location)
        if match:
            return match.group(1)
    return None


def _format_marker(violation: Violation) -> str:
    """Сформировать текст маркера для нарушения (inline-режим)."""
    return f"[{violation.check_code}: {violation.message}]"


def _format_comment_text(violation: Violation) -> str:
    """Сформировать полный текст для OOXML-комментария.

    Включает код проверки, сообщение и suggestion, если он есть.
    """
    text = f"{violation.check_code}: {violation.message}"
    if violation.suggestion:
        text += f" → {violation.suggestion}"
    return text


def _insert_marker_run(paragraph: object, text: str) -> None:
    """Вставить run с маркером в начало параграфа.

    Использует низкоуровневое API python-docx: создаёт новый run и
    переносит его в начало списка дочерних элементов параграфа.
    """
    # paragraph — это docx.text.paragraph.Paragraph; пользуемся его публичным
    # API add_run, затем «выкручиваем» получившийся <w:r> в начало <w:p>.
    run = paragraph.add_run(text + " ")  # type: ignore[attr-defined]
    run.italic = True
    run.font.color.rgb = _MARKER_COLOR

    # Переместить элемент run в начало <w:p>. Все существующие runs сдвинутся.
    p_elem = paragraph._element  # type: ignore[attr-defined]
    r_elem = run._element
    p_elem.remove(r_elem)
    # Вставить после <w:pPr>, если он есть; иначе в самое начало.
    ppr = p_elem.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr"
    )
    if ppr is not None:
        ppr.addnext(r_elem)
    else:
        p_elem.insert(0, r_elem)


def _is_document_level(location: str) -> bool:
    """Понять, относится ли нарушение к документу в целом, а не к блоку."""
    if not location:
        return True
    return any(location.startswith(prefix) for prefix in _DOCUMENT_LEVEL_PREFIXES)


def _resolve_paragraph_index(location: str, total: int) -> int | None:
    """Сопоставить location → 0-based индекс параграфа в <w:body>.

    Возвращает None, если location не разрешается в конкретный параграф.
    """
    block_id = _extract_block_id(location)
    if block_id is None:
        return None
    match = _PARAGRAPH_ID_INDEX.match(block_id)
    if not match:
        return None
    idx_1_based = int(match.group(1))
    idx_0_based = idx_1_based - 1
    if 0 <= idx_0_based < total:
        return idx_0_based
    return None


def _annotate_inline(
    input_path: Path,
    output_path: Path,
    violations: list[Violation],
) -> int:
    """Старый режим: inline-маркеры в начале параграфов."""
    docx_doc = DocxDocument(str(input_path))
    paragraphs = list(docx_doc.paragraphs)

    if not paragraphs:
        # Документ без параграфов — создаём один пустой, чтобы было куда писать.
        docx_doc.add_paragraph("")
        paragraphs = list(docx_doc.paragraphs)

    inserted = 0
    for violation in violations:
        idx = _resolve_paragraph_index(violation.location, len(paragraphs))
        # Если location не разрешается в параграф (document-level или неизвестный
        # ID) — пометка падает в первый параграф документа.
        target = paragraphs[idx] if idx is not None else paragraphs[0]
        _insert_marker_run(target, _format_marker(violation))
        inserted += 1

    docx_doc.save(str(output_path))
    return inserted


def _build_comments_xml(comments: list[tuple[int, str, str]]) -> bytes:
    """Собрать XML тело ``word/comments.xml``.

    ``comments`` — список ``(comment_id, author, text)``.
    """
    root = etree.Element(f"{{{_W_NS}}}comments", nsmap=_NSMAP_W)
    iso_date = datetime.now(UTC).strftime("%Y-%m-%dT%H:%M:%SZ")
    for cid, author, text in comments:
        comment = etree.SubElement(
            root,
            f"{{{_W_NS}}}comment",
            attrib={
                f"{{{_W_NS}}}id": str(cid),
                f"{{{_W_NS}}}author": author,
                f"{{{_W_NS}}}date": iso_date,
                f"{{{_W_NS}}}initials": "GF",
            },
        )
        p = etree.SubElement(comment, f"{{{_W_NS}}}p")
        r = etree.SubElement(p, f"{{{_W_NS}}}r")
        t = etree.SubElement(r, f"{{{_W_NS}}}t")
        t.set(
            "{http://www.w3.org/XML/1998/namespace}space",
            "preserve",
        )
        t.text = text
    return cast(
        bytes,
        etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone=True
        ),
    )


def _insert_comment_refs(
    document_xml: bytes, plan: list[tuple[int, int]]
) -> bytes:
    """Вставить ``<w:commentRangeStart/End>`` и reference-run в document.xml.

    ``plan`` — список ``(paragraph_index_0_based, comment_id)``. Для одного
    параграфа может быть несколько записей: они обрамляют один и тот же
    набор runs независимыми диапазонами (вложенные `commentRange`-ы Word
    поддерживает корректно).
    """
    # Сгруппируем по индексу параграфа, сохранив порядок добавления.
    by_para: dict[int, list[int]] = {}
    for para_idx, cid in plan:
        by_para.setdefault(para_idx, []).append(cid)

    root = etree.fromstring(document_xml)
    body = root.find(f"{{{_W_NS}}}body")
    if body is None:
        return document_xml

    # Собираем все <w:p> внутри <w:body>. Сюда попадают и параграфы внутри
    # ячеек таблиц — но нумерация Violation.location у нас плоская, по
    # порядку появления <w:p> в body (см. parser/docx_parser.py:923).
    paragraphs = body.findall(f".//{{{_W_NS}}}p")

    for para_idx, cids in by_para.items():
        if para_idx < 0 or para_idx >= len(paragraphs):
            continue
        p_elem = paragraphs[para_idx]

        # Определяем точку, после которой следуют content-runs (после <w:pPr>).
        ppr = p_elem.find(f"{{{_W_NS}}}pPr")
        insert_after_ppr = ppr is not None

        for cid in cids:
            start = etree.Element(f"{{{_W_NS}}}commentRangeStart")
            start.set(f"{{{_W_NS}}}id", str(cid))
            end = etree.Element(f"{{{_W_NS}}}commentRangeEnd")
            end.set(f"{{{_W_NS}}}id", str(cid))
            ref_run = etree.Element(f"{{{_W_NS}}}r")
            rpr = etree.SubElement(ref_run, f"{{{_W_NS}}}rPr")
            rstyle = etree.SubElement(rpr, f"{{{_W_NS}}}rStyle")
            rstyle.set(f"{{{_W_NS}}}val", "CommentReference")
            ref = etree.SubElement(ref_run, f"{{{_W_NS}}}commentReference")
            ref.set(f"{{{_W_NS}}}id", str(cid))

            # Вставляем commentRangeStart сразу после <w:pPr> (или в начало <w:p>).
            if insert_after_ppr:
                assert ppr is not None
                ppr.addnext(start)
            else:
                p_elem.insert(0, start)
            # commentRangeEnd и reference-run — в самый конец <w:p>.
            p_elem.append(end)
            p_elem.append(ref_run)

    return cast(
        bytes,
        etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone=True
        ),
    )


def _update_content_types(content_types_xml: bytes) -> bytes:
    """Добавить Override для ``/word/comments.xml`` в ``[Content_Types].xml``.

    Если override уже есть — возвращает исходный XML без изменений.
    """
    root = etree.fromstring(content_types_xml)
    override_tag = f"{{{_CT_NS}}}Override"
    for override in root.findall(override_tag):
        if override.get("PartName") == "/word/comments.xml":
            return content_types_xml
    new_override = etree.SubElement(root, override_tag)
    new_override.set("PartName", "/word/comments.xml")
    new_override.set("ContentType", _COMMENTS_CONTENT_TYPE)
    return cast(
        bytes,
        etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone=True
        ),
    )


def _update_document_rels(rels_xml: bytes) -> tuple[bytes, bool]:
    """Добавить relationship на ``comments.xml`` в ``word/_rels/document.xml.rels``.

    Возвращает (новый XML, был ли добавлен relationship).
    Если relationship уже существует — возвращает (исходный XML, False).
    """
    root = etree.fromstring(rels_xml)
    rel_tag = f"{{{_PR_NS}}}Relationship"
    existing_ids: set[str] = set()
    for rel in root.findall(rel_tag):
        if rel.get("Type") == _COMMENTS_REL_TYPE:
            return rels_xml, False
        rid = rel.get("Id")
        if rid:
            existing_ids.add(rid)

    # Подобрать свободный rId.
    n = 1
    while f"rId{n}" in existing_ids:
        n += 1
    new_id = f"rId{n}"

    new_rel = etree.SubElement(root, rel_tag)
    new_rel.set("Id", new_id)
    new_rel.set("Type", _COMMENTS_REL_TYPE)
    new_rel.set("Target", "comments.xml")
    return (
        cast(
            bytes,
            etree.tostring(
                root, xml_declaration=True, encoding="UTF-8", standalone=True
            ),
        ),
        True,
    )


def _replace_zip_entries(
    src_path: Path, dst_path: Path, overrides: dict[str, bytes]
) -> None:
    """Создать новый zip-архив на основе ``src_path``, заменив/добавив записи.

    ``overrides`` — словарь ``{имя файла внутри zip: новое содержимое}``.
    Записи, не упомянутые в overrides, копируются как есть.

    Если запись уже была в исходном архиве — её содержимое заменяется.
    Если нет — добавляется. Это позволяет одновременно
    переписать ``document.xml`` и создать новую ``word/comments.xml``.
    """
    written: set[str] = set()
    with (
        zipfile.ZipFile(src_path, "r") as src_zip,
        zipfile.ZipFile(dst_path, "w", compression=zipfile.ZIP_DEFLATED) as dst_zip,
    ):
        for item in src_zip.infolist():
            name = item.filename
            if name in overrides:
                dst_zip.writestr(name, overrides[name])
                written.add(name)
            else:
                dst_zip.writestr(item, src_zip.read(name))
        # Добавляем новые записи, которых не было в исходном архиве.
        for name, data in overrides.items():
            if name not in written:
                dst_zip.writestr(name, data)


def _annotate_comments(
    input_path: Path,
    output_path: Path,
    violations: list[Violation],
    *,
    author: str = "gostforge",
) -> int:
    """Новый режим: настоящие OOXML-комментарии Word.

    Реализация — прямые манипуляции с zip-архивом .docx без использования
    внутреннего API python-docx. Это надёжнее: формат части ``comments.xml``
    и её relationship стабильны (описаны в ECMA-376), а API python-docx
    для комментариев официально не существует.

    Шаги:
    1. Прочитать ``word/document.xml`` из исходного .docx.
    2. Для каждого violation определить целевой 0-based индекс <w:p>.
    3. Сгенерировать ``word/comments.xml`` с по одному ``<w:comment>`` на нарушение.
    4. Вставить пары ``<w:commentRangeStart/End>`` и ``<w:r><w:commentReference/></w:r>``
       в соответствующие <w:p> в document.xml.
    5. Обновить ``[Content_Types].xml`` (Override для comments.xml).
    6. Обновить ``word/_rels/document.xml.rels`` (Relationship на comments.xml).
    7. Записать новый .docx в output_path.
    """
    # Прочитать нужные части исходного архива.
    with zipfile.ZipFile(input_path, "r") as z:
        names = set(z.namelist())
        if "word/document.xml" not in names:
            raise ValueError(f"{input_path}: не похоже на .docx — нет word/document.xml")
        document_xml = z.read("word/document.xml")
        content_types_xml = z.read("[Content_Types].xml")
        if "word/_rels/document.xml.rels" in names:
            rels_xml = z.read("word/_rels/document.xml.rels")
        else:
            # Минимально корректный rels-документ.
            rels_xml = (
                b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                b'<Relationships xmlns="' + _PR_NS.encode("ascii") + b'"/>'
            )

    # Подсчитать число <w:p> внутри <w:body>.
    doc_root = etree.fromstring(document_xml)
    body = doc_root.find(f"{{{_W_NS}}}body")
    total_paragraphs = (
        len(body.findall(f".//{{{_W_NS}}}p")) if body is not None else 0
    )

    if total_paragraphs == 0 and len(violations) > 0:
        # Документ без параграфов — комментарии негде ставить.
        # Создадим один пустой <w:p>, в который привяжем все комментарии.
        if body is None:
            body = etree.SubElement(doc_root, f"{{{_W_NS}}}body")
        etree.SubElement(body, f"{{{_W_NS}}}p")
        total_paragraphs = 1
        document_xml = cast(
            bytes,
            etree.tostring(
                doc_root, xml_declaration=True, encoding="UTF-8", standalone=True
            ),
        )

    # Спланировать вставки: (comment_id, para_index, text).
    comments_data: list[tuple[int, str, str]] = []
    insert_plan: list[tuple[int, int]] = []
    for cid, violation in enumerate(violations):
        idx = _resolve_paragraph_index(violation.location, total_paragraphs)
        if idx is None:
            idx = 0  # fallback: первый параграф
        if total_paragraphs == 0:
            # Параграфов нет, и не удалось создать (теоретически невозможно
            # после ветки выше) — пропускаем, иначе сломаем документ.
            continue
        comments_data.append((cid, author, _format_comment_text(violation)))
        insert_plan.append((idx, cid))

    if not comments_data:
        # Нет нарушений → просто скопировать исходник в output.
        shutil.copy(str(input_path), str(output_path))
        return 0

    # Собрать новые куски XML.
    new_document_xml = _insert_comment_refs(document_xml, insert_plan)
    comments_xml = _build_comments_xml(comments_data)
    new_content_types_xml = _update_content_types(content_types_xml)
    new_rels_xml, _ = _update_document_rels(rels_xml)

    overrides: dict[str, bytes] = {
        "word/document.xml": new_document_xml,
        "word/comments.xml": comments_xml,
        "[Content_Types].xml": new_content_types_xml,
        "word/_rels/document.xml.rels": new_rels_xml,
    }

    # Пересобрать архив.
    _replace_zip_entries(input_path, output_path, overrides)
    return len(comments_data)


def annotate_docx(
    input_path: str | Path,
    output_path: str | Path,
    profile: Profile,
    *,
    style: Literal["inline", "comments"] = "comments",
) -> int:
    """Прогнать .docx через парсер и валидатор, записать копию с пометками.

    Поведение:
    1. ``parse_docx(input_path)`` → Document
    2. ``validate(doc, profile)`` → list[Violation]
    3. Открыть ``input_path`` и для каждого Violation вставить пометку
       в соответствующий параграф (или в первый — для document-level).

    Параметр ``style``:

    * ``"comments"`` (по умолчанию) — настоящие OOXML-комментарии Word.
      Создаётся часть ``word/comments.xml`` и пары
      ``<w:commentRangeStart/End>`` + run с ``<w:commentReference/>``.
      Word и LibreOffice отображают их как боковые выноски, на которые
      можно отвечать и которые можно резолвить.

    * ``"inline"`` — старое поведение: в начало проблемного параграфа
      вставляется TextRun с текстом ``[CODE: message]`` курсивом, красным.

    Возвращает число вставленных комментариев/маркеров.
    """
    input_path = Path(input_path)
    output_path = Path(output_path)

    if not input_path.exists():
        raise FileNotFoundError(input_path)

    model = parse_docx(input_path)
    violations = validate(model, profile)

    if style == "inline":
        return _annotate_inline(input_path, output_path, violations)
    if style == "comments":
        return _annotate_comments(input_path, output_path, violations)
    raise ValueError(f"Unknown annotation style: {style!r}")
