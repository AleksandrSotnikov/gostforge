"""Рендер основной надписи (штампа ЕСКД, ГОСТ 2.104) в нижний колонтитул.

Форма 1 — для заглавного листа; форма 2а — сокращённая для последующих.
Штамп строится таблицей в footer-е, поэтому повторяется на каждой
странице секции (как и любой колонтитул).

Компоновка граф приближена к ГОСТ 2.104 (ширина 185 мм, канонические
ширины колонок), но не претендует на пиксельную точность формы —
геометрия может уточняться. Содержательно покрыты основные графы:
1 (наименование), 2 (обозначение), 4 (литера), 5 (масса), 6 (масштаб),
7 (лист), 8 (листов), 9 (организация), 11/13 (фамилии/даты ролей).
"""

from __future__ import annotations

from typing import Any

from docx.enum.table import (
    WD_CELL_VERTICAL_ALIGNMENT,
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt
from lxml import etree  # type: ignore[import-untyped]

from gostforge.model import TitleBlock

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Канонические ширины колонок основной надписи (мм), сумма = 185.
# Левый блок изменений/ролей: 7+10+23+15+10 = 65. Правый блок: 70 (графа 1
# наименование) + 25 + 25 (графы 4/7/8) = 120.
_FORM1_COLS_MM = [7, 10, 23, 15, 10, 70, 25, 25]
_FORM1_LABEL_PT = 7.0  # мелкий текст граф (компактно, как в ГОСТ-форме)
_FORM1_TITLE_PT = 9.0  # наименование / обозначение
# Высота строки штампа: ГОСТ 2.104 — строки ~5 мм, итог формы 2 ≈ 35–40 мм.
_FORM_ROW_MM = 5.0


def _set_table_borders(table: Any) -> None:
    """Сплошные рамки 0.5 pt на все стороны и внутренние линии таблицы."""
    tbl = table._element
    tbl_pr = tbl.find(f"{{{W_NS}}}tblPr")
    if tbl_pr is None:
        tbl_pr = etree.SubElement(tbl, f"{{{W_NS}}}tblPr")
    existing = tbl_pr.find(f"{{{W_NS}}}tblBorders")
    if existing is not None:
        tbl_pr.remove(existing)
    borders = etree.SubElement(tbl_pr, f"{{{W_NS}}}tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = etree.SubElement(borders, f"{{{W_NS}}}{side}")
        el.set(f"{{{W_NS}}}val", "single")
        el.set(f"{{{W_NS}}}sz", "4")  # 0.5 pt
        el.set(f"{{{W_NS}}}space", "0")
        el.set(f"{{{W_NS}}}color", "auto")


def _make_compact(table: Any, row_mm: float = _FORM_ROW_MM) -> None:
    """Сделать таблицу штампа компактной: нулевые отступы ячеек и
    фиксированная высота строк.

    Без этого Word авто-растягивает строки по содержимому и добавляет
    поля ячеек — штамп получается «огромным». Задаём точную высоту строк
    (``trHeight … exact``) и обнуляём ``tblCellMar``, чтобы форма
    укладывалась в канонические ~5 мм на строку.
    """
    tbl = table._element
    tbl_pr = tbl.find(f"{{{W_NS}}}tblPr")
    if tbl_pr is None:
        tbl_pr = etree.SubElement(tbl, f"{{{W_NS}}}tblPr")
    # Нулевые поля ячеек (top/bottom = 0, left/right = ~0.5 мм для воздуха).
    existing = tbl_pr.find(f"{{{W_NS}}}tblCellMar")
    if existing is not None:
        tbl_pr.remove(existing)
    cell_mar = etree.SubElement(tbl_pr, f"{{{W_NS}}}tblCellMar")
    for side, twips in (("top", 0), ("bottom", 0), ("left", 28), ("right", 28)):
        el = etree.SubElement(cell_mar, f"{{{W_NS}}}{side}")
        el.set(f"{{{W_NS}}}w", str(twips))
        el.set(f"{{{W_NS}}}type", "dxa")
    # Фиксированная высота строк.
    for row in table.rows:
        row.height = Mm(row_mm)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


def _set_cell(
    cell: Any,
    text: str,
    *,
    bold: bool = False,
    size_pt: float = _FORM1_LABEL_PT,
    align: str = "center",
) -> None:
    """Записать текст в ячейку штампа с нужным шрифтом/выравниванием.

    Вертикальное выравнивание — по центру; перезаписывает существующий
    единственный параграф ячейки, чтобы не плодить пустые строки.
    """
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    para.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    # Очищаем существующие run-ы.
    for run in list(para.runs):
        run._element.getparent().remove(run._element)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def _set_cell_sheet(cell: Any, tb: TitleBlock, *, prefix: str = "Лист ") -> None:
    """Записать графу 7 «Лист»: статичный номер или авто-поле PAGE.

    Если ``tb.sheet`` задан явно — пишем «{prefix}N» как обычный текст. Если
    пусто — вставляем ``prefix`` + OOXML-поле ``PAGE``, чтобы номер листа
    подставлялся автоматически по фактической странице (для основной
    надписи, повторяющейся в колонтитуле всех листов секции).
    """
    if tb.sheet:
        _set_cell(cell, f"{prefix}{tb.sheet}".strip())
        return
    _set_cell(cell, prefix)
    para = cell.paragraphs[0]
    fld = etree.SubElement(para._p, f"{{{W_NS}}}fldSimple")
    fld.set(f"{{{W_NS}}}instr", "PAGE")
    run = etree.SubElement(fld, f"{{{W_NS}}}r")
    rpr = etree.SubElement(run, f"{{{W_NS}}}rPr")
    rfonts = etree.SubElement(rpr, f"{{{W_NS}}}rFonts")
    rfonts.set(f"{{{W_NS}}}ascii", "Times New Roman")
    rfonts.set(f"{{{W_NS}}}hAnsi", "Times New Roman")
    sz = etree.SubElement(rpr, f"{{{W_NS}}}sz")
    sz.set(f"{{{W_NS}}}val", str(int(_FORM1_LABEL_PT * 2)))  # half-points
    rt = etree.SubElement(run, f"{{{W_NS}}}t")
    rt.text = ""


def _set_fixed_layout(table: Any, cols_mm: list[int]) -> None:
    """Зафиксировать ширины колонок (fixed layout, без автоподбора)."""
    table.autofit = False
    table.allow_autofit = False
    for i, width_mm in enumerate(cols_mm):
        for row in table.rows:
            row.cells[i].width = Mm(width_mm)


def _write_form1(footer: Any, tb: TitleBlock) -> Any:
    """Построить основную надпись «Форма 2» (заглавный лист текстового
    документа, ГОСТ 2.104) — компактная таблица 185×35 мм в footer-е.

    Раскладка граф:
    * слева сверху — шапка блока изменений «Изм.|Лист|№ докум.|Подп.|Дата»;
    * слева ниже — роли (графа 10/11/13): «Разраб./Пров./Н.контр./Утв.» с
      фамилией и датой;
    * справа сверху — обозначение документа (графа 2);
    * справа в центре — наименование (графа 1);
    * справа узкие графы — Лит. (4), Лист (7, авто-поле PAGE), Листов (8);
    * справа снизу — организация (графа 9).
    """
    rows = 7
    table = footer.add_table(rows=rows, cols=8, width=Mm(sum(_FORM1_COLS_MM)))
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    _set_fixed_layout(table, _FORM1_COLS_MM)

    # --- Левый блок: шапка изменений + роли --------------------------------
    for col, label in enumerate(("Изм.", "Лист", "№ докум.", "Подп.", "Дата")):
        _set_cell(table.cell(0, col), label)
    # Роли в строках 1..6: метка (c0-c1), фамилия (c2), дата (c4).
    roles = tb.roles[:6]
    for i in range(1, rows):
        table.cell(i, 0).merge(table.cell(i, 1))
        role = roles[i - 1] if i - 1 < len(roles) else None
        if role is not None:
            _set_cell(table.cell(i, 0), role.role, align="left")
            _set_cell(table.cell(i, 2), role.name)
            _set_cell(table.cell(i, 4), role.date)
        else:
            _set_cell(table.cell(i, 0), "")

    # --- Правый блок -------------------------------------------------------
    # Обозначение (графа 2) — верхняя строка на всю ширину правого блока.
    table.cell(0, 5).merge(table.cell(0, 7))
    _set_cell(table.cell(0, 5), tb.designation, bold=True, size_pt=_FORM1_TITLE_PT)
    # Наименование (графа 1) — c5, строки 1..3.
    table.cell(1, 5).merge(table.cell(3, 5))
    _set_cell(table.cell(1, 5), tb.title, bold=True, size_pt=_FORM1_TITLE_PT)
    # Узкие графы 4/7/8 — c6/c7, строки 1..3.
    _set_cell(table.cell(1, 6), "Лит.")
    _set_cell(table.cell(1, 7), "Лист")
    _set_cell(table.cell(2, 6), tb.stage)
    _set_cell_sheet(table.cell(2, 7), tb, prefix="")  # авто-PAGE
    _set_cell(table.cell(3, 6), "Листов")
    _set_cell(table.cell(3, 7), tb.sheets_total)
    # Организация (графа 9) — c5-c7, строки 4..6.
    table.cell(4, 5).merge(table.cell(rows - 1, 7))
    _set_cell(table.cell(4, 5), tb.organization, size_pt=_FORM1_LABEL_PT)

    _set_table_borders(table)
    _make_compact(table)
    return table


# Форма 2а — сокращённая (последующие листы): обозначение + номер листа.
_FORM2A_COLS_MM = [165, 20]


def _write_form2a(footer: Any, tb: TitleBlock) -> Any:
    """Построить форму 2а (последующие листы) — узкая таблица 185 мм."""
    table = footer.add_table(rows=1, cols=2, width=Mm(sum(_FORM2A_COLS_MM)))
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    _set_fixed_layout(table, _FORM2A_COLS_MM)
    _set_cell(table.cell(0, 0), tb.designation, bold=True, align="left")
    _set_cell_sheet(table.cell(0, 1), tb)
    _set_table_borders(table)
    _make_compact(table, row_mm=8.0)
    return table


def write_title_block(footer: Any, title_block: TitleBlock) -> Any | None:
    """Записать основную надпись в footer. Возвращает таблицу или None.

    Если штамп выключен (`enabled = False`) — ничего не пишет.
    """
    if not title_block.enabled:
        return None
    if title_block.form == "form2a":
        return _write_form2a(footer, title_block)
    return _write_form1(footer, title_block)
