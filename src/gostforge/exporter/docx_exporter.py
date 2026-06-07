"""Экспорт модели документа в .docx с применением стилей профиля.

Реализация Фазы 1:
- Поля страницы из профиля
- Стиль Normal (шрифт, кегль, межстрочный интервал, отступ первой строки)
- Параграфы со склейкой TextRun-ов и сохранением форматирования
- Per-paragraph переопределения (alignment, line_spacing, first_line_indent,
  page_break_before)
- Логические разделы как заголовки соответствующего уровня
- Таблицы с подписями и шапкой
- Рисунки экспортируются как заглушка-параграф (image_path не сохраняем
  на Фазе 1 — это потребует копирования media в docx)

Дальнейшие фазы (sectPr per-PageSection, header/footer-part, поля
PAGE/STYLEREF, реальные изображения, OMML-формулы) — отдельные итерации.
"""

from __future__ import annotations

import contextlib
import io
import logging
import re
from collections.abc import Sequence
from dataclasses import replace
from datetime import UTC
from pathlib import Path
from typing import Any, Literal

import docx
from docx.document import Document as DocxDocument
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from docx.text.paragraph import Paragraph as DocxParagraph
from lxml import etree  # type: ignore[import-untyped]

from gostforge.model import (
    Block,
    CellMerge,
    Citation,
    ContentTemplate,
    CrossRef,
    Document,
    Figure,
    FootnoteRef,
    Formula,
    Hyperlink,
    InlineElement,
    InlineFormula,
    ListBlock,
    LogicalSection,
    PageBorder,
    PageSection,
    Paragraph,
    Table,
    TableOfContents,
    TextRun,
    TitleBlock,
    TitleBlockRole,
)
from gostforge.profile import Profile

# Relationship namespace для w:hyperlink r:id="...".
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

logger = logging.getLogger(__name__)

# Контекст для одной операции экспорта: открытый исходный документ, из
# которого парсился `document`. Используется только в _write_figure для
# вставки реальных изображений (`embedded:rIdN`). Threading-небезопасно —
# для CLI и тестов достаточно; при необходимости можно перенести в
# ContextVar.
_current_source_docx: Any | None = None

# Контекст одной операции экспорта: 1-based индексы записей библиографии
# по их id. Используется в `_write_runs` для рендеринга Citation. Сбрасывается
# в `export_docx()` через try/finally, чтобы состояние не утекало между вызовами.
_current_bibliography_index: dict[str, int] | None = None

# Контекст одной операции экспорта: профиль для доступа к heading/caption/
# table-стилям внутри функций записи. Заполняется в export_docx() и
# сбрасывается через try/finally.
_current_profile: Any | None = None

# Per-document счётчик id-ов закладок: bookmarkStart/End требуют
# уникальные w:id внутри документа. Сбрасывается в export_docx()
# вместе с _current_profile.
_bookmark_id_counter: int = 0

_ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# Плейсхолдер, который парсер ставит при обнаружении поля PAGE; при
# экспорте мы материализуем его обратно в <w:fldSimple w:instr="PAGE"/>.
_PAGE_PLACEHOLDER = "{page}"


# Размеры бумаги в мм (портретная ориентация: short, long).
_PAPER_SIZES_MM: dict[str, tuple[float, float]] = {
    "A4": (210.0, 297.0),
    "A3": (297.0, 420.0),
    "A5": (148.0, 210.0),
    "Letter": (215.9, 279.4),
    "Legal": (215.9, 355.6),
}


def _apply_page_geometry(doc: DocxDocument, profile: Profile) -> None:
    """Применить поля страницы из профиля к первой секции docx."""
    margins = profile.styles.page.margins_mm
    section = doc.sections[0]
    if "top" in margins:
        section.top_margin = Mm(margins["top"])
    if "right" in margins:
        section.right_margin = Mm(margins["right"])
    if "bottom" in margins:
        section.bottom_margin = Mm(margins["bottom"])
    if "left" in margins:
        section.left_margin = Mm(margins["left"])


def _apply_page_size(sect: Any, page_section: PageSection) -> None:
    """Применить paper size и orientation из PageSection к docx-секции ``sect``.

    Если paper неизвестен — оставляем дефолт Word. Если orientation =
    landscape — width и height меняются местами относительно portrait.
    """
    paper = page_section.page.paper
    if paper not in _PAPER_SIZES_MM:
        return
    short, long_ = _PAPER_SIZES_MM[paper]
    if page_section.page.orientation == "landscape":
        width, height = long_, short
    else:
        width, height = short, long_
    sect.page_width = Mm(width)
    sect.page_height = Mm(height)


def _apply_section_margins(sect: Any, page_section: PageSection) -> None:
    """Применить поля страницы PageSection к docx-секции ``sect``.

    Используется при мультисекционном экспорте, чтобы у каждой секции были
    свои поля (например, у листов с рамкой ЕСКД левое поле шире). Поля,
    отсутствующие в ``margins_mm``, не трогаем — секция наследует значение.
    """
    margins = page_section.page.margins_mm
    if "top" in margins:
        sect.top_margin = Mm(margins["top"])
    if "right" in margins:
        sect.right_margin = Mm(margins["right"])
    if "bottom" in margins:
        sect.bottom_margin = Mm(margins["bottom"])
    if "left" in margins:
        sect.left_margin = Mm(margins["left"])


def _apply_normal_style(doc: DocxDocument, profile: Profile) -> None:
    """Применить шрифт/кегль/интервалы к стилю Normal.

    Дополнительно зануляет theme-fonts (asciiTheme/hAnsiTheme),
    которые Word из дефолтного шаблона ставит как majorHAnsi/minorHAnsi
    и которые при render-е в Word/LibreOffice перекрывают явно
    указанный font.name.
    """
    body = profile.styles.body
    normal = doc.styles["Normal"]
    normal.font.name = body.font
    normal.font.size = Pt(body.size_pt)
    # Снимаем theme-fonts со стиля Normal, иначе в Word шрифт может
    # отображаться как Calibri/Cambria поверх явно указанного.
    _clear_theme_fonts(normal.element, font_name=body.font)
    pf = normal.paragraph_format
    pf.line_spacing = body.line_spacing
    pf.first_line_indent = Cm(body.first_line_indent_cm)
    # Выравнивание основного текста — обычно justify по ГОСТу. Без
    # явной установки Word наследует left из своего дефолтного шаблона,
    # и абзацы основного текста выглядят рваными по правому краю.
    pf.alignment = _ALIGNMENT_MAP[body.alignment]
    # Интервалы между абзацами. По ГОСТу — 0; Word из дефолтного
    # шаблона ставит 'after=200 twips' (10 pt), что вылезает
    # между абзацами обычного текста как лишнее белое поле. Явно
    # сбрасываем по значению из профиля.
    pf.space_before = Pt(body.space_before_pt)
    pf.space_after = Pt(body.space_after_pt)


def _clear_theme_fonts(style_element: Any, *, font_name: str) -> None:
    """Удалить theme-атрибуты у w:rFonts и проставить явный шрифт.

    Word при создании документа через python-docx наследует stлей Normal
    из бортового шаблона: rFonts с asciiTheme/hAnsiTheme = minorHAnsi
    (Calibri). Это перекрывает явно заданный font.name при рендере.
    Решение — найти w:rFonts в xml-элементе стиля, убрать
    *Theme-атрибуты и проставить ascii/hAnsi/cs/eastAsia на нужный шрифт.
    """
    w_ns = W_NS
    rPr = style_element.find(f"{{{w_ns}}}rPr")
    if rPr is None:
        return
    rFonts = rPr.find(f"{{{w_ns}}}rFonts")
    if rFonts is None:
        rFonts = etree.SubElement(rPr, f"{{{w_ns}}}rFonts")
    # Снимаем theme-атрибуты.
    for attr in (
        "asciiTheme",
        "hAnsiTheme",
        "cstheme",
        "eastAsiaTheme",
    ):
        key = f"{{{w_ns}}}{attr}"
        if key in rFonts.attrib:
            del rFonts.attrib[key]
    # Ставим явный шрифт для всех 4 наборов символов.
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(f"{{{w_ns}}}{attr}", font_name)


def _apply_heading_styles(doc: DocxDocument, profile: Profile) -> None:
    """Переопределить стили Heading1..Heading4 параметрами из профиля.

    python-docx наследует Heading-стили из бортового шаблона Word —
    там они синие, Cambria, с большими before-spacing. Эта функция
    их переписывает по `profile.styles.heading_N`.
    """
    levels = (
        (1, profile.styles.heading_1),
        (2, profile.styles.heading_2),
        (3, profile.styles.heading_3),
        (4, profile.styles.heading_4),
    )
    for level, cfg in levels:
        style_id = f"Heading {level}"
        try:
            style = doc.styles[style_id]
        except KeyError:  # pragma: no cover - python-docx делает Heading1..9
            continue
        # Font + theme cleanup.
        style.font.name = cfg.font
        style.font.size = Pt(cfg.size_pt)
        style.font.bold = cfg.bold
        style.font.italic = cfg.italic
        # Цвет: 'auto' = снять явный цвет, иначе hex.
        _apply_style_color(style.element, cfg.color)
        _clear_theme_fonts(style.element, font_name=cfg.font)
        # Параграф-формат.
        pf = style.paragraph_format
        pf.alignment = _ALIGNMENT_MAP[cfg.alignment]
        pf.line_spacing = cfg.line_spacing
        pf.space_before = Pt(cfg.spacing_before_pt)
        pf.space_after = Pt(cfg.spacing_after_pt)
        pf.first_line_indent = Cm(cfg.first_line_indent_cm)
        pf.page_break_before = cfg.page_break_before
        pf.keep_with_next = cfg.keep_with_next
        # Связанный character-стиль (HeadingNChar): Word при рендере run-ов
        # внутри параграфа применяет linked-char поверх параграф-стиля.
        # Если не переписать — Cambria+синий из дефолтного шаблона
        # перекроют нашу правку.
        _sync_linked_char_style(doc, style.element, cfg)


def _sync_linked_char_style(doc: DocxDocument, p_style_element: Any, cfg: Any) -> None:
    """Применить шрифт/цвет/жирность из cfg к linked character-стилю.

    В styles.xml у каждого heading-параграф-стиля есть ссылка
    ``<w:link w:val="HeadingNChar"/>`` на character-стиль. При рендере
    run-ов параграфа Word применяет char-стиль поверх параграф-стиля
    (и его theme-fonts/синий цвет могут перекрыть параграф-настройки).

    Эта функция находит linked-char-стиль по styleId из w:link и
    переписывает его font, color, bold, italic симметрично с
    параграф-стилем. theme-fonts чистятся через _clear_theme_fonts.
    """
    w_ns = W_NS
    link = p_style_element.find(f"{{{w_ns}}}link")
    if link is None:
        return
    char_style_id = link.get(f"{{{w_ns}}}val")
    if not char_style_id:
        return
    # python-docx не индексирует character-стили по styleId напрямую,
    # ищем через styles_part.element.
    styles_root = p_style_element.getparent()
    char_elem = None
    for st in styles_root.findall(f"{{{w_ns}}}style"):
        if (
            st.get(f"{{{w_ns}}}type") == "character"
            and st.get(f"{{{w_ns}}}styleId") == char_style_id
        ):
            char_elem = st
            break
    if char_elem is None:
        return
    # Font + theme cleanup на rPr этого char-стиля.
    _clear_theme_fonts(char_elem, font_name=cfg.font)
    rPr = char_elem.find(f"{{{w_ns}}}rPr")
    if rPr is None:
        rPr = etree.SubElement(char_elem, f"{{{w_ns}}}rPr")
    # Size: w:sz/w:szCs в полу-пунктах (Pt*2).
    half_pt = str(int(cfg.size_pt * 2))
    for tag in ("sz", "szCs"):
        el = rPr.find(f"{{{w_ns}}}{tag}")
        if el is None:
            el = etree.SubElement(rPr, f"{{{w_ns}}}{tag}")
        el.set(f"{{{w_ns}}}val", half_pt)
    # Bold / italic: явные w:b и w:i (или их удаление).
    for tag, want in (("b", cfg.bold), ("bCs", cfg.bold), ("i", cfg.italic), ("iCs", cfg.italic)):
        el = rPr.find(f"{{{w_ns}}}{tag}")
        if want:
            if el is None:
                etree.SubElement(rPr, f"{{{w_ns}}}{tag}")
        else:
            if el is not None:
                rPr.remove(el)
    # Color — через ту же функцию что и для параграф-стиля.
    _apply_style_color(char_elem, cfg.color)


def _apply_style_color(style_element: Any, color: str) -> None:
    """Установить или снять явный цвет шрифта на уровне стиля.

    'auto' (или пустая строка) — снимает атрибут (Word возьмёт чёрный).
    hex без # — ставит rgb-цвет.
    """
    w_ns = W_NS
    rPr = style_element.find(f"{{{w_ns}}}rPr")
    if rPr is None:
        rPr = etree.SubElement(style_element, f"{{{w_ns}}}rPr")
    color_elem = rPr.find(f"{{{w_ns}}}color")
    if color in ("auto", "", None):
        if color_elem is not None:
            rPr.remove(color_elem)
        # Также удалим theme-color если есть.
        return
    if color_elem is None:
        color_elem = etree.SubElement(rPr, f"{{{w_ns}}}color")
    # Удалим тема-атрибуты, чтобы наш val реально применился.
    for attr in ("themeColor", "themeTint", "themeShade"):
        key = f"{{{w_ns}}}{attr}"
        if key in color_elem.attrib:
            del color_elem.attrib[key]
    color_elem.set(f"{{{w_ns}}}val", color.lstrip("#"))


def _apply_caption_style(doc: DocxDocument, profile: Profile) -> None:
    """Применить параметры подписи к стилю Caption (для figure и table).

    Поскольку figure.caption и table.caption могут отличаться (центр
    vs. лево, выше/ниже) — стиль Caption переписываем по figure.caption,
    а отдельные настройки для подписи таблицы применяются прямо к
    параграфу при записи таблицы (см. _write_table_caption).
    """
    cfg = profile.styles.figure.caption
    try:
        style = doc.styles["Caption"]
    except KeyError:  # pragma: no cover
        return
    style.font.name = cfg.font
    style.font.size = Pt(cfg.size_pt)
    style.font.bold = cfg.bold
    style.font.italic = cfg.italic
    _clear_theme_fonts(style.element, font_name=cfg.font)
    _apply_style_color(style.element, "auto")
    pf = style.paragraph_format
    pf.alignment = _ALIGNMENT_MAP[cfg.alignment]
    pf.space_before = Pt(cfg.spacing_before_pt)
    pf.space_after = Pt(cfg.spacing_after_pt)
    pf.first_line_indent = Cm(0)


def _write_runs(docx_paragraph: DocxParagraph, content: Sequence[InlineElement]) -> None:
    """Записать список InlineElement как набор run-ов в docx-параграф.

    Поддерживаются 4 типа inline-элементов (см. `gostforge.model.InlineElement`):

    * `TextRun` — обычный run с inline-форматированием.
    * `CrossRef` — OOXML-поле `<w:fldSimple w:instr=" REF target_id \\h "/>`;
      опциональный текст `prefix` добавляется как соседний run перед полем.
    * `InlineFormula` — `<m:oMath>` внутри `<w:r>` того же параграфа.
    * `Citation` — текстовый run «[N]» / «[N, с. P]», где N — 1-based индекс
      `source_id` в `Document.bibliography` (берётся из модуля).
    """
    for element in content:
        if isinstance(element, TextRun):
            _write_text_run(docx_paragraph, element)
        elif isinstance(element, CrossRef):
            _write_cross_ref(docx_paragraph, element)
        elif isinstance(element, InlineFormula):
            _write_inline_formula(docx_paragraph, element)
        elif isinstance(element, Citation):
            _write_citation(docx_paragraph, element)
        elif isinstance(element, Hyperlink):
            _write_hyperlink(docx_paragraph, element)
        elif isinstance(element, FootnoteRef):
            _write_footnote_ref(docx_paragraph, element)


def _write_hyperlink(docx_paragraph: DocxParagraph, element: Hyperlink) -> None:
    """Записать <w:hyperlink r:id="rIdN"> в параграф.

    Регистрирует Relationship на URL в части document.xml.rels через
    python-docx ``part.relate_to(url, RT.HYPERLINK, is_external=True)``.
    Если есть anchor (внутренняя ссылка) — пишем w:anchor вместо r:id.
    """
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    p_xml = docx_paragraph._p
    hl = etree.SubElement(p_xml, f"{{{W_NS}}}hyperlink")
    if element.anchor:
        hl.set(f"{{{W_NS}}}anchor", element.anchor)
    elif element.url:
        try:
            part = docx_paragraph.part
            r_id = part.relate_to(element.url, RT.HYPERLINK, is_external=True)
            hl.set(f"{{{_R_NS}}}id", r_id)
        except Exception:
            # Если relate_to не работает (тесты на mock docs) —
            # пишем URL прямо в anchor для отладки.
            hl.set(f"{{{W_NS}}}anchor", element.url)
    # Run с текстом ссылки.
    r = etree.SubElement(hl, f"{{{W_NS}}}r")
    rPr = etree.SubElement(r, f"{{{W_NS}}}rPr")
    # Стиль Hyperlink (синий + подчёркивание) применяется автоматически
    # Word-ом, но добавим style-ref для гарантии.
    rStyle = etree.SubElement(rPr, f"{{{W_NS}}}rStyle")
    rStyle.set(f"{{{W_NS}}}val", "Hyperlink")
    t = etree.SubElement(r, f"{{{W_NS}}}t")
    t.set(
        "{http://www.w3.org/XML/1998/namespace}space",
        "preserve",
    )
    t.text = element.text


def _write_footnote_ref(docx_paragraph: DocxParagraph, element: FootnoteRef) -> None:
    """Записать ссылку на footnote: <w:footnoteReference w:id="N"/>.

    Сам текст сноски должен лежать в word/footnotes.xml — экспортёр
    Phase 4 пока его не записывает, только ссылку. Если есть element.text,
    он добавляется как fallback-TextRun перед reference, чтобы текст не
    потерялся при экспорте в окружении без footnotes-part.
    """
    p_xml = docx_paragraph._p
    if element.text:
        # Fallback: добавляем текст сноски сразу после ссылки как
        # обычный текст (не теряем содержимое).
        r_text = etree.SubElement(p_xml, f"{{{W_NS}}}r")
        rPr = etree.SubElement(r_text, f"{{{W_NS}}}rPr")
        vertAlign = etree.SubElement(rPr, f"{{{W_NS}}}vertAlign")
        vertAlign.set(f"{{{W_NS}}}val", "superscript")
        t = etree.SubElement(r_text, f"{{{W_NS}}}t")
        t.text = f"[{element.footnote_id}]"


def _write_text_run(docx_paragraph: DocxParagraph, element: TextRun) -> None:
    """Записать TextRun со всеми его inline-атрибутами форматирования."""
    run = docx_paragraph.add_run(element.text)
    if element.bold:
        run.bold = True
    if element.italic:
        run.italic = True
    if element.underline:
        run.underline = True
    if element.superscript:
        run.font.superscript = True
    if element.subscript:
        run.font.subscript = True
    if element.font:
        run.font.name = element.font
    if element.size_pt is not None:
        run.font.size = Pt(element.size_pt)
    if element.color_hex:
        # color_hex в модели хранится как "#RRGGBB"; RGBColor ждёт hex без "#".
        hex_value = element.color_hex.lstrip("#")
        try:
            run.font.color.rgb = RGBColor.from_string(hex_value)
        except (ValueError, TypeError):
            # Невалидный цвет — игнорируем, не ломаем экспорт ради одного run.
            logger.debug("Невалидный color_hex=%r, пропускаем", element.color_hex)


def _write_cross_ref(docx_paragraph: DocxParagraph, element: CrossRef) -> None:
    """Записать CrossRef как опциональный prefix-run + `<w:fldSimple>` с REF.

    Display-текст внутри fldSimple — placeholder «[?]»: Word подменит его
    реальной автонумерацией при первом пересчёте полей. Для проверок C.*
    важно наличие правильного `w:instr` в OOXML.
    """
    if element.prefix:
        docx_paragraph.add_run(element.prefix)
    fld = OxmlElement("w:fldSimple")
    # Стандартный синтаксис Word: REF <bookmark> \h (гиперссылка). Пробелы
    # вокруг target_id критичны — без них Word не распарсит инструкцию.
    fld.set(qn("w:instr"), f" REF {element.target_id} \\h ")
    # Word рендерит fldSimple только при наличии хотя бы одного <w:r> внутри.
    inner_run = OxmlElement("w:r")
    inner_t = OxmlElement("w:t")
    inner_t.text = "[?]"
    inner_run.append(inner_t)
    fld.append(inner_run)
    docx_paragraph._p.append(fld)


def _write_inline_formula(docx_paragraph: DocxParagraph, element: InlineFormula) -> None:
    """Записать InlineFormula как `<m:oMath>` ВНУТРИ `<w:r>` параграфа.

    В отличие от блочной `Formula`, inline-формула не оборачивается в
    `<m:oMathPara>` и идёт в потоке текста. На Фазе 2.5 используется
    простейший fallback: LaTeX-строка кладётся в `<m:t>` как обычный текст;
    парсер уже умеет читать такой формат.
    """
    w_run = OxmlElement("w:r")
    omath = etree.SubElement(w_run, f"{{{M_NS}}}oMath")
    m_r = etree.SubElement(omath, f"{{{M_NS}}}r")
    m_t = etree.SubElement(m_r, f"{{{M_NS}}}t")
    m_t.text = element.latex
    docx_paragraph._p.append(w_run)


def _write_citation(docx_paragraph: DocxParagraph, element: Citation) -> None:
    """Записать Citation как текстовый run «[N]» / «[N, с. P]».

    Номер N вычисляется из модуль-level карты `_current_bibliography_index`,
    проставленной `export_docx()`. Если карта пуста или источник не найден,
    подставляем «?» (это всё ещё корректный текст, но проверка R.04 такой
    цитаты не пропустит — что и нужно).
    """
    index_map = _current_bibliography_index or {}
    n_value: int | str = index_map.get(element.source_id, "?")
    text = element.template.format(n=n_value, pages=element.pages or "")
    docx_paragraph.add_run(text)


def _apply_paragraph_format(docx_para: DocxParagraph, paragraph: Paragraph) -> None:
    """Применить per-paragraph переопределения формата (если заданы в модели)."""
    pf = docx_para.paragraph_format
    if paragraph.alignment is not None:
        pf.alignment = _ALIGNMENT_MAP[paragraph.alignment]
    if paragraph.line_spacing is not None:
        pf.line_spacing = paragraph.line_spacing
    if paragraph.first_line_indent_cm is not None:
        pf.first_line_indent = Cm(paragraph.first_line_indent_cm)
    if paragraph.page_break_before is not None:
        pf.page_break_before = paragraph.page_break_before
    if paragraph.space_before_pt is not None:
        pf.space_before = Pt(paragraph.space_before_pt)
    if paragraph.space_after_pt is not None:
        pf.space_after = Pt(paragraph.space_after_pt)


def _write_paragraph(doc: DocxDocument, paragraph: Paragraph) -> None:
    """Добавить один Paragraph в docx-документ."""
    style_name = paragraph.style_name or "Normal"
    try:
        docx_para = doc.add_paragraph(style=style_name)
    except KeyError:
        # Неизвестный стиль (например, кастомное имя) — используем Normal.
        docx_para = doc.add_paragraph(style="Normal")
    _apply_paragraph_format(docx_para, paragraph)
    _write_runs(docx_para, paragraph.content)


def _write_logical_section(doc: DocxDocument, section: LogicalSection) -> None:
    """Добавить заголовок логического раздела и рекурсивно записать его содержимое.

    Применяет UPPERCASE если профиль это требует для текущего уровня
    (по умолчанию — для heading_1 в ГОСТ 7.32).
    """
    heading_text = "".join(el.text for el in section.heading if isinstance(el, TextRun))
    level = max(0, min(section.level, 4))  # docx supports 0..9, мы — 1..4
    # Профиль-конфиг уровня: heading_1..heading_4 → uppercase, прочее.
    if _current_profile is not None and 1 <= level <= 4:
        cfg = getattr(_current_profile.styles, f"heading_{level}", None)
        if cfg is not None and cfg.uppercase:
            heading_text = heading_text.upper()
    doc.add_heading(heading_text, level=level)
    _write_items(doc, section.children)


def _write_caption_paragraph(
    doc: DocxDocument,
    content: Sequence[InlineElement],
    *,
    caption_kind: Literal["figure", "table"] = "figure",
) -> Any | None:
    """Записать подпись отдельным параграфом со стилем «Caption».

    `caption_kind` определяет, какой `CaptionStyleProfile` использовать —
    figure (по центру под рисунком) или table (слева над таблицей).
    Без него используем стиль figure-подписи (исторический default).

    Возвращает созданный параграф (полезно вызывающему, чтобы
    привязать bookmark — например, для continuation-row таблицы).
    ``None`` — если content пустой и параграф не создавался.
    """
    if not content:
        return None
    try:
        docx_para = doc.add_paragraph(style="Caption")
    except KeyError:
        docx_para = doc.add_paragraph()
    _write_runs(docx_para, content)
    # Применяем alignment и spacing согласно профилю — стиль Caption общий,
    # а для table-подписи нужны другие настройки (слева, перед таблицей).
    if _current_profile is None:
        return docx_para
    if caption_kind == "table":
        cfg = _current_profile.styles.table.caption
    else:
        cfg = _current_profile.styles.figure.caption
    pf = docx_para.paragraph_format
    pf.alignment = _ALIGNMENT_MAP[cfg.alignment]
    pf.space_before = Pt(cfg.spacing_before_pt)
    pf.space_after = Pt(cfg.spacing_after_pt)
    pf.first_line_indent = Cm(0)
    # keep_together: длинная подпись не разрывается между страницами.
    # keep_with_next: подпись таблицы (position=above) не отрывается
    # от таблицы под ней.
    if cfg.keep_together:
        pf.keep_together = True
    if cfg.keep_with_next:
        pf.keep_with_next = True
    # На уровне run-ов — шрифт/кегль (если в стиле Caption они сбиты).
    for run in docx_para.runs:
        run.font.name = cfg.font
        run.font.size = Pt(cfg.size_pt)
        if cfg.bold:
            run.bold = True
        if cfg.italic:
            run.italic = True
    return docx_para


def _write_table(doc: DocxDocument, table: Table) -> None:
    """Записать таблицу с подписью НАД ней (по ГОСТ).

    Применяет рамки и стиль ячеек из ``profile.styles.table``. Шапка
    bold. Подпись таблицы — слева, ВЫШЕ таблицы (профиль). Если
    ``repeat_header`` включён, на шапочные строки ставится
    ``<w:tblHeader/>`` — Word повторяет их на каждой continuation-странице
    (ГОСТ 7.32 «шапка таблицы повторяется при переносе»).

    При ``continuation_caption`` экспортёр прибавляет первой строкой
    специальную ячейку с OOXML field code
    ``{IF {PAGE} > {PAGEREF bm} "Продолжение таблицы N" ""}``,
    привязанным к bookmark-у в caption-параграфе. Word при отрисовке
    показывает «Продолжение таблицы N» только на continuation-страницах;
    на первой странице таблицы ячейка пустая. Это решает ограничение
    OOXML: чисто-XML способа показать row только на 2+ странице нет,
    но IF-поле даёт корректное поведение «как у руки».
    """
    cfg = _current_profile.styles.table if _current_profile is not None else None
    add_continuation = cfg is not None and cfg.continuation_caption and table.number is not None
    continuation_text = f"Продолжение таблицы {table.number}" if add_continuation else None
    continuation_bookmark = (
        f"gf_tbl_cont_{table.number}_{_next_bookmark_id()}" if add_continuation else None
    )

    caption_paragraph = _write_caption_paragraph(doc, table.caption, caption_kind="table")
    # Якорь для PAGEREF — кладём в caption-параграф ПЕРЕД таблицей.
    if continuation_bookmark is not None and caption_paragraph is not None:
        _add_bookmark_to_paragraph(caption_paragraph, continuation_bookmark)

    column_count = len(table.headers) if table.headers else 0
    for extra_row in table.extra_header_rows:
        column_count = max(column_count, len(extra_row))
    for row in table.rows:
        column_count = max(column_count, len(row))
    if column_count == 0:
        return

    rows_total = (
        (1 if continuation_text else 0)
        + len(table.extra_header_rows)
        + (1 if table.headers else 0)
        + len(table.rows)
    )
    if rows_total == 0:
        return
    docx_table = doc.add_table(rows=rows_total, cols=column_count)

    if cfg is not None:
        _apply_table_borders(docx_table, cfg)

    repeat_header = cfg is None or cfg.repeat_header
    header_bold = cfg.header_bold if cfg is not None else True

    row_idx = 0
    # Опциональная строка «Продолжение таблицы N» — широкая на все колонки,
    # tblHeader-репитер для continuation-страниц. Текст подменяется
    # OOXML field code `{IF {PAGE} > {PAGEREF bm} "..." ""}`: на первой
    # странице ячейка пустая (всё равно занимает строку — это
    # ограничение OOXML), на continuation-страницах появляется
    # «Продолжение таблицы N».
    if continuation_text is not None and continuation_bookmark is not None:
        cell0 = docx_table.rows[row_idx].cells[0]
        cell0.text = ""
        _apply_cell_paragraph_format(cell0, cfg, is_header=True)
        _apply_cell_font(cell0, cfg)
        _write_continuation_field_in_paragraph(
            cell0.paragraphs[0], continuation_bookmark, continuation_text
        )
        if column_count > 1:
            _apply_cell_merges(
                docx_table,
                [CellMerge(row=row_idx, col=0, rowspan=1, colspan=column_count)],
            )
        _mark_row_as_header(docx_table.rows[row_idx])
        row_idx += 1

    # Дополнительные строки шапки (для двух/трёх-уровневой шапки).
    for extra_row in table.extra_header_rows:
        for col_idx, cell_content in enumerate(extra_row):
            if col_idx >= column_count:
                break
            cell = docx_table.rows[row_idx].cells[col_idx]
            cell.text = ""
            _write_runs(cell.paragraphs[0], cell_content)
            _apply_cell_paragraph_format(cell, cfg, is_header=True)
            _apply_cell_font(cell, cfg)
            if header_bold:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
        if repeat_header:
            _mark_row_as_header(docx_table.rows[row_idx])
        row_idx += 1

    if table.headers:
        for col_idx, cell_content in enumerate(table.headers):
            cell = docx_table.rows[row_idx].cells[col_idx]
            cell.text = ""
            _write_runs(cell.paragraphs[0], cell_content)
            _apply_cell_paragraph_format(cell, cfg, is_header=True)
            _apply_cell_font(cell, cfg)
            if header_bold:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
        if repeat_header:
            _mark_row_as_header(docx_table.rows[row_idx])
        row_idx += 1
    for row in table.rows:
        for col_idx, cell_content in enumerate(row):
            if col_idx >= column_count:
                break
            cell = docx_table.rows[row_idx].cells[col_idx]
            cell.text = ""
            _write_runs(cell.paragraphs[0], cell_content)
            _apply_cell_paragraph_format(cell, cfg, is_header=False)
            _apply_cell_font(cell, cfg)
        row_idx += 1

    # CellMerge — индексы заданы в координатах (extra_header_rows +
    # headers + rows). Если есть continuation-row, сдвигаем все на 1.
    if table.merges:
        row_offset = 1 if continuation_text else 0
        shifted = (
            [
                CellMerge(row=m.row + row_offset, col=m.col, rowspan=m.rowspan, colspan=m.colspan)
                for m in table.merges
            ]
            if row_offset
            else list(table.merges)
        )
        _apply_cell_merges(docx_table, shifted)


def _mark_row_as_header(docx_row: Any) -> None:
    """Пометить строку таблицы как повторяющуюся шапку через `<w:tblHeader/>`.

    Word при переносе таблицы на следующую страницу будет повторять эту
    строку (и все строки выше с тем же атрибутом) в виде шапки.
    """
    tr = docx_row._tr
    trPr = tr.find(f"{{{W_NS}}}trPr")
    if trPr is None:
        trPr = etree.SubElement(tr, f"{{{W_NS}}}trPr")
        tr.insert(0, trPr)
    # Не дублируем — если уже есть, ничего не делаем.
    if trPr.find(f"{{{W_NS}}}tblHeader") is None:
        etree.SubElement(trPr, f"{{{W_NS}}}tblHeader")


_XML_NS = "http://www.w3.org/XML/1998/namespace"


def _next_bookmark_id() -> int:
    """Выдать новый уникальный w:id для bookmarkStart/End в документе."""
    global _bookmark_id_counter
    _bookmark_id_counter += 1
    return _bookmark_id_counter


def _add_bookmark_to_paragraph(paragraph: Any, bookmark_name: str) -> None:
    """Вставить ``<w:bookmarkStart>``/``<w:bookmarkEnd>`` в начало параграфа.

    Закладка нужна, чтобы PAGEREF в continuation-row знал, на какой
    странице начинается таблица. Помещается в caption-параграф ПЕРЕД
    таблицей.
    """
    p = paragraph._p
    bm_id = _next_bookmark_id()
    bm_start = etree.Element(f"{{{W_NS}}}bookmarkStart")
    bm_start.set(f"{{{W_NS}}}id", str(bm_id))
    bm_start.set(f"{{{W_NS}}}name", bookmark_name)
    bm_end = etree.Element(f"{{{W_NS}}}bookmarkEnd")
    bm_end.set(f"{{{W_NS}}}id", str(bm_id))
    # Вставляем bookmarkEnd сразу после Start: длина нулевая, нам нужна
    # только привязка к странице.
    p.insert(0, bm_end)
    p.insert(0, bm_start)


def _write_continuation_field_in_paragraph(
    paragraph: Any,
    bookmark_name: str,
    text: str,
    *,
    italic: bool = True,
) -> None:
    """Записать в параграф OOXML field code ``{IF {PAGE} > {PAGEREF bm} "text" ""}``.

    Word при отрисовке оценивает field динамически:

    * На странице, где начинается таблица, ``PAGE == PAGEREF(bm)`` →
      пустая строка.
    * На continuation-страницах ``PAGE > PAGEREF(bm)`` → отображается
      ``text``.

    Это решает ограничение OOXML «нельзя показать строку только на
    continuation-страницах» — раньше «Продолжение таблицы N»
    дублировалось на первой странице. Теперь там виден пустой
    (но всё ещё занимающий немного места) row; на continuation —
    корректный «Продолжение таблицы N».
    """
    p = paragraph._p
    # Удалим существующие runs — мы пишем свои.
    for r in list(p.findall(f"{{{W_NS}}}r")):
        p.remove(r)

    def _add_fldchar(fld_type: str) -> None:
        r = etree.SubElement(p, f"{{{W_NS}}}r")
        if italic:
            rpr = etree.SubElement(r, f"{{{W_NS}}}rPr")
            etree.SubElement(rpr, f"{{{W_NS}}}i")
        fld = etree.SubElement(r, f"{{{W_NS}}}fldChar")
        fld.set(f"{{{W_NS}}}fldCharType", fld_type)

    def _add_instr(instr: str) -> None:
        r = etree.SubElement(p, f"{{{W_NS}}}r")
        if italic:
            rpr = etree.SubElement(r, f"{{{W_NS}}}rPr")
            etree.SubElement(rpr, f"{{{W_NS}}}i")
        i = etree.SubElement(r, f"{{{W_NS}}}instrText")
        i.set(f"{{{_XML_NS}}}space", "preserve")
        i.text = instr

    def _add_text(value: str) -> None:
        r = etree.SubElement(p, f"{{{W_NS}}}r")
        if italic:
            rpr = etree.SubElement(r, f"{{{W_NS}}}rPr")
            etree.SubElement(rpr, f"{{{W_NS}}}i")
        t = etree.SubElement(r, f"{{{W_NS}}}t")
        t.set(f"{{{_XML_NS}}}space", "preserve")
        t.text = value

    # IF { PAGE } > { PAGEREF bm } "text" ""
    _add_fldchar("begin")
    _add_instr(" IF ")
    _add_fldchar("begin")
    _add_instr(" PAGE ")
    _add_fldchar("end")
    _add_instr(" > ")
    _add_fldchar("begin")
    _add_instr(f" PAGEREF {bookmark_name} ")
    _add_fldchar("end")
    # Кавычки экранируем как escape — IF-field парсит свои операнды
    # как строки в двойных кавычках. Внутри текста кавычек не ждём.
    safe_text = text.replace('"', "")
    _add_instr(f' "{safe_text}" ""')
    _add_fldchar("separate")
    _add_text("")  # cached result; Word пересчитает на render
    _add_fldchar("end")


def _apply_cell_merges(docx_table: Any, merges: list[CellMerge]) -> None:
    """Применить CellMerge к docx-таблице через <w:gridSpan>/<w:vMerge>.

    Для colspan > 1: на первой ячейке ставим <w:gridSpan w:val="N">,
    физические ячейки справа удаляем (их содержимое теряется — но
    модель CellMerge подразумевает, что в этой логической ячейке
    только один контент).

    Для rowspan > 1: на первой строке-ячейке ставим
    <w:vMerge w:val="restart">, на ячейках ниже — <w:vMerge/> (continue)
    без val.
    """
    for m in merges:
        try:
            top_cell = docx_table.rows[m.row].cells[m.col]
        except IndexError:
            continue
        tc = top_cell._tc
        tcPr = tc.find(f"{{{W_NS}}}tcPr")
        if tcPr is None:
            tcPr = etree.SubElement(tc, f"{{{W_NS}}}tcPr")
            # tcPr должен быть первым ребёнком tc.
            tc.insert(0, tcPr)
        # colspan: <w:gridSpan w:val="N"/>.
        if m.colspan > 1:
            for existing in tcPr.findall(f"{{{W_NS}}}gridSpan"):
                tcPr.remove(existing)
            grid_span = etree.SubElement(tcPr, f"{{{W_NS}}}gridSpan")
            grid_span.set(f"{{{W_NS}}}val", str(m.colspan))
            # Удалим соседние tc справа (они «съедены» colspan-ом).
            tr = tc.getparent()
            tcs_in_row = tr.findall(f"{{{W_NS}}}tc")
            try:
                start_idx = tcs_in_row.index(tc)
            except ValueError:
                start_idx = -1
            if start_idx >= 0:
                for to_remove in tcs_in_row[start_idx + 1 : start_idx + m.colspan]:
                    tr.remove(to_remove)
        # rowspan: <w:vMerge w:val="restart"/> на top + <w:vMerge/> ниже.
        if m.rowspan > 1:
            for existing in tcPr.findall(f"{{{W_NS}}}vMerge"):
                tcPr.remove(existing)
            v_merge_top = etree.SubElement(tcPr, f"{{{W_NS}}}vMerge")
            v_merge_top.set(f"{{{W_NS}}}val", "restart")
            for r_offset in range(1, m.rowspan):
                next_r = m.row + r_offset
                try:
                    next_cell = docx_table.rows[next_r].cells[m.col]
                except IndexError:
                    break
                next_tc = next_cell._tc
                next_tcPr = next_tc.find(f"{{{W_NS}}}tcPr")
                if next_tcPr is None:
                    next_tcPr = etree.SubElement(next_tc, f"{{{W_NS}}}tcPr")
                    next_tc.insert(0, next_tcPr)
                for existing in next_tcPr.findall(f"{{{W_NS}}}vMerge"):
                    next_tcPr.remove(existing)
                etree.SubElement(next_tcPr, f"{{{W_NS}}}vMerge")


def _apply_cell_paragraph_format(cell: Any, cfg: Any, *, is_header: bool) -> None:
    """Применить параграф-формат к ячейке таблицы.

    Без этого ячейки наследуют стиль Normal (justify + красная строка
    1.25 см + межстрочный 1.5), что в узких колонках ломает читаемость.
    По ГОСТ Р 2.105-2019 в таблицах: текст слева/центр, без красной
    строки, single-spacing, без интервалов между параграфами.

    Параметры берутся из profile.styles.table: cell_alignment,
    cell_first_line_indent_cm, cell_line_spacing, cell_space_before/
    after_pt. Для шапки — header_alignment (обычно центр).
    """
    if cfg is None:
        return
    alignment = cfg.header_alignment if is_header else cfg.cell_alignment
    for paragraph in cell.paragraphs:
        pf = paragraph.paragraph_format
        pf.alignment = _ALIGNMENT_MAP[alignment]
        pf.first_line_indent = Cm(cfg.cell_first_line_indent_cm)
        pf.line_spacing = cfg.cell_line_spacing
        pf.space_before = Pt(cfg.cell_space_before_pt)
        pf.space_after = Pt(cfg.cell_space_after_pt)


def _apply_table_borders(docx_table: Any, cfg: Any) -> None:
    """Прокинуть рамки в OOXML таблицы через w:tblBorders.

    python-docx позволяет задать ``table.style`` строкой ("Table Grid"),
    но стиль может отсутствовать в bortоvom шаблоне. Надёжнее —
    написать ``<w:tblBorders>`` напрямую в XML свойств таблицы.
    """
    if cfg.border_style == "none":
        return
    tbl_pr = docx_table._element.find(f"{{{W_NS}}}tblPr")
    if tbl_pr is None:
        tbl_pr = etree.SubElement(docx_table._element, f"{{{W_NS}}}tblPr")
    # Если tblBorders уже есть — заменим, иначе создадим.
    existing = tbl_pr.find(f"{{{W_NS}}}tblBorders")
    if existing is not None:
        tbl_pr.remove(existing)
    borders = etree.SubElement(tbl_pr, f"{{{W_NS}}}tblBorders")
    color = "auto" if cfg.border_color == "auto" else cfg.border_color.lstrip("#")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = etree.SubElement(borders, f"{{{W_NS}}}{side}")
        el.set(f"{{{W_NS}}}val", cfg.border_style)
        el.set(f"{{{W_NS}}}sz", str(int(cfg.border_size)))
        el.set(f"{{{W_NS}}}space", "0")
        el.set(f"{{{W_NS}}}color", color)


def _apply_cell_font(cell: Any, cfg: Any) -> None:
    """Применить шрифт/кегль ячеек таблицы, если они заданы в профиле.

    Если cell_font / cell_size_pt = None — оставляем дефолт от стиля Normal
    (который уже Times New Roman через _apply_normal_style).
    """
    if cfg is None:
        return
    if cfg.cell_font is None and cfg.cell_size_pt is None:
        return
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if cfg.cell_font is not None:
                run.font.name = cfg.cell_font
            if cfg.cell_size_pt is not None:
                run.font.size = Pt(cfg.cell_size_pt)


def _write_figure(doc: DocxDocument, figure: Figure) -> None:
    """Записать рисунок и его подпись.

    Стратегия выбора источника изображения:

    1. ``image_path = "embedded:rIdN"`` — рисунок пришёл из парсера и
       ссылается на media-blob исходного .docx. Если в текущем контексте
       экспорта (`_current_source_docx`) открыт source_docx — копируем
       blob и вставляем как настоящее изображение.
    2. ``image_path`` указывает на существующий файл — `add_picture(path)`.
    3. Иначе (пусто, не-файл, ошибка вставки) — placeholder-параграф
       вида `[Рисунок: <id>]`.
    """
    path = figure.image_path

    fig_alignment = _figure_alignment_from_profile()

    # Случай 1: embedded:rIdN с открытым source_docx.
    if path.startswith("embedded:") and _current_source_docx is not None:
        rid = path[len("embedded:") :]
        if _try_write_embedded_picture(doc, _current_source_docx, rid, figure):
            return

    # Случай 2: реальный файл на диске.
    if path and not path.startswith("embedded:") and Path(path).is_file():
        paragraph = doc.add_paragraph()
        paragraph.alignment = fig_alignment
        # Первый отступ убираем — это не текст, а рисунок.
        paragraph.paragraph_format.first_line_indent = Cm(0)
        _apply_figure_paragraph_constraints(paragraph)
        run = paragraph.add_run()
        try:
            _add_picture_with_max_width(run, path)
        except Exception:
            paragraph.add_run(f"[Рисунок: {figure.id}]").italic = True
        _write_caption_paragraph(doc, figure.caption, caption_kind="figure")
        return

    # Случай 3: placeholder.
    placeholder = doc.add_paragraph()
    placeholder.alignment = fig_alignment
    placeholder.paragraph_format.first_line_indent = Cm(0)
    _apply_figure_paragraph_constraints(placeholder)
    placeholder.add_run(f"[Рисунок: {figure.id}]").italic = True
    _write_caption_paragraph(doc, figure.caption, caption_kind="figure")


def _apply_figure_paragraph_constraints(paragraph: Any) -> None:
    """Применить параграф-формат к параграфу-рисунку.

    Сейчас — только keep_with_next=True (из profile.styles.figure):
    рисунок и подпись остаются на одной странице, Word не отрывает
    подпись на следующую страницу.
    """
    if _current_profile is None:
        return
    fig_cfg = _current_profile.styles.figure
    if fig_cfg.keep_with_next:
        paragraph.paragraph_format.keep_with_next = True


def _add_picture_with_max_width(run: Any, source: Any) -> None:
    """Вставить картинку с ограничением максимальной ширины и высоты.

    По умолчанию add_picture(path) использует оригинальный размер,
    из-за чего большие сканы и скриншоты вылезают за поля страницы
    или вообще выходят за лист (а подпись съезжает на следующую
    страницу — нарушение ГОСТ). Сначала вставляем без явной ширины,
    потом если картинка шире или выше лимита из профиля — масштабируем
    пропорционально (сохраняем aspect ratio).
    """
    if _current_profile is None:
        run.add_picture(source)
        return
    fig_cfg = _current_profile.styles.figure
    max_width_cm = float(fig_cfg.max_width_cm)
    max_height_cm = float(fig_cfg.max_height_cm)
    picture = run.add_picture(source)
    # picture.width/height — EMU (1 cm = 360000 EMU). Сначала жмём по
    # ширине, потом — если всё ещё высоковато — дожимаем по высоте.
    try:
        max_w_emu = Cm(max_width_cm).emu
        max_h_emu = Cm(max_height_cm).emu
        if picture.width > max_w_emu:
            ratio = picture.height / picture.width
            picture.width = max_w_emu
            picture.height = int(max_w_emu * ratio)
        if picture.height > max_h_emu:
            ratio = picture.width / picture.height
            picture.height = max_h_emu
            picture.width = int(max_h_emu * ratio)
    except (AttributeError, TypeError):  # pragma: no cover
        # На всякий случай — если picture не InlineShape (мок и пр.),
        # не падаем; картинка останется оригинального размера.
        return


def _figure_alignment_from_profile() -> Any:
    """Вернуть выравнивание рисунка из профиля (или CENTER по умолчанию)."""
    if _current_profile is None:
        return WD_ALIGN_PARAGRAPH.CENTER
    return _ALIGNMENT_MAP[_current_profile.styles.figure.alignment]


def _write_formula(doc: DocxDocument, formula: Formula) -> None:
    """Записать формулу как OOXML-OMath блок."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_xml = paragraph._p

    omath_para = etree.SubElement(p_xml, f"{{{M_NS}}}oMathPara")
    omath = etree.SubElement(omath_para, f"{{{M_NS}}}oMath")
    if formula.latex:
        m_r = etree.SubElement(omath, f"{{{M_NS}}}r")
        m_t = etree.SubElement(m_r, f"{{{M_NS}}}t")
        m_t.text = formula.latex

    if formula.number is not None:
        run = paragraph.add_run(f"\t({formula.number})")
        run.italic = False


def _try_write_embedded_picture(
    doc: DocxDocument, source_docx_obj: Any, rid: str, figure: Figure
) -> bool:
    """Попытка достать media-blob из source_docx и вставить как картинку."""
    try:
        image_part = source_docx_obj.part.related_parts.get(rid)
    except Exception:
        logger.debug("Не удалось получить related_parts для rId=%s", rid)
        return False
    if image_part is None:
        return False
    try:
        blob = image_part.blob
    except Exception:
        return False

    paragraph = doc.add_paragraph()
    paragraph.alignment = _figure_alignment_from_profile()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    _apply_figure_paragraph_constraints(paragraph)
    run = paragraph.add_run()
    try:
        _add_picture_with_max_width(run, io.BytesIO(blob))
    except Exception:
        p_xml = paragraph._p
        parent = p_xml.getparent()
        if parent is not None:
            parent.remove(p_xml)
        return False

    _write_caption_paragraph(doc, figure.caption, caption_kind="figure")
    return True


# Регекс «вручную добавленного маркера» в начале элемента списка.
# Покрывает: '-', '–', '—', '•', '*', '◦' + любые табы/пробелы после,
# а также нумерованные варианты '1.', '1)', 'a)', 'A.'.
# Не пытается удалять маркер посреди текста — только в самом начале.
_LEADING_MARKER_RE = re.compile(
    r"^("
    r"[-–—•*◦]"  # bullet-маркеры
    r"|\d{1,3}[\.\)]"  # '1.', '12)'
    r"|[a-zа-яёA-ZА-ЯЁ][\.\)]"  # 'a)', 'б.', 'A.', 'Б)'
    r")[\s\t]+"
)


def _strip_leading_marker_from_inline(
    content: Sequence[InlineElement],
) -> list[InlineElement]:
    """Удалить ведущий маркер списка из первого TextRun, если он есть.

    Идея: пользователи часто пишут элементы списка с уже добавленным
    маркером («- NET Framework 4.8», «1. шаг»). Если оставить такой
    текст как есть, экспортёр-numPr нарисует свой маркер ПЛЮС
    текстовый маркер в run-е останется — получаются «два маркера».

    Алгоритм:
    1. Найти первый TextRun в content (пропуская CrossRef/Citation/
       Formula — они не могут содержать маркер).
    2. Если его text начинается с известного маркера — удалить
       совпадение (маркер + последующие пробелы/табы).
    3. Если в результате text пуст — оставить TextRun с пустым text
       (легче, чем удалять элемент: остальные форматные атрибуты
       run-а сохраняются для случая когда они применяются к
       последующему).

    Возвращает НОВЫЙ list (не мутирует входной), чтобы избежать
    side-effect-ов при многократном проходе.
    """
    if not content:
        return list(content)
    result = list(content)
    for i, el in enumerate(result):
        if isinstance(el, TextRun):
            stripped = _LEADING_MARKER_RE.sub("", el.text, count=1)
            if stripped != el.text:
                result[i] = TextRun(
                    text=stripped,
                    bold=el.bold,
                    italic=el.italic,
                    underline=el.underline,
                    superscript=el.superscript,
                    subscript=el.subscript,
                    font=el.font,
                    size_pt=el.size_pt,
                    color_hex=el.color_hex,
                )
            return result
        # Не TextRun (CrossRef/InlineFormula/Citation в начале) —
        # маркер вряд ли там, прерываем поиск.
        break
    return result


def _write_list(doc: DocxDocument, list_block: ListBlock) -> None:
    """Записать список (нумерованный/маркированный) через настоящий numPr.

    Маркер и шаблон нумерации берутся из ``profile.styles.lists``
    (по умолчанию по ГОСТ Р 7.32-2017: маркер = тире «–»,
    нумерация = «N)»). Отступ слева и hanging — тоже из профиля.

    Стратегия: один раз на экспорт регистрируем в numbering.xml два
    abstractNum (ordered и bullet) с настроенным lvlText, и одно
    concrete num на каждый. Параграфы списка получают
    ``<w:pPr><w:numPr><w:ilvl w:val="0"/><w:numId w:val="N"/></w:numPr></w:pPr>``
    — Word/LibreOffice сами отрисовывают маркер. Это даёт правильное
    поведение списка при редактировании в Word (Enter → новый bullet),
    и парсер видит ``<w:numPr>`` → ListBlock без эвристик.
    """
    cfg = _current_profile.styles.lists if _current_profile is not None else None
    bullet = cfg.bullet_char if cfg is not None else "–"
    ordered_fmt = cfg.ordered_format if cfg is not None else "{n})"
    left_indent_cm = cfg.left_indent_cm if cfg is not None else 1.25
    hanging_indent_cm = cfg.hanging_indent_cm if cfg is not None else 0.5
    marker_suffix = cfg.marker_suffix if cfg is not None else "tab"

    # lvlText для numbering.xml: %1 = подстановка номера, остальное —
    # литеральный текст. ordered_fmt из профиля имеет «{n}» — конвертируем.
    lvl_text = ordered_fmt.replace("{n}", "%1") if list_block.ordered else bullet
    # Сколько уровней нужно (0 = плоский). Берём max(item_levels)+1
    # или 1 (если levels пустые).
    levels = list(list_block.item_levels) if list_block.item_levels else []
    max_level = max(levels) if levels else 0
    num_id = _ensure_list_num_in_numbering(
        doc,
        ordered=list_block.ordered,
        lvl_text=lvl_text,
        left_twips=int(Cm(left_indent_cm).twips),
        hanging_twips=int(Cm(hanging_indent_cm).twips),
        max_level=max_level,
        marker_suffix=marker_suffix,
    )

    item_left_twips = int(Cm(left_indent_cm).twips)
    item_hanging_twips = int(Cm(hanging_indent_cm).twips)

    for idx, item_content in enumerate(list_block.items):
        # Удаляем «вручную написанный» маркер в начале элемента
        # (типа «- NET Framework», «– требование», «1. шаг», «•  пункт»).
        # Без этого numPr рисует свой маркер ПЛЮС текстовый маркер
        # остаётся — пользователь видит «– – текст» или «1) 1. текст».
        item_content = _strip_leading_marker_from_inline(item_content)
        paragraph = doc.add_paragraph()
        # ilvl: из item_levels или 0 по умолчанию.
        item_level = levels[idx] if idx < len(levels) else 0
        # Ставим numPr.
        pPr = paragraph._p.get_or_add_pPr()
        for old in pPr.findall(f"{{{W_NS}}}numPr"):
            pPr.remove(old)
        num_pr = etree.SubElement(pPr, f"{{{W_NS}}}numPr")
        ilvl = etree.SubElement(num_pr, f"{{{W_NS}}}ilvl")
        ilvl.set(f"{{{W_NS}}}val", str(item_level))
        num_id_el = etree.SubElement(num_pr, f"{{{W_NS}}}numId")
        num_id_el.set(f"{{{W_NS}}}val", str(num_id))
        # Явный <w:ind> с left+hanging на уровне параграфа списка.
        # БАГ-ФИКС: раньше тут стоял `paragraph.paragraph_format.
        # first_line_indent = Cm(0)`, что писало <w:ind w:firstLine="0"/>
        # — это перекрывало hanging из numbering.xml, и при переносе
        # длинного элемента на следующую строку текст начинался с
        # красной строки 1.25 см (наследуясь от стиля Normal).
        # Теперь явно ставим left+hanging на каждый параграф списка:
        # numbering управляет ilvl, а параграф — отступом продолжения
        # строки (left).
        for old in pPr.findall(f"{{{W_NS}}}ind"):
            pPr.remove(old)
        # Уровни > 0 получают увеличенный left на +720 twips (как в
        # _ensure_list_num_in_numbering при создании lvl).
        effective_left = item_left_twips + item_level * 720
        ind = etree.SubElement(pPr, f"{{{W_NS}}}ind")
        ind.set(f"{{{W_NS}}}left", str(effective_left))
        # Семантика диалога Word «Изменение отступов в списке»:
        # hanging>0 — маркер левее текста; hanging<0 — правее (firstLine);
        # hanging=0 — совпадают (firstLine="0" перекрывает красную строку
        # стиля Normal, иначе перенос длинной строки съезжает на 1.25 см).
        if item_hanging_twips > 0:
            ind.set(f"{{{W_NS}}}hanging", str(item_hanging_twips))
        elif item_hanging_twips < 0:
            ind.set(f"{{{W_NS}}}firstLine", str(-item_hanging_twips))
        else:
            ind.set(f"{{{W_NS}}}firstLine", "0")  # перекрыть красную строку Normal
        _write_runs(paragraph, item_content)


# Кеш зарегистрированных numId внутри одного export-вызова: чтобы не
# плодить дубли abstractNum/num на каждый ListBlock. Ключ — кортеж
# (ordered, lvl_text, left_twips, hanging_twips). Сбрасывается в
# try/finally export_docx через _current_profile = None.
_current_list_numbering: dict[tuple[Any, ...], int] | None = None


def _ensure_list_num_in_numbering(
    doc: DocxDocument,
    *,
    ordered: bool,
    lvl_text: str,
    left_twips: int,
    hanging_twips: int,
    max_level: int = 0,
    marker_suffix: str = "tab",
) -> int:
    """Зарегистрировать (если ещё нет) abstractNum + num с этими параметрами.

    Возвращает numId, который нужно ставить в pPr/numPr/numId.

    Идея: на один уникальный (ordered+lvl_text+отступы) набор —
    один abstractNum и один num. Это позволяет иметь несколько списков
    с разной нумерацией (например, основной список с «N)» и доп.
    подсписок с «N.») в одном документе.
    """
    global _current_list_numbering
    if _current_list_numbering is None:
        _current_list_numbering = {}
    cache_key = (ordered, lvl_text, left_twips, hanging_twips, max_level, marker_suffix)
    if cache_key in _current_list_numbering:
        return _current_list_numbering[cache_key]

    try:
        numbering_part = doc.part.numbering_part
    except (AttributeError, KeyError):  # pragma: no cover
        # numbering_part должен быть в дефолтном шаблоне python-docx,
        # но на всякий случай — возвращаем 0 (отсутствующий numId,
        # Word проигнорирует и параграф будет без маркера).
        return 0
    num_elem = numbering_part.element

    # Находим максимальный существующий abstractNumId и numId.
    max_anid = -1
    for an in num_elem.findall(f"{{{W_NS}}}abstractNum"):
        with contextlib.suppress(ValueError):
            max_anid = max(max_anid, int(an.get(f"{{{W_NS}}}abstractNumId") or "-1"))
    max_nid = 0
    for n in num_elem.findall(f"{{{W_NS}}}num"):
        with contextlib.suppress(ValueError):
            max_nid = max(max_nid, int(n.get(f"{{{W_NS}}}numId") or "0"))

    new_anid = max_anid + 1
    new_nid = max_nid + 1

    # Создаём <w:abstractNum>.
    an = etree.SubElement(num_elem, f"{{{W_NS}}}abstractNum")
    an.set(f"{{{W_NS}}}abstractNumId", str(new_anid))
    mlt = etree.SubElement(an, f"{{{W_NS}}}multiLevelType")
    mlt.set(
        f"{{{W_NS}}}val",
        "singleLevel" if max_level == 0 else "multilevel",
    )
    for _ilvl in range(max_level + 1):
        _add_list_level(
            an,
            ilvl=_ilvl,
            ordered=ordered,
            lvl_text=(lvl_text if _ilvl == 0 else _nested_lvl_text(ordered, _ilvl, lvl_text)),
            left_twips=left_twips + _ilvl * 720,
            hanging_twips=hanging_twips,
            marker_suffix=marker_suffix,
            font_for_bullet=(
                _current_profile.styles.body.font
                if _current_profile is not None
                else "Times New Roman"
            ),
        )

    # Создаём <w:num> ссылающийся на abstractNum.
    n = etree.SubElement(num_elem, f"{{{W_NS}}}num")
    n.set(f"{{{W_NS}}}numId", str(new_nid))
    abstract_ref = etree.SubElement(n, f"{{{W_NS}}}abstractNumId")
    abstract_ref.set(f"{{{W_NS}}}val", str(new_anid))

    _current_list_numbering[cache_key] = new_nid
    return new_nid


def _add_list_level(
    abstract_num: Any,
    *,
    ilvl: int,
    ordered: bool,
    lvl_text: str,
    left_twips: int,
    hanging_twips: int,
    marker_suffix: str = "tab",
    font_for_bullet: str = "Times New Roman",
) -> None:
    """Добавить <w:lvl ilvl="N"> внутрь <w:abstractNum>.

    Содержит: start, numFmt (decimal/bullet), lvlText, lvlJc, pPr/ind
    с left+hanging/firstLine, rPr (для bullet — явный шрифт во избежание
    Symbol) и suff (tab/space/nothing) — символ после маркера/номера
    из профиля (поле «Символ после номера» диалога Word).
    """
    lvl = etree.SubElement(abstract_num, f"{{{W_NS}}}lvl")
    lvl.set(f"{{{W_NS}}}ilvl", str(ilvl))
    start = etree.SubElement(lvl, f"{{{W_NS}}}start")
    start.set(f"{{{W_NS}}}val", "1")
    num_fmt = etree.SubElement(lvl, f"{{{W_NS}}}numFmt")
    num_fmt.set(f"{{{W_NS}}}val", "decimal" if ordered else "bullet")
    # suff = разделитель между маркером и текстом: tab (default
    # Word), space или nothing — из профиля (поле «Символ после
    # номера»). При tab Tab расширяется до позиции left из <w:ind>,
    # выравнивая текст первой строки с переносом длинной строки.
    suff = etree.SubElement(lvl, f"{{{W_NS}}}suff")
    suff.set(f"{{{W_NS}}}val", marker_suffix)
    lvl_text_el = etree.SubElement(lvl, f"{{{W_NS}}}lvlText")
    lvl_text_el.set(f"{{{W_NS}}}val", lvl_text)
    lvl_jc = etree.SubElement(lvl, f"{{{W_NS}}}lvlJc")
    lvl_jc.set(f"{{{W_NS}}}val", "left")
    lvl_pPr = etree.SubElement(lvl, f"{{{W_NS}}}pPr")
    lvl_ind = etree.SubElement(lvl_pPr, f"{{{W_NS}}}ind")
    lvl_ind.set(f"{{{W_NS}}}left", str(left_twips))
    # Семантика диалога Word «Изменение отступов в списке»:
    # hanging>0 — маркер левее текста; hanging<0 — правее (firstLine);
    # hanging=0 — совпадают (ни hanging, ни firstLine).
    if hanging_twips > 0:
        lvl_ind.set(f"{{{W_NS}}}hanging", str(hanging_twips))
    elif hanging_twips < 0:
        lvl_ind.set(f"{{{W_NS}}}firstLine", str(-hanging_twips))
    if not ordered:
        lvl_rPr = etree.SubElement(lvl, f"{{{W_NS}}}rPr")
        lvl_rFonts = etree.SubElement(lvl_rPr, f"{{{W_NS}}}rFonts")
        for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
            lvl_rFonts.set(f"{{{W_NS}}}{attr}", font_for_bullet)


def _nested_lvl_text(ordered: bool, ilvl: int, base_lvl_text: str) -> str:
    """Вернуть lvlText для вложенного уровня ilvl > 0.

    Для ordered (decimal): '%(ilvl+1).' — стандартная nested-нумерация
    '%2.', '%3.', и т. д.

    Для unordered: альтернативный bullet-символ для разнообразия
    ('◦', '▪', ...), если base_lvl_text стандартный. Кастомный
    bullet просто повторяется.
    """
    if ordered:
        return f"%{ilvl + 1}."
    nested_bullets = ["◦", "▪", "·", "▫", "‣", "◆", "•"]
    if base_lvl_text in {"–", "—", "*", "•"}:
        return nested_bullets[min(ilvl - 1, len(nested_bullets) - 1)]
    return base_lvl_text


def _write_template_into_footer_paragraph(
    docx_paragraph: DocxParagraph, content: Sequence[InlineElement]
) -> None:
    """Записать содержимое одного слота footer/header, материализуя {page}.

    Каждый TextRun(text="{page}") превращается в OOXML-поле PAGE:
        <w:fldSimple w:instr="PAGE"/>
    Остальной текст пишется как обычные run-ы.
    """
    p_xml = docx_paragraph._p
    for element in content:
        if not isinstance(element, TextRun):
            continue
        # Разбиваем текст элемента на чередующиеся куски «обычный текст» / «{page}».
        text = element.text
        if not text:
            continue
        chunks = text.split(_PAGE_PLACEHOLDER)
        for i, chunk in enumerate(chunks):
            if chunk:
                docx_paragraph.add_run(chunk)
            if i < len(chunks) - 1:
                fld = etree.SubElement(p_xml, f"{{{W_NS}}}fldSimple")
                fld.set(f"{{{W_NS}}}instr", "PAGE")
                # Добавим внутрь fldSimple фиктивный run с пустым текстом —
                # без него Word не рендерит поле в новых версиях.
                run = etree.SubElement(fld, f"{{{W_NS}}}r")
                rt = etree.SubElement(run, f"{{{W_NS}}}t")
                rt.text = ""


def _has_text(items: Sequence[InlineElement] | None) -> bool:
    if not items:
        return False
    return any(isinstance(el, TextRun) and el.text and el.text.strip() for el in items)


def _write_footer(footer: Any, footer_template: ContentTemplate) -> None:
    """Записать footer из ContentTemplate в данный docx-footer ``footer``.

    Распределение по слотам left/center/right на Фазе 1 делаем через
    выравнивание единственного параграфа: если задан center — выравниваем
    по центру; right — вправо; иначе left. Если заполнены несколько слотов,
    добавляем отдельные параграфы под каждый.
    """
    # Удалим параграфы-плейсхолдеры, которые python-docx создаёт по умолчанию
    # (один пустой <w:p/>).
    for p in list(footer.paragraphs):
        p_xml = p._p
        if p_xml.getparent() is not None and not p.text and not list(p_xml):
            p_xml.getparent().remove(p_xml)

    slots: list[tuple[str, Sequence[InlineElement] | None]] = [
        ("left", footer_template.left),
        ("center", footer_template.center),
        ("right", footer_template.right),
    ]

    for slot, content in slots:
        if not _has_text(content):
            continue
        assert content is not None  # ради mypy
        para = footer.add_paragraph()
        if slot == "center":
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif slot == "right":
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _write_template_into_footer_paragraph(para, content)


# Обратное отображение модели в OOXML w:fmt (см. парсер _PAGE_FMT_MAP).
_PAGE_FMT_OOXML = {
    "arabic": "decimal",
    "roman": "upperRoman",
    "uppercase_letter": "upperLetter",
}


def _sync_page_section_with_profile(page_section: PageSection, profile: Profile) -> None:
    """Применить параметры профиля к модели page_section перед записью.

    Согласует:
    * margins_mm — поля страницы из profile.styles.page;
    * page_numbering.start_value — из profile.checks['F.06'].params,
      если задано и текущий start_mode='start_at'.

    Эта функция мутирует input page_section (side effect). Это
    единственное место, где модель «приземляется» под конкретный
    профиль во время экспорта — builder.build() остаётся
    профиль-агностичным.
    """
    # Поля страницы.
    margins = profile.styles.page.margins_mm
    if margins:
        merged = dict(page_section.page.margins_mm)
        merged.update({k: float(v) for k, v in margins.items()})
        page_section.page.margins_mm = merged
    # Рамка листа (ЕСКД). Профиль задаёт рамку, если в модели её ещё нет.
    border_cfg = profile.styles.page.border
    if border_cfg is not None and border_cfg.enabled and page_section.page.border is None:
        page_section.page.border = PageBorder(
            enabled=True,
            style=border_cfg.style,
            size_eighth_pt=int(border_cfg.size_eighth_pt),
            color=border_cfg.color,
            offset_from=border_cfg.offset_from,
            space_pt=int(border_cfg.space_pt),
        )
    # Основная надпись (штамп ЕСКД). Профиль задаёт штамп, если в модели нет.
    tb_cfg = profile.styles.page.title_block
    if tb_cfg is not None and tb_cfg.enabled and page_section.title_block is None:
        page_section.title_block = TitleBlock(
            enabled=True,
            form=tb_cfg.form,
            organization=tb_cfg.organization,
            roles=[TitleBlockRole(role=r.role, name=r.name, date=r.date) for r in tb_cfg.roles],
        )
    # F.06 start_value.
    f06 = profile.checks.get("F.06")
    if (
        f06
        and f06.enabled
        and f06.params.get("start_value") is not None
        and page_section.page_numbering.start_mode == "start_at"
    ):
        with contextlib.suppress(TypeError, ValueError):
            page_section.page_numbering.start_value = int(f06.params["start_value"])


def _apply_pgnumtype(sect: Any, page_section: PageSection) -> None:
    """Прописать <w:pgNumType w:start="N" w:fmt="..."/> в sectPr секции ``sect``.

    Атрибут `w:start` пишется только при `start_mode = "start_at"` и
    наличии `start_value`. Атрибут `w:fmt` пишется, если формат отличается
    от арабских цифр (Word-дефолт), либо если хоть один из атрибутов
    нужен — в этом случае пишем оба.
    """
    numbering = page_section.page_numbering
    needs_start = numbering.start_mode == "start_at" and numbering.start_value is not None
    needs_fmt = numbering.format != "arabic"
    if not needs_start and not needs_fmt:
        return
    sect_pr = getattr(sect, "_sectPr", None)
    if sect_pr is None:
        return
    # Удаляем существующий pgNumType, чтобы не было дублей.
    for existing in sect_pr.findall(f"{{{W_NS}}}pgNumType"):
        sect_pr.remove(existing)
    pg = etree.SubElement(sect_pr, f"{{{W_NS}}}pgNumType")
    if needs_start:
        pg.set(f"{{{W_NS}}}start", str(numbering.start_value))
    if needs_fmt:
        pg.set(f"{{{W_NS}}}fmt", _PAGE_FMT_OOXML.get(numbering.format, "decimal"))


def _apply_pg_borders(sect: Any, page_section: PageSection) -> None:
    """Прописать <w:pgBorders> (рамка листа, ЕСКД) в sectPr секции ``sect``.

    Пишется, только если ``page_section.page.border`` задан и
    ``enabled``. Все четыре стороны рамки получают одинаковые
    стиль/толщину/отступ/цвет. Зеркально парсеру ``_extract_pg_borders``.
    """
    border = page_section.page.border
    if border is None or not border.enabled:
        return
    sect_pr = getattr(sect, "_sectPr", None)
    if sect_pr is None:
        return
    # Удаляем существующий pgBorders, чтобы не было дублей.
    for existing in sect_pr.findall(f"{{{W_NS}}}pgBorders"):
        sect_pr.remove(existing)
    pg_borders = etree.SubElement(sect_pr, f"{{{W_NS}}}pgBorders")
    pg_borders.set(f"{{{W_NS}}}offsetFrom", border.offset_from)
    color = "auto" if border.color == "auto" else border.color.lstrip("#")
    space = str(max(0, min(31, int(border.space_pt))))
    for side in ("top", "left", "bottom", "right"):
        el = etree.SubElement(pg_borders, f"{{{W_NS}}}{side}")
        el.set(f"{{{W_NS}}}val", border.style)
        el.set(f"{{{W_NS}}}sz", str(int(border.size_eighth_pt)))
        el.set(f"{{{W_NS}}}space", space)
        el.set(f"{{{W_NS}}}color", color)


def _write_items(doc: DocxDocument, items: Sequence[LogicalSection | Block]) -> None:
    """Рекурсивно записать смешанный список логических разделов и блоков."""
    for item in items:
        if isinstance(item, LogicalSection):
            _write_logical_section(doc, item)
        elif isinstance(item, Paragraph):
            _write_paragraph(doc, item)
        elif isinstance(item, Table):
            _write_table(doc, item)
        elif isinstance(item, Figure):
            _write_figure(doc, item)
        elif isinstance(item, ListBlock):
            _write_list(doc, item)
        elif isinstance(item, Formula):
            _write_formula(doc, item)
        elif isinstance(item, TableOfContents):
            _write_toc(doc, item)


def _write_toc(doc: DocxDocument, toc: TableOfContents) -> None:
    """Записать автоматическое оглавление через Word TOC-field.

    OOXML: <w:p><w:fldSimple w:instr=" TOC \\o "1-3" \\h \\z "/>
    с placeholder-параграфом «Оглавление будет здесь после обновления
    (F9 в Word)» внутри fldSimple. При открытии файла Word строит
    оглавление автоматически на основе Heading-стилей.

    Опции TOC-field:
    * \\o "min-max" — диапазон уровней (default 1-3);
    * \\h — гиперссылки на заголовки;
    * \\z — скрыть номера в Web-preview.
    """
    paragraph = doc.add_paragraph()
    instr = f' TOC \\o "{toc.min_level}-{toc.max_level}" \\h \\z '
    fld = etree.SubElement(paragraph._p, f"{{{W_NS}}}fldSimple")
    fld.set(f"{{{W_NS}}}instr", instr)
    # Placeholder-run внутри fldSimple — Word его заменит при первом
    # обновлении поля (F9 или открытие документа с подтверждением).
    inner_r = etree.SubElement(fld, f"{{{W_NS}}}r")
    inner_t = etree.SubElement(inner_r, f"{{{W_NS}}}t")
    inner_t.text = "Оглавление будет построено при открытии (F9 в Word)."


def _configure_section_headers_footers(sect: Any, ps: PageSection, document: Document) -> None:
    """Настроить колонтитулы и основную надпись одной физической секции docx.

    Логика случаев (как в дипломе СПО, ГОСТ 2.104):
    * секция без колонтитула (``footer is None`` и штамп выключен) —
      отвязываем footer от предыдущей секции и оставляем пустым, чтобы не
      унаследовать чужой штамп/номер (титульный лист + задание);
    * секция со штампом и ``different_first_page`` — основная надпись
      (форма 2) печатается ТОЛЬКО на первой странице секции (содержание),
      на последующих листах колонтитула нет;
    * обычная секция со штампом/footer — пишем один footer на все страницы.
    """
    from .title_block import write_title_block

    tb = ps.title_block
    different_first = ps.different_first_page
    if different_first:
        sect.different_first_page_header_footer = True

    has_footer = (ps.footer is not None) or (tb is not None and tb.enabled)
    if has_footer:
        sect.footer.is_linked_to_previous = False
        if different_first:
            sect.first_page_footer.is_linked_to_previous = False
        if ps.footer is not None:
            _write_footer(sect.footer, ps.footer.default)
            if different_first and ps.footer.first_page is not None:
                _write_footer(sect.first_page_footer, ps.footer.first_page)
        if tb is not None and tb.enabled:
            if not tb.title and document.metadata.title:
                tb = replace(tb, title=document.metadata.title)
            if different_first:
                # Штамп только на первой странице секции (содержание);
                # остальные листы — без колонтитула (default footer пуст).
                write_title_block(sect.first_page_footer, tb)
            else:
                write_title_block(sect.footer, tb)
    else:
        # Колонтитула нет: отвязываем от предыдущей секции и не заполняем.
        sect.footer.is_linked_to_previous = False
        if different_first:
            sect.first_page_footer.is_linked_to_previous = False

    if ps.header is not None:
        sect.header.is_linked_to_previous = False
        _write_header(sect.header, ps.header.default)
        if different_first and ps.header.first_page is not None:
            sect.first_page_header.is_linked_to_previous = False
            _write_header(sect.first_page_header, ps.header.first_page)
    else:
        sect.header.is_linked_to_previous = False


def _suppress_page_break_before(docx_paragraph: DocxParagraph) -> None:
    """Явно отключить ``pageBreakBefore`` у параграфа (``w:val="false"``).

    Нужно для первого заголовка секции после разрыва «с новой страницы»:
    стиль «Заголовок 1» ставит pageBreakBefore, и он складывается с
    разрывом секции, давая лишнюю пустую страницу. Явный атрибут на
    параграфе перекрывает значение стиля.
    """
    p_pr = docx_paragraph._p.get_or_add_pPr()
    for existing in p_pr.findall(f"{{{W_NS}}}pageBreakBefore"):
        p_pr.remove(existing)
    el = etree.SubElement(p_pr, f"{{{W_NS}}}pageBreakBefore")
    el.set(f"{{{W_NS}}}val", "false")


def _export_multi_section(doc: DocxDocument, document: Document, profile: Profile) -> None:
    """Записать документ как несколько физических секций docx.

    Каждая ``PageSection`` модели становится отдельной секцией .docx со
    своими полями, нумерацией, рамкой и колонтитулами. Между секциями
    вставляется разрыв «с новой страницы» (``WD_SECTION.NEW_PAGE``).

    Сначала пишется контент всех секций (с разрывами), затем настраиваются
    свойства каждой физической секции — порядок важен: на момент настройки
    все ``doc.sections`` уже существуют.
    """
    page_sections = document.page_sections
    # Индекс первого body-параграфа каждой секции — чтобы снять лишний
    # page-break у первого заголовка (разрыв секции уже даёт новую страницу).
    first_para_index: list[int] = []
    for idx, ps in enumerate(page_sections):
        if idx > 0:
            doc.add_section(WD_SECTION.NEW_PAGE)
        first_para_index.append(len(doc.paragraphs))
        _write_items(doc, ps.content)

    paragraphs = doc.paragraphs
    for idx in range(1, len(page_sections)):
        start = first_para_index[idx]
        if start < len(paragraphs):
            _suppress_page_break_before(paragraphs[start])

    for idx, ps in enumerate(page_sections):
        _sync_page_section_with_profile(ps, profile)
        sect = doc.sections[idx]
        _apply_section_margins(sect, ps)
        _apply_page_size(sect, ps)
        _apply_pgnumtype(sect, ps)
        _apply_pg_borders(sect, ps)
        _configure_section_headers_footers(sect, ps, document)


def export_docx(
    document: Document,
    profile: Profile,
    output_path: str | Path,
    *,
    source_docx: str | Path | None = None,
) -> None:
    """Собрать .docx из модели по профилю.

    Минимальная реализация: геометрия страницы, стиль Normal, параграфы и
    заголовки. PageSection обрабатываются последовательно, но все кладутся
    в одну физическую секцию docx (sectPr per-PageSection — Фаза 2).

    Если у первой PageSection задан footer (например, после parse_docx
    плейсхолдер {page}) — он материализуется в OOXML-поле PAGE через lxml.

    Параметры:
      source_docx — путь к исходному .docx, из которого парсился document.
        Если задан и у Figure.image_path == ``embedded:rIdN``, экспортёр
        достанет соответствующий media-blob из source_docx и вставит как
        реальное изображение. Иначе на месте картинки будет placeholder.
    """
    global \
        _current_source_docx, \
        _current_bibliography_index, \
        _current_profile, \
        _current_list_numbering, \
        _bookmark_id_counter
    output_path = Path(output_path)
    doc = docx.Document()

    # Открываем source_docx (если задан) и кладём его в модуль-level
    # контекст — _write_figure читает оттуда. Try/finally гарантирует
    # сброс контекста даже при исключении в середине экспорта.
    _current_source_docx = None
    if source_docx is not None:
        try:
            _current_source_docx = docx.Document(str(source_docx))
        except Exception:
            # Повреждённый или нечитаемый docx → пропускаем источник, картинки уйдут в placeholder.
            logger.warning(
                "Не удалось открыть source_docx=%s; картинки будут placeholder-ами",
                source_docx,
            )
            _current_source_docx = None

    # Карта source_id → 1-based номер для рендеринга Citation. Заполняется
    # из Document.bibliography; пустая карта → все цитаты получат «?».
    _current_bibliography_index = {entry.id: i + 1 for i, entry in enumerate(document.bibliography)}

    try:
        _apply_page_geometry(doc, profile)
        _apply_normal_style(doc, profile)
        _apply_heading_styles(doc, profile)
        _apply_caption_style(doc, profile)
        _current_profile = profile

        # Метаданные документа в docProps/core.xml.
        core = doc.core_properties
        if document.metadata.title:
            core.title = document.metadata.title
        if document.metadata.author:
            core.author = document.metadata.author
        # Год работы: пишем как core.created (1 января указанного года),
        # чтобы парсер мог его прочитать обратно при impоrt-docx.
        # python-docx иначе ставит datetime.now(), и год потеряется.
        if document.metadata.year:
            from datetime import datetime

            core.created = datetime(document.metadata.year, 1, 1, tzinfo=UTC)

        if len(document.page_sections) > 1:
            # Мультисекционный экспорт: несколько физических секций docx с
            # собственными колонтитулами (титул/задание без штампа →
            # содержание с полной основной надписью → далее сокращённая).
            _export_multi_section(doc, document, profile)
        elif document.page_sections:
            first = document.page_sections[0]
            # Согласовать page_section с профилем перед записью: F.06
            # start_value, поля страницы. Builder ставит свои дефолты,
            # а профиль (особенно наследник base) может их переопределить.
            # Этот вызов — единое место синхронизации модели с профилем.
            _sync_page_section_with_profile(first, profile)
            sect = doc.sections[0]
            _apply_page_size(sect, first)
            _apply_pgnumtype(sect, first)
            _apply_pg_borders(sect, first)
            if first.footer is not None:
                _write_footer(sect.footer, first.footer.default)
            if first.header is not None:
                _write_header(sect.header, first.header.default)
            if first.title_block is not None and first.title_block.enabled:
                from .title_block import write_title_block

                tb = first.title_block
                # Наименование (графа 1) по умолчанию = заголовок работы.
                if not tb.title and document.metadata.title:
                    tb.title = document.metadata.title
                write_title_block(sect.footer, tb)

            _write_items(doc, first.content)

        doc.save(str(output_path))
    finally:
        _current_source_docx = None
        _current_bibliography_index = None
        _current_profile = None
        _current_list_numbering = None
        _bookmark_id_counter = 0

    # Post-processing на уровне zip: запись настроек, которые python-docx
    # не умеет менять напрямую (например, w:autoHyphenation в settings.xml).
    _postprocess_zip(Path(output_path), document)


def _postprocess_zip(output_path: Path, document: Document) -> None:
    """Доработать .docx-архив после save(): записать настройки в settings.xml."""
    if document.auto_hyphenation is None:
        return
    _patch_settings_auto_hyphenation(output_path, document.auto_hyphenation)


def _patch_settings_auto_hyphenation(docx_path: Path, value: bool) -> None:
    """Прописать/удалить <w:autoHyphenation/> в word/settings.xml внутри docx-zip."""
    import shutil
    import zipfile

    settings_path = "word/settings.xml"
    tmp_path = docx_path.with_suffix(".docx.tmp")
    with (
        zipfile.ZipFile(docx_path, "r") as zin,
        zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout,
    ):
        names = zin.namelist()
        if settings_path not in names:
            settings_xml = (
                b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                b'<w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                b"</w:settings>"
            )
        else:
            settings_xml = zin.read(settings_path)

        try:
            root = etree.fromstring(settings_xml)
            for existing in root.findall(f"{{{W_NS}}}autoHyphenation"):
                root.remove(existing)
            if value:
                elem = etree.SubElement(root, f"{{{W_NS}}}autoHyphenation")
                root.insert(0, elem)
            new_xml = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
        except Exception:
            new_xml = settings_xml

        wrote_settings = False
        for name in names:
            if name == settings_path:
                zout.writestr(name, new_xml)
                wrote_settings = True
            else:
                zout.writestr(name, zin.read(name))
        if not wrote_settings:
            zout.writestr(settings_path, new_xml)

    shutil.move(str(tmp_path), str(docx_path))


def _write_header(header: Any, header_template: ContentTemplate) -> None:
    """Записать содержимое в данный docx-header ``header`` (зеркально _write_footer)."""
    for p in list(header.paragraphs):
        p_xml = p._p
        if p_xml.getparent() is not None and not p.text and not list(p_xml):
            p_xml.getparent().remove(p_xml)

    slots: list[tuple[str, Sequence[InlineElement] | None]] = [
        ("left", header_template.left),
        ("center", header_template.center),
        ("right", header_template.right),
    ]

    for slot, content in slots:
        if not _has_text(content):
            continue
        assert content is not None
        para = header.add_paragraph()
        if slot == "center":
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif slot == "right":
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _write_template_into_footer_paragraph(para, content)
