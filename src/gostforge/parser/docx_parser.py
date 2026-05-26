"""Эвристический парсер .docx в модель документа.

Использует python-docx для базового разбора и lxml для нестандартных случаев
(колонтитулы, поля, ссылки). Работает эвристически: тип блока определяется
по стилю абзаца, паттернам в тексте, форматированию.

Поэтапная стратегия:
- Фаза 0: распознавать только параграфы и базовые свойства страницы
- Фаза 1: рисунки, таблицы, заголовки, нумерация
- Фаза 2: список литературы, перекрёстные ссылки, формулы
- Фаза 3: полный разбор колонтитулов, приложений
"""

# ruff: noqa: RUF001, RUF002, RUF003

from __future__ import annotations

import re
from collections.abc import Sequence
from pathlib import Path
from typing import TYPE_CHECKING, Any, Literal, cast

import docx  # type: ignore[import-not-found]
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-not-found]
from docx.table import Table as DocxTableCls  # type: ignore[import-not-found]
from docx.text.paragraph import Paragraph as DocxParagraphCls  # type: ignore[import-not-found]
from lxml import etree  # type: ignore[import-untyped]

from gostforge.model import (
    BibliographyEntry,
    Block,
    Citation,
    ContentTemplate,
    CrossRef,
    Document,
    DocumentMetadata,
    Figure,
    FootnoteRef,
    Formula,
    HeaderConfig,
    Hyperlink,
    InlineElement,
    InlineFormula,
    ListBlock,
    LogicalSection,
    PageGeometry,
    PageSection,
    Paragraph,
    ParagraphAlignment,
    Table,
    TextRun,
)

# Relationship-namespace для <w:hyperlink r:id="...">.
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

if TYPE_CHECKING:
    DocxDocument = Any
    DocxSection = Any
    DocxParagraph = Any
    DocxTable = Any
    DocxCell = Any


# Пространство имён OOXML wordprocessingml
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
# Пространство имён DrawingML и Office Relationships — нужны, чтобы достать
# rId из <a:blip r:embed="rIdN"/> внутри <w:drawing>.
A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NSMAP = {"w": W_NS, "m": M_NS, "a": A_NS, "r": R_NS}

# Шаблон номера формулы в скобках в конце текста параграфа: «(3)» или «(3.1)».
_FORMULA_NUMBER_RE = re.compile(r"\((\d+(?:\.\d+)?)\)\s*$")

# Соответствие WD_ALIGN_PARAGRAPH → строковая литералка модели
_ALIGN_MAP: dict[int, ParagraphAlignment] = {
    int(WD_ALIGN_PARAGRAPH.LEFT): "left",
    int(WD_ALIGN_PARAGRAPH.RIGHT): "right",
    int(WD_ALIGN_PARAGRAPH.CENTER): "center",
    int(WD_ALIGN_PARAGRAPH.JUSTIFY): "justify",
}

# Регэксп заголовков Word: "Heading 1", "Heading 2" и т.д.
_HEADING_RE = re.compile(r"^Heading\s+(\d+)$")

# Шаблоны подписей по тексту: «Рисунок 1 — ...», «Рис. 1 ...»; «Таблица 1 — ...».
_FIGURE_CAPTION_TEXT_RE = re.compile(r"^Рис(?:унок)?\.?\s+\d")
_TABLE_CAPTION_TEXT_RE = re.compile(r"^Таблица\s+\d")

# Множество имён Word-стилей подписи (нормализуется в lowercase).
_CAPTION_STYLE_NAMES = {"caption", "image caption", "figure caption", "table caption"}

# Заголовки, признаваемые началом раздела со списком литературы (нормализация
# к нижнему регистру и сжатию пробелов выполняется отдельно).
_BIBLIOGRAPHY_HEADINGS: set[str] = {
    "список использованных источников",
    "список литературы",
    "библиографический список",
    "список источников",
}

# Эвристики определения типа библиографической записи.
_BIB_URL_RE = re.compile(r"https?://", re.IGNORECASE)
_BIB_STANDARD_RE = re.compile(r"\bГОСТ\b")
_BIB_ARTICLE_RE = re.compile(r"(?:^|\s)//\s|\bЖурнал\b|журн\.|\bNo\.|№")
_BIB_THESIS_RE = re.compile(r"\bдис\.|\bдиссертация\b|\bавтореф\.", re.IGNORECASE)
_BIB_CONFERENCE_RE = re.compile(r"\bконференция\b|\bматериалы\b|\bсб\.\s*ст\.", re.IGNORECASE)
_BIB_LAW_RE = re.compile(r"\bзакон\b|\bфедер\.\s*закон\b|\bпостановление\b", re.IGNORECASE)

# Регэкспы для извлечения структурных полей библиографической записи
# по ГОСТ Р 7.0.100-2018. Поля заполняются опционально — отсутствие
# совпадения не считается ошибкой парсера (валидаторы R.* проверят сами).
_BIB_AUTHOR_RU_RE = re.compile(r"^[А-ЯЁ][а-яё]+\s[А-ЯЁ]\.\s?[А-ЯЁ]\.")
_BIB_AUTHOR_EN_RE = re.compile(r"^[A-Z][a-z]+,?\s[A-Z]\.\s?[A-Z]\.")
_BIB_YEAR_RE = re.compile(r"\b(?:19|20)\d{2}\b")
_BIB_URL_FULL_RE = re.compile(r"https?://\S+")
_BIB_DOI_RE = re.compile(r"10\.\d{4,9}/[-._;()/:A-Za-z0-9]+")
_BIB_ACCESS_DATE_RE = re.compile(r"дата\s+обращения:\s*(\d{1,2}\.\d{1,2}\.\d{4})", re.IGNORECASE)
# Эвристика места издания: «— Москва :», «— Санкт-Петербург :»,
# «— Нижний Новгород :». Принимаем составные через дефис или пробел.
_BIB_PLACE_RE = re.compile(r"—\s*([А-ЯЁ][а-яё]+(?:[ \-][А-ЯЁ][а-яё]+)?)\s*:")
_BIB_CYRILLIC_START_RE = re.compile(r"^[А-ЯЁа-яё]")
_BIB_LATIN_START_RE = re.compile(r"^[A-Za-z]")


def _parse_bibliography_fields(text: str) -> dict[str, str]:
    """Извлечь структурные поля из библиографической записи.

    Возвращает словарь только с найденными полями (опциональные —
    отсутствие паттерна означает, что поле просто не добавляется).

    Извлекаемые поля по ГОСТ Р 7.0.100-2018:
      - author: «Фамилия И. О.» в начале (русский или латинский вариант);
      - year: четырёхзначный год (1900-2099);
      - url: первая встреченная http(s)-ссылка;
      - doi: идентификатор DOI вида 10.NNNN/...;
      - access_date: «дата обращения: ДД.ММ.ГГГГ» (в любом регистре);
      - place: место издания между «— » и «:» (для книг);
      - language: «ru», если запись начинается с кириллицы; «en» — если с латиницы.
    """
    fields: dict[str, str] = {}

    author_ru = _BIB_AUTHOR_RU_RE.match(text)
    if author_ru is not None:
        fields["author"] = author_ru.group(0).strip()
    else:
        author_en = _BIB_AUTHOR_EN_RE.match(text)
        if author_en is not None:
            fields["author"] = author_en.group(0).strip()

    year_match = _BIB_YEAR_RE.search(text)
    if year_match is not None:
        fields["year"] = year_match.group(0)

    url_match = _BIB_URL_FULL_RE.search(text)
    if url_match is not None:
        # Отрезаем закрывающие скобки / точку в конце URL — они обычно
        # принадлежат окружающему тексту, а не самому адресу.
        url = url_match.group(0).rstrip(").,;")
        fields["url"] = url

    doi_match = _BIB_DOI_RE.search(text)
    if doi_match is not None:
        fields["doi"] = doi_match.group(0).rstrip(").,;")

    access_match = _BIB_ACCESS_DATE_RE.search(text)
    if access_match is not None:
        fields["access_date"] = access_match.group(1)

    place_match = _BIB_PLACE_RE.search(text)
    if place_match is not None:
        fields["place"] = place_match.group(1).strip()

    if _BIB_CYRILLIC_START_RE.match(text):
        fields["language"] = "ru"
    elif _BIB_LATIN_START_RE.match(text):
        fields["language"] = "en"

    return fields


def parse_docx(path: str | Path) -> Document:
    """Прочитать .docx и вернуть модель документа.

    Возвращает Document с одной PageSection (id="main"), куда уложены
    LogicalSection-ы (по заголовкам) и Paragraph-ы.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(path)

    docx_doc = cast("DocxDocument", docx.Document(str(path)))

    metadata = _extract_metadata(docx_doc, fallback_title=path.stem)
    page_section = _extract_page_section(docx_doc)
    _populate_content(docx_doc, page_section)
    bibliography = _extract_bibliography([page_section])
    auto_hyphenation = _extract_auto_hyphenation(docx_doc)
    _populate_image_dpi(docx_doc, page_section)
    # Citation post-processing должен выполняться ПОСЛЕ заполнения
    # bibliography, потому что эвристика «[N]» требует знать допустимый
    # диапазон N (валидный 1-based индекс в Document.bibliography).
    _annotate_citations([page_section], bibliography)
    # Группировка маркированных параграфов в ListBlock. Делаем ПОСЛЕ
    # извлечения библиографии, чтобы не сгруппировать references как
    # элементы списка (в bib-разделе каждая запись = отдельный параграф).
    _group_text_marker_lists(
        page_section, exclude_section_ids=_bibliography_section_ids([page_section])
    )
    comments = _extract_comments(docx_doc)
    # Сноски: извлекаем содержимое word/footnotes.xml и подставляем
    # текст в FootnoteRef-элементы внутри параграфов.
    footnotes = _extract_footnotes(docx_doc)
    if footnotes:
        _attach_footnote_text([page_section], footnotes)

    return Document(
        metadata=metadata,
        page_sections=[page_section],
        bibliography=bibliography,
        auto_hyphenation=auto_hyphenation,
        comments=comments,
    )


def _extract_footnotes(docx_doc: DocxDocument) -> dict[str, str]:
    """Извлечь содержимое сносок из word/footnotes.xml.

    Возвращает {footnote_id: text}. Идентификатор — целое число
    как строка. Содержимое — склейка всех ``<w:t>`` внутри
    ``<w:footnote w:id="N">``.

    Если footnotes-part отсутствует (документ без сносок) —
    пустой dict.
    """
    out: dict[str, str] = {}
    rels = getattr(docx_doc.part, "rels", None)
    if rels is None:
        return out
    footnotes_part = None
    target_reltype = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
    try:
        for rel in rels.values():
            if getattr(rel, "reltype", None) == target_reltype:
                footnotes_part = rel.target_part
                break
    except (etree.XMLSyntaxError, KeyError, AttributeError):
        return out
    if footnotes_part is None:
        return out

    blob = getattr(footnotes_part, "blob", None)
    if not blob:
        return out
    try:
        root = etree.fromstring(blob)
    except etree.XMLSyntaxError:
        return out

    for fn in root.findall(f"{{{W_NS}}}footnote"):
        fn_id = fn.get(f"{{{W_NS}}}id", "")
        if not fn_id:
            continue
        # Стандартные системные сноски (id=-1, 0): «separator» и
        # «continuationSeparator» — пропускаем, они не текстовые.
        fn_type = fn.get(f"{{{W_NS}}}type")
        if fn_type in {"separator", "continuationSeparator"}:
            continue
        texts: list[str] = []
        for t in fn.iter(f"{{{W_NS}}}t"):
            if t.text:
                texts.append(t.text)
        out[fn_id] = "".join(texts).strip()

    return out


def _attach_footnote_text(page_sections: list[PageSection], footnotes: dict[str, str]) -> None:
    """Заполнить FootnoteRef.text из карты footnotes.

    Парсер inline-элементов создаёт FootnoteRef только с id; полный
    текст подставляется здесь, чтобы избежать передачи footnotes-map
    через сигнатуры _build_paragraph и далее.
    """

    def walk(items: list) -> None:
        for item in items:
            if hasattr(item, "content"):  # Paragraph
                for el in item.content:
                    if isinstance(el, FootnoteRef):
                        el.text = footnotes.get(el.footnote_id, "")
            if hasattr(item, "children"):  # LogicalSection
                walk(item.children)

    for ps in page_sections:
        walk(ps.content)


def _extract_comments(docx_doc: DocxDocument) -> list[Any]:
    """Извлечь комментарии рецензента из word/comments.xml.

    Возвращает list[Comment]. Если comments-part отсутствует
    (документ без комментариев) — пустой список.

    python-docx не предоставляет высокоуровневого API для
    комментариев — обходимся прямым lxml-доступом к part-у.
    """
    from gostforge.model import Comment  # noqa: PLC0415

    out: list[Comment] = []
    # python-docx предоставляет part через related_parts; ищем
    # CommentsPart по reltype.
    rels = getattr(docx_doc.part, "rels", None)
    if rels is None:
        return out
    comments_part = None
    target_reltype = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
    try:
        for rel in rels.values():
            if getattr(rel, "reltype", None) == target_reltype:
                # rel.target_part может бросить XMLSyntaxError, если
                # comments.xml сломан — поймаем здесь и вернём [].
                comments_part = rel.target_part
                break
    except (etree.XMLSyntaxError, KeyError, AttributeError):
        return out
    if comments_part is None:
        return out

    blob = getattr(comments_part, "blob", None)
    if not blob:
        return out
    try:
        root = etree.fromstring(blob)
    except etree.XMLSyntaxError:
        return out

    for cm in root.findall(f"{{{W_NS}}}comment"):
        cid = cm.get(f"{{{W_NS}}}id", "")
        author = cm.get(f"{{{W_NS}}}author", "")
        date = cm.get(f"{{{W_NS}}}date", "")
        # Текст: собираем все <w:t> внутри.
        texts: list[str] = []
        for t in cm.iter(f"{{{W_NS}}}t"):
            if t.text:
                texts.append(t.text)
        text = "\n".join(texts).strip()
        # Без section_id — для MVP. Привязка к разделу через
        # commentRangeStart/End — отложена (требует обхода document.xml
        # с трекингом текущей секции).
        out.append(Comment(id=cid, author=author, date=date, text=text))

    return out


def _populate_image_dpi(docx_doc: DocxDocument, page_section: PageSection) -> None:
    """Для каждой Figure с image_path вида 'embedded:rIdN' извлечь DPI.

    Использует Pillow для определения разрешения media-file. Если Pillow
    недоступен, оставляет dpi=None. DPI вычисляется как min(x_dpi, y_dpi)
    из info['dpi'] или из физических размеров EMU vs пикселей (упрощённо —
    только info['dpi']).
    """
    try:
        from PIL import Image  # type: ignore[import-not-found]
    except ImportError:
        return

    def _iter_figures(items: list[LogicalSection | Block]) -> list[Figure]:
        out: list[Figure] = []
        for item in items:
            if isinstance(item, Figure):
                out.append(item)
            elif isinstance(item, LogicalSection):
                out.extend(_iter_figures(item.children))
        return out

    figures = _iter_figures(page_section.content)
    for fig in figures:
        if not fig.image_path.startswith("embedded:"):
            continue
        rid = fig.image_path[len("embedded:") :]
        try:
            image_part = docx_doc.part.related_parts.get(rid)
            if image_part is None:
                continue
            blob = image_part.blob
            with Image.open(__import__("io").BytesIO(blob)) as im:
                dpi_value = im.info.get("dpi")
                if dpi_value:
                    fig.dpi = int(min(dpi_value))
        except Exception:  # noqa: BLE001 — повреждённый media-file не должен валить парсер
            continue


def _extract_auto_hyphenation(docx_doc: DocxDocument) -> bool | None:
    """Прочитать `<w:autoHyphenation/>` из word/settings.xml.

    По OOXML autoHyphenation — toggle-элемент: его наличие включает
    автоматический перенос. Отсутствие = выключен. На Фазе 1
    интерпретируем строго: элемент найден → True, иначе False.
    """
    settings_part = None
    for rel in docx_doc.part.rels.values():
        if rel.reltype.endswith("/relationships/settings"):
            settings_part = rel.target_part
            break
    if settings_part is None:
        return None
    settings_xml = settings_part.blob
    try:
        root = etree.fromstring(settings_xml)
    except Exception:  # noqa: BLE001
        return None
    auto_hyph = root.find(f"{{{W_NS}}}autoHyphenation")
    if auto_hyph is None:
        return False
    val = auto_hyph.get(f"{{{W_NS}}}val")
    # По OOXML toggle: отсутствие val или val="1"/"true" → on; "0"/"false" → off.
    if val in {"0", "false"}:
        return False
    return True


# --- метаданные --------------------------------------------------------------


def _extract_metadata(docx_doc: DocxDocument, *, fallback_title: str) -> DocumentMetadata:
    """Извлечь title/author/year из docProps/core.xml."""
    core = docx_doc.core_properties
    title = (core.title or "").strip() or fallback_title
    author = (core.author or "").strip()
    year: int | None = None
    if core.created is not None:
        year = core.created.year
    return DocumentMetadata(title=title, author=author, year=year)


# --- секция страницы и колонтитулы -------------------------------------------


def _extract_page_section(docx_doc: DocxDocument) -> PageSection:
    """Построить PageSection из первой docx-секции (поля + footer)."""
    sect: DocxSection = docx_doc.sections[0]

    margins_mm = {
        "top": _length_to_mm(sect.top_margin),
        "right": _length_to_mm(sect.right_margin),
        "bottom": _length_to_mm(sect.bottom_margin),
        "left": _length_to_mm(sect.left_margin),
    }

    paper = _detect_paper_size(sect)
    orientation = _detect_orientation(sect)

    page_section = PageSection(
        id="main",
        name="Основная часть",
        type="main",
        page=PageGeometry(margins_mm=margins_mm, paper=paper, orientation=orientation),
    )

    # Стартовая страница нумерации: <w:pgNumType w:start="N"/> в sectPr.
    # Если атрибут задан — это эквивалентно start_mode = "start_at".
    start_value = _extract_page_number_start(sect)
    if start_value is not None:
        page_section.page_numbering.start_mode = "start_at"
        page_section.page_numbering.start_value = start_value

    # Формат нумерации: <w:pgNumType w:fmt="decimal|upperRoman|..."/>.
    fmt = _extract_page_number_format(sect)
    if fmt is not None:
        page_section.page_numbering.format = fmt

    footer_config = _extract_footer(sect)
    if footer_config is not None:
        page_section.footer = footer_config
        page_section.page_numbering.visible = True

    header_config = _extract_header(sect)
    if header_config is not None:
        page_section.header = header_config
        # Если PAGE-поле было в шапке, тоже считаем нумерацию видимой
        if _template_has_text(header_config.default, "{page}"):
            page_section.page_numbering.visible = True

    return page_section


def _template_has_text(template: ContentTemplate | None, needle: str) -> bool:
    """Проверка: есть ли в каком-либо слоте шаблона текст `needle`."""
    if template is None:
        return False
    for slot in (template.left, template.center, template.right):
        if slot is None:
            continue
        for el in slot:
            if isinstance(el, TextRun) and needle in el.text:
                return True
    return False


def _extract_page_number_start(sect: DocxSection) -> int | None:
    """Прочитать <w:pgNumType w:start="N"/> из sectPr секции и вернуть N.

    Возвращает None, если элемент отсутствует или значение не парсится.
    """
    sect_pr = getattr(sect, "_sectPr", None)
    if sect_pr is None:
        return None
    pg_num_type = sect_pr.find(f"{{{W_NS}}}pgNumType")
    if pg_num_type is None:
        return None
    start_attr = pg_num_type.get(f"{{{W_NS}}}start")
    if start_attr is None:
        return None
    try:
        return int(start_attr)
    except (TypeError, ValueError):
        return None


# OOXML поддерживает много форматов (decimalZero, hindiNumbers, и т.д.),
# но в модели мы фиксируем только три обиходных. Остальные значения
# отображаем в ближайший аналог (decimal → arabic, прочие — None).
_PAGE_FMT_MAP: dict[str, Literal["arabic", "roman", "uppercase_letter"]] = {
    "decimal": "arabic",
    "decimalZero": "arabic",
    "lowerRoman": "roman",
    "upperRoman": "roman",
    "lowerLetter": "uppercase_letter",
    "upperLetter": "uppercase_letter",
}


def _extract_page_number_format(
    sect: DocxSection,
) -> Literal["arabic", "roman", "uppercase_letter"] | None:
    """Прочитать <w:pgNumType w:fmt="..."/> и вернуть значение для модели.

    Возвращает None, если элемент или атрибут отсутствуют, либо если
    OOXML-значение не покрывается моделью.
    """
    sect_pr = getattr(sect, "_sectPr", None)
    if sect_pr is None:
        return None
    pg_num_type = sect_pr.find(f"{{{W_NS}}}pgNumType")
    if pg_num_type is None:
        return None
    fmt_attr = pg_num_type.get(f"{{{W_NS}}}fmt")
    if fmt_attr is None:
        return None
    return _PAGE_FMT_MAP.get(fmt_attr)


def _length_to_mm(length: object | None) -> float:
    """Перевести python-docx Length в миллиметры (округление до 0.1 мм)."""
    if length is None:
        # Дефолты Word подставит сам python-docx; если значение отсутствует
        # после чтения секции — возвращаем 0.0, чтобы валидатор увидел это.
        return 0.0
    mm = float(length.mm)  # type: ignore[attr-defined]
    return round(mm, 1)


# Известные размеры бумаги в мм. Сравнение с допуском ±2 мм.
_PAPER_SIZES_MM: dict[str, tuple[float, float]] = {
    "A4": (210.0, 297.0),
    "A3": (297.0, 420.0),
    "A5": (148.0, 210.0),
    "Letter": (215.9, 279.4),
    "Legal": (215.9, 355.6),
}


def _detect_paper_size(sect: DocxSection) -> str:
    """Определить формат бумаги по page_width/page_height.

    Сравнивает с известными размерами (A4, A3, ...) с допуском ±2 мм.
    Если ничего не подошло — возвращает «Unknown».
    """
    width_mm = _length_to_mm(sect.page_width)
    height_mm = _length_to_mm(sect.page_height)
    # Нормализуем (короткая сторона — первая), чтобы сравнение не зависело от ориентации
    short, long = sorted((width_mm, height_mm))
    for name, (w, h) in _PAPER_SIZES_MM.items():
        if abs(short - w) <= 2.0 and abs(long - h) <= 2.0:
            return name
    return "Unknown"


def _detect_orientation(sect: DocxSection) -> Literal["portrait", "landscape"]:
    """Определить ориентацию страницы по соотношению ширина/высота.

    Если width > height → landscape, иначе portrait. Атрибут sect.orientation
    бывает None для дефолтных значений, поэтому считаем по фактическим
    размерам через _length_to_mm.
    """
    width_mm = _length_to_mm(sect.page_width)
    height_mm = _length_to_mm(sect.page_height)
    if width_mm > height_mm:
        return "landscape"
    return "portrait"


def _extract_footer(sect: DocxSection) -> HeaderConfig | None:
    """Прочитать footer секции и собрать HeaderConfig, если есть поле PAGE."""
    return _extract_header_or_footer(sect.footer)


def _extract_header(sect: DocxSection) -> HeaderConfig | None:
    """Прочитать header секции и собрать HeaderConfig, если есть поле PAGE.

    Симметрично _extract_footer. Поле PAGE в header — это альтернатива
    положению номера страницы (top_*); проверка F.04 такие случаи
    обрабатывает.
    """
    return _extract_header_or_footer(sect.header)


def _extract_header_or_footer(container: Any) -> HeaderConfig | None:
    """Общая логика парсинга header/footer: ищем поле PAGE и кладём в подходящий слот."""
    if container is None:
        return None

    template = ContentTemplate()
    found = False

    for fp in container.paragraphs:
        if not _paragraph_has_page_field(fp):
            continue
        alignment = _alignment_to_literal(fp.paragraph_format.alignment)
        slot_runs = [TextRun(text="{page}")]
        # Распределение по слоту определяется выравниванием параграфа в
        # колонтитуле. None и justify трактуем как center — это типичный
        # случай отображения номера страницы.
        if alignment == "right":
            template.right = slot_runs
        elif alignment == "left":
            template.left = slot_runs
        else:
            template.center = slot_runs
        found = True
        # На Фазе 1 ограничиваемся одним маркером {page} в колонтитуле.
        break

    if not found:
        return None

    return HeaderConfig(default=template)


def _paragraph_has_page_field(fp: DocxParagraph) -> bool:
    """Найти в абзаце поле PAGE (как fldSimple, так и через fldChar+instrText)."""
    p_xml = fp._p

    # Вариант 1: <w:fldSimple w:instr="PAGE">
    for fld in p_xml.findall(f"{{{W_NS}}}fldSimple"):
        instr = fld.get(f"{{{W_NS}}}instr") or ""
        if _instr_is_page(instr):
            return True

    # Вариант 2: пара fldChar + instrText
    instr_chunks: list[str] = []
    in_field = False
    for elem in p_xml.iter():
        tag = etree.QName(elem.tag).localname
        if tag == "fldChar":
            ftype = elem.get(f"{{{W_NS}}}fldCharType")
            if ftype == "begin":
                in_field = True
                instr_chunks = []
            elif ftype == "end" and in_field:
                combined = "".join(instr_chunks)
                if _instr_is_page(combined):
                    return True
                in_field = False
        elif tag == "instrText" and in_field:
            instr_chunks.append(elem.text or "")
    return False


def _instr_is_page(instr: str) -> bool:
    """Проверить, что инструкция поля содержит PAGE (но не NUMPAGES)."""
    tokens = instr.strip().split()
    if not tokens:
        return False
    return tokens[0].upper() == "PAGE"


# --- контент: параграфы и заголовки ------------------------------------------


class _Counters:
    """Сквозные счётчики id блоков (общие на весь документ)."""

    def __init__(self) -> None:
        self.heading = 0
        self.paragraph = 0
        self.figure = 0
        self.table = 0
        self.list = 0
        self.formula = 0
        self.toc = 0


# Имена Word-стилей, обозначающих элемент списка (case-insensitive).
_LIST_STYLE_PREFIXES = ("list paragraph", "list number", "list bullet", "list")


def _paragraph_list_kind(dp: DocxParagraph) -> str | None:
    """Определить, является ли параграф элементом списка.

    Возвращает 'ordered' | 'bulleted' | None. Эвристика:
    1. Стиль абзаца 'List Number*' → ordered, 'List Bullet*' → bulleted.
    2. <w:numPr> → разрешаем numId через numbering.xml: ищем abstractNum
       по numId, читаем w:numFmt уровня 0 — 'bullet' → bulleted,
       'decimal'/'lowerLetter'/etc → ordered.
    3. Иначе None.
    """
    p_xml = dp._p
    num_pr = p_xml.find(f"{{{W_NS}}}pPr/{{{W_NS}}}numPr")
    style_name = (dp.style.name if dp.style is not None else "").lower()
    if style_name.startswith("list number"):
        return "ordered"
    if style_name.startswith("list bullet"):
        return "bulleted"
    if num_pr is not None:
        kind_via_num = _resolve_numpr_kind(dp, num_pr)
        if kind_via_num is not None:
            return kind_via_num
        # numbering.xml не доступен или не нашлось — берём ordered как
        # совместимое поведение (так делает дефолтный шаблон Word).
        return "ordered"
    return None


def _resolve_numpr_kind(dp: DocxParagraph, num_pr: Any) -> str | None:
    """Определить ordered/bulleted по numId через numbering.xml.

    Идём: numPr/numId → numbering_part: num[numId] → abstractNumId →
    abstractNum: lvl[ilvl=0]/numFmt. 'bullet' → bulleted, прочее → ordered.
    """
    w_ns = W_NS
    num_id_el = num_pr.find(f"{{{w_ns}}}numId")
    if num_id_el is None:
        return None
    num_id = num_id_el.get(f"{{{w_ns}}}val")
    if not num_id:
        return None
    try:
        numbering_part = dp.part.numbering_part
    except (AttributeError, KeyError):
        return None
    numbering_elem = getattr(numbering_part, "element", None)
    if numbering_elem is None:
        return None
    # Найти <w:num w:numId="N">.
    num_match = None
    for n in numbering_elem.findall(f"{{{w_ns}}}num"):
        if n.get(f"{{{w_ns}}}numId") == num_id:
            num_match = n
            break
    if num_match is None:
        return None
    abstract_ref = num_match.find(f"{{{w_ns}}}abstractNumId")
    if abstract_ref is None:
        return None
    abstract_id = abstract_ref.get(f"{{{w_ns}}}val")
    # Найти соответствующий abstractNum.
    for an in numbering_elem.findall(f"{{{w_ns}}}abstractNum"):
        if an.get(f"{{{w_ns}}}abstractNumId") == abstract_id:
            lvl = an.find(f"{{{w_ns}}}lvl")
            if lvl is None:
                return None
            num_fmt = lvl.find(f"{{{w_ns}}}numFmt")
            if num_fmt is None:
                return None
            fmt_val = num_fmt.get(f"{{{w_ns}}}val", "")
            if fmt_val == "bullet":
                return "bulleted"
            return "ordered"
    return None


def _paragraph_num_id(dp: DocxParagraph) -> str | None:
    """Вернуть numId параграфа, если у него есть <w:numPr>. Иначе None."""
    p_xml = dp._p
    num_pr = p_xml.find(f"{{{W_NS}}}pPr/{{{W_NS}}}numPr")
    if num_pr is None:
        return None
    num_id_el = num_pr.find(f"{{{W_NS}}}numId")
    if num_id_el is None:
        return None
    return num_id_el.get(f"{{{W_NS}}}val")


# Regex-маркеры элементов текстовых списков, написанных через builder/
# экспортёр без <w:numPr> (например, кастомный bullet_char из профиля).
# Группа 1 — собственно символ маркера, группа 2 — текст после маркера.
# Поддерживаются:
#   * «– <текст>» / «— <текст>»  — bullet (тире из ГОСТ),
#   * «• <текст>» / «* <текст>» / «◦ <текст>»  — bullet (другие маркеры),
#   * «1) <текст>» / «1. <текст>»  — ordered.
_BULLET_MARKERS = re.compile(r"^([–—•*◦])\s+(.+)$")
_ORDERED_MARKERS = re.compile(r"^(\d{1,3})[\.\)]\s+(.+)$")


def _group_text_marker_lists(page_section: PageSection, *, exclude_section_ids: set[str]) -> None:
    """Сгруппировать подряд идущие маркированные параграфы в ListBlock.

    Builder/экспортёр сейчас пишет списки как обычные параграфы с
    текстовым префиксом-маркером (из ``profile.styles.lists.bullet_char``,
    например «– »). Парсер уже корректно собирает их как Paragraph,
    но визуально это список — без этой постобработки round-trip
    «builder → export → parse → builder» теряет list-блоки.

    Алгоритм:
    1. Обходим рекурсивно содержимое page_section (внутри LogicalSection).
    2. Скользящим окном ищем последовательность Paragraph-ов, каждый
       из которых начинается с одинакового lead-маркера. Минимум 2
       параграфа подряд — иначе одиночный «– слово» в обычном тексте
       не превратится в список из одного элемента.
    3. Параграфы заменяются на один ``ListBlock`` с теми же items
       (без маркера в начале).

    Секции из ``exclude_section_ids`` пропускаются — обычно это
    библиографический раздел, где каждый параграф — отдельная
    запись, а не элемент списка.
    """
    from gostforge.model import LogicalSection, ListBlock, Paragraph  # noqa: PLC0415

    def process(items: list[Any], in_excluded_section: bool) -> list[Any]:
        result: list[Any] = []
        i = 0
        while i < len(items):
            item = items[i]
            if isinstance(item, LogicalSection):
                excluded = in_excluded_section or item.id in exclude_section_ids
                item.children = process(item.children, excluded)
                result.append(item)
                i += 1
                continue
            if in_excluded_section or not isinstance(item, Paragraph):
                result.append(item)
                i += 1
                continue
            # Пытаемся захватить run-серию параграфов с одинаковым
            # типом маркера (bulleted или ordered) и одним и тем же
            # символом маркера для bullet-варианта.
            group = _try_consume_marker_run(items, i)
            if group is None:
                result.append(item)
                i += 1
                continue
            list_block, consumed = group
            result.append(list_block)
            i += consumed
        return result

    for ps_attr in ("content",):
        items = getattr(page_section, ps_attr)
        setattr(
            page_section,
            ps_attr,
            process(items, in_excluded_section=False),
        )


def _try_consume_marker_run(items: list[Any], start: int) -> tuple[Any, int] | None:
    """Попытаться захватить серию маркированных параграфов с позиции ``start``.

    Возвращает (ListBlock, consumed) при успехе или None.
    """
    from gostforge.model import Block, ListBlock, Paragraph, TextRun  # noqa: PLC0415

    first = items[start]
    if not isinstance(first, Paragraph):
        return None
    first_text = _paragraph_plain_text(first)
    bullet_match = _BULLET_MARKERS.match(first_text)
    ordered_match = _ORDERED_MARKERS.match(first_text)
    if bullet_match:
        marker_char = bullet_match.group(1)
        ordered = False
    elif ordered_match:
        marker_char = None  # для ordered не привязываемся к конкретному номеру
        ordered = True
    else:
        return None

    items_texts: list[str] = []
    consumed = 0
    expected_num = 1 if ordered else None
    for idx in range(start, len(items)):
        item = items[idx]
        if not isinstance(item, Paragraph):
            break
        text = _paragraph_plain_text(item)
        if ordered:
            m = _ORDERED_MARKERS.match(text)
            if m is None:
                break
            try:
                num = int(m.group(1))
            except ValueError:
                break
            # Допускаем сбой нумерации — лишь бы шла серия. Строгая
            # проверка n+1 ломалась бы на «1) ..., 2) ..., 5) ...» —
            # это всё равно список.
            items_texts.append(m.group(2))
            consumed += 1
        else:
            m = _BULLET_MARKERS.match(text)
            if m is None or m.group(1) != marker_char:
                break
            items_texts.append(m.group(2))
            consumed += 1

    if consumed < 2:
        # Одиночный «– слово» не считаем списком (ложные срабатывания
        # на тире-сепаратор в обычном тексте).
        return None

    list_block = ListBlock(
        id=f"list-{first.id}",
        ordered=ordered,
        items=[[TextRun(text=t)] for t in items_texts],
    )
    return list_block, consumed


def _paragraph_plain_text(p: Any) -> str:
    """Склейка text всех TextRun параграфа (helper для list-grouping)."""
    from gostforge.model import TextRun  # noqa: PLC0415

    return "".join(el.text for el in p.content if isinstance(el, TextRun)).strip()


def _bibliography_section_ids(page_sections: list[PageSection]) -> set[str]:
    """Найти id всех LogicalSection-ов с заголовком «Список ...».

    Используется чтобы пропустить bib-секцию при группировке маркированных
    параграфов в ListBlock (в bib каждая запись = отдельный параграф,
    а не элемент списка).
    """
    from gostforge.model import LogicalSection, TextRun  # noqa: PLC0415

    aliases = {
        "список использованных источников",
        "список литературы",
        "литература",
        "список источников",
        "библиографический список",
        "references",
    }

    result: set[str] = set()

    def walk(items: list[Any]) -> None:
        for it in items:
            if isinstance(it, LogicalSection):
                heading = (
                    "".join(el.text for el in it.heading if isinstance(el, TextRun)).strip().lower()
                )
                if heading in aliases:
                    result.add(it.id)
                walk(it.children)

    for ps in page_sections:
        walk(ps.content)
    return result


def _populate_content(docx_doc: DocxDocument, page_section: PageSection) -> None:
    """Пройти по телу документа в порядке появления и наполнить PageSection.

    Алгоритм Фазы 1:
    1. Итерируемся по `body.iterchildren()` и обрабатываем каждый элемент
       по тегу: `<w:p>` — параграф/заголовок/рисунок, `<w:tbl>` — таблица.
       `<w:sectPr>` и прочие служебные элементы — пропускаем.
    2. Заголовок (Heading N) открывает новую LogicalSection.
    3. Параграфы и таблицы кладутся либо в текущую LogicalSection, либо
       прямо в `page_section.content` (если ни одного раздела ещё не было).
    4. После того как линейная последовательность собрана, делается
       единый проход «склейки подписей»: рисункам подпись прикрепляется
       снизу, таблицам — сверху; параграфы-подписи удаляются.
    """
    current_section: LogicalSection | None = None
    # Стек открытых секций для построения иерархии (глава → подраздел
    # → пункт). См. логику ниже при обработке заголовков.
    section_stack: list[LogicalSection] = []
    counters = _Counters()
    body = docx_doc.element.body

    # Буфер для группировки последовательных list-параграфов в один ListBlock.
    # Элемент: (kind, num_id, content). num_id используется для разделения
    # разных списков подряд (например, нумерованный + маркированный
    # сразу — не сольются в один ListBlock).
    list_buffer: list[tuple[str, str | None, list[InlineElement]]] = []

    def flush_list_buffer() -> None:
        """Если в буфере накопились list-параграфы — сделать из них ListBlock.

        Дополнительно разрезает буфер по сменам (kind, num_id), чтобы
        два подряд идущих списка разных типов (например, ordered и
        bulleted) не слились в один.
        """
        if not list_buffer:
            return
        # Группируем по (kind, num_id).
        groups: list[tuple[str, list[list[InlineElement]]]] = []
        cur_key: tuple[str, str | None] | None = None
        for kind, num_id, content in list_buffer:
            key = (kind, num_id)
            if cur_key != key:
                groups.append((kind, []))
                cur_key = key
            groups[-1][1].append(content)
        for kind, items in groups:
            counters.list += 1
            list_block = ListBlock(
                id=f"list-{counters.list}",
                ordered=(kind == "ordered"),
                items=items,
            )
            _append_block(list_block, page_section, current_section)
        list_buffer.clear()

    for child in body.iterchildren():
        tag = etree.QName(child.tag).localname
        if tag == "p":
            dp = DocxParagraphCls(child, docx_doc)
            heading_level = _heading_level(dp)
            if heading_level is not None:
                flush_list_buffer()
                counters.heading += 1
                # Собираем heading через ту же логику inline-разбора, что и
                # обычный параграф: парсер пройдётся по <w:r>-ам, выполнит
                # style-cascade (font/size/bold/italic от стиля Heading{N})
                # и заполнит TextRun реальными атрибутами форматирования.
                # Без этого H.01/H.02 «слепы» к стиле-уровневым настройкам
                # синего Cambria из дефолтного шаблона Word.
                heading_paragraph = _build_paragraph(dp, idx=counters.heading)
                heading_content = heading_paragraph.content or [TextRun(text=dp.text)]
                section = LogicalSection(
                    id=f"sec-{counters.heading}",
                    level=heading_level,
                    heading=heading_content,
                )
                # Поддерживаем стек открытых секций для правильной
                # иерархии. Новая секция уровня L закрывает (pop) все
                # открытые секции с level>=L и становится child-ом
                # текущей вершины стека (или попадает на top-level,
                # если стек пуст).
                # Без этого S.07/H.06 ложно срабатывали на главах:
                # подразделы лежали на top-level рядом с главой, а у
                # главы children=[].
                while section_stack and section_stack[-1].level >= heading_level:
                    section_stack.pop()
                if section_stack:
                    section_stack[-1].children.append(section)
                else:
                    page_section.content.append(section)
                section_stack.append(section)
                current_section = section
                continue

            # Распознаём элемент списка ещё до построения Paragraph
            list_kind = _paragraph_list_kind(dp)
            if list_kind is not None:
                # Собираем inline-контент параграфа как InlineElement-список
                runs: list[InlineElement] = [TextRun(text=run.text) for run in dp.runs if run.text]
                if not runs and dp.text:
                    runs = [TextRun(text=dp.text)]
                if runs:
                    num_id = _paragraph_num_id(dp)
                    list_buffer.append((list_kind, num_id, runs))
                continue

            # Обычный параграф или figure — сначала закрываем накопленный список
            flush_list_buffer()
            block = _block_from_paragraph(dp, counters)
            _append_block(block, page_section, current_section)
        elif tag == "tbl":
            flush_list_buffer()
            counters.table += 1
            table = _build_table(DocxTableCls(child, docx_doc), idx=counters.table)
            _append_block(table, page_section, current_section)
        # sectPr и прочие — игнорируем.

    # Не забыть закрыть финальный список (если документ закончился им)
    flush_list_buffer()

    # Склейка подписей: для content страницы и для каждой LogicalSection.
    _glue_captions(page_section.content)
    for section in _iter_logical_sections(page_section.content):
        _glue_captions(section.children)


def _append_block(
    block: Block,
    page_section: PageSection,
    current_section: LogicalSection | None,
) -> None:
    """Положить блок в текущую LogicalSection либо прямо в content страницы."""
    if current_section is not None:
        current_section.children.append(block)
    else:
        page_section.content.append(block)


def _iter_logical_sections(
    items: list[LogicalSection | Block],
) -> list[LogicalSection]:
    """Рекурсивно собрать все LogicalSection (всех уровней) — для склейки."""
    result: list[LogicalSection] = []
    for item in items:
        if isinstance(item, LogicalSection):
            result.append(item)
            result.extend(_iter_logical_sections(item.children))
    return result


def _block_from_paragraph(dp: DocxParagraph, counters: _Counters) -> Block:
    """Конвертировать <w:p> в Figure / Formula / Paragraph.

    Приоритет распознавания:
    1. <w:txbxContent> (текстбокс внутри drawing) → Paragraph с текстом.
       Часто титульный лист или вынос помещают в текстбокс — без этой
       ветки текст потерялся бы (drawing считался бы пустым рисунком).
    2. <w:drawing> → Figure (рисунок).
    3. <m:oMathPara> → блочная Formula (формула как отдельный блок).
       Голый <m:oMath> без <m:oMathPara> считается inline-формулой
       и обрабатывается в _build_paragraph.
    4. Иначе — обычный Paragraph (возможно с InlineFormula в content).
    """
    # Текстбокс: извлекаем текст из <w:txbxContent>, если он непустой.
    # Это важнее, чем считать параграф рисунком — текст не должен
    # теряться.
    txbx_text = _extract_textbox_text(dp._p)
    if txbx_text:
        counters.paragraph += 1
        return Paragraph(
            id=f"p-{counters.paragraph}",
            content=[TextRun(text=txbx_text)],
            style_name=dp.style.name if dp.style is not None else None,
        )
    drawings = dp._p.findall(f".//{{{W_NS}}}drawing")
    if drawings:
        counters.figure += 1
        image_path = _extract_drawing_rid(drawings[0])
        alignment = _alignment_to_literal(dp.paragraph_format.alignment)
        return Figure(
            id=f"fig-{counters.figure}",
            image_path=image_path,
            caption=[],
            alignment=alignment,
        )
    # TOC-field: <w:fldSimple w:instr=" TOC \o \"1-3\" \h \z "/>.
    # При парсинге чужого .docx распознаём TOC-поле и возвращаем
    # TableOfContents-блок, чтобы он переживал round-trip
    # docx → state → docx без замены на текстовый плейсхолдер.
    for fld in dp._p.findall(f".//{{{W_NS}}}fldSimple"):
        instr = (fld.get(f"{{{W_NS}}}instr") or "").strip()
        if instr.startswith("TOC"):
            counters.toc = getattr(counters, "toc", 0) + 1
            from gostforge.model import TableOfContents  # noqa: PLC0415

            min_lvl, max_lvl = 1, 3
            m = re.search(r'\\o\s+"(\d+)-(\d+)"', instr)
            if m:
                try:
                    min_lvl = int(m.group(1))
                    max_lvl = int(m.group(2))
                except ValueError:
                    pass
            return TableOfContents(
                id=f"toc-{counters.toc}",
                min_level=min_lvl,
                max_level=max_lvl,
            )
    # Блочная формула: <m:oMathPara> существует. Тогда весь параграф
    # становится Formula. Внутри <m:oMathPara> может быть несколько
    # <m:oMath>; склеиваем их текст.
    omml_text = _extract_block_omml_text(dp._p)
    if omml_text is not None:
        counters.formula += 1
        return Formula(
            id=f"formula-{counters.formula}",
            latex=omml_text,
            number=_extract_formula_number(dp),
        )
    counters.paragraph += 1
    return _build_paragraph(dp, idx=counters.paragraph)


def _extract_textbox_text(p_elem: Any) -> str:
    """Извлечь текст из <w:txbxContent> внутри параграфа.

    Текстбоксы в OOXML лежат как
    ``<w:drawing>…<wps:txbx><w:txbxContent><w:p>…</w:p></w:txbxContent>``
    (DrawingML) или ``<w:pict>…<v:textbox><w:txbxContent>`` (VML).
    Извлекаем склейку всех <w:t> внутри любого <w:txbxContent>.

    Возвращает пустую строку, если текстбокса нет или он пуст.
    """
    parts: list[str] = []
    # txbxContent в любом namespace (DrawingML wps или VML) — ищем по
    # localname через iter с проверкой тега.
    for el in p_elem.iter():
        if etree.QName(el.tag).localname == "txbxContent":
            for t in el.iter(f"{{{W_NS}}}t"):
                if t.text:
                    parts.append(t.text)
    return "".join(parts).strip()


def _extract_drawing_rid(drawing_elem: Any) -> str:
    """Извлечь relationship-id изображения из <w:drawing>.

    Внутри <w:drawing> есть <a:blip r:embed="rIdN"/> (или r:link для
    внешних ссылок). Возвращает идентификатор вида 'embedded:rIdN'.
    Если не нашли — пустую строку (тогда экспортёр напишет placeholder).
    """
    for blip in drawing_elem.iter(f"{{{A_NS}}}blip"):
        embed = blip.get(f"{{{R_NS}}}embed")
        if embed:
            return f"embedded:{embed}"
        link = blip.get(f"{{{R_NS}}}link")
        if link:
            return f"linked:{link}"
    return ""


def _extract_block_omml_text(p_elem: Any) -> str | None:
    """Склейка `<m:t>` из `<m:oMath>`, лежащих внутри `<m:oMathPara>`.

    На Фазе 2.5 различаем блочные и inline-формулы:
      - `<m:oMath>` внутри `<m:oMathPara>` → блочная формула, весь
        параграф конвертируется в Formula.
      - `<m:oMath>` напрямую как ребёнок `<w:p>` (без `<m:oMathPara>`) →
        inline-формула, обрабатывается в _build_paragraph.

    Возвращает None, если блочной формулы в параграфе нет.
    Это не настоящий LaTeX — берём только видимый математический
    текст (операнды и константы), как и в предыдущих фазах.
    """
    omath_para_elements = p_elem.findall(f".//{{{M_NS}}}oMathPara")
    if not omath_para_elements:
        return None
    parts: list[str] = []
    for omath_para in omath_para_elements:
        for t in omath_para.findall(f".//{{{M_NS}}}t"):
            if t.text:
                parts.append(t.text)
    return "".join(parts)


def _extract_inline_formula_latex(omath_elem: Any) -> str:
    """Собрать LaTeX из `<m:oMath>` для inline-формулы.

    Снимает обрамляющие `$...$` (если экспортёр их добавил), оставляет
    содержимое `<m:t>` как есть.
    """
    parts: list[str] = []
    for t in omath_elem.findall(f".//{{{M_NS}}}t"):
        if t.text:
            parts.append(t.text)
    latex = "".join(parts)
    # Снять долларовые ограничители, если экспортёр их применил.
    if len(latex) >= 2 and latex.startswith("$") and latex.endswith("$"):
        latex = latex[1:-1]
    return latex


def _extract_formula_number(dp: DocxParagraph) -> int | None:
    """Извлечь номер формулы из паттерна «(N)» / «(N.M)» в конце параграфа."""
    text = dp.text or ""
    match = _FORMULA_NUMBER_RE.search(text)
    if match is None:
        return None
    raw = match.group(1).split(".")[0]
    try:
        return int(raw)
    except ValueError:
        return None


def _build_table(dtable: DocxTable, *, idx: int) -> Table:
    """Сконвертировать docx-таблицу в модель Table.

    Headers — первый ряд cells; rows — остальные ряды. Каждая ячейка
    хранится как list[InlineElement] из текста первого параграфа
    ячейки (без атрибутов форматирования — Фаза 1).

    Заполняет ``merges: list[CellMerge]`` информацией об объединённых
    ячейках: ``<w:vMerge>`` (вертикальное), ``<w:gridSpan>``
    (горизонтальное). Координаты row/col 0-based от верха таблицы.
    """
    from gostforge.model import CellMerge  # noqa: PLC0415

    rows_raw = list(dtable.rows)
    headers: list[list[InlineElement]] = []
    body_rows: list[list[list[InlineElement]]] = []
    merges: list[CellMerge] = []

    if rows_raw:
        headers = [_cell_inline(cell) for cell in rows_raw[0].cells]
        for row in rows_raw[1:]:
            body_rows.append([_cell_inline(cell) for cell in row.cells])
        merges = _extract_cell_merges(rows_raw)

    return Table(
        id=f"t-{idx}",
        caption=[],
        headers=headers,
        rows=body_rows,
        merges=merges,
    )


def _extract_cell_merges(rows_raw: list[Any]) -> list[Any]:
    """Извлечь информацию об объединённых ячейках из таблицы.

    Алгоритм:
    1. Для каждой ячейки читаем <w:tcPr>/<w:gridSpan w:val="N"/> (colspan).
    2. <w:tcPr>/<w:vMerge w:val="restart"/> начинает вертикальное
       объединение; <w:vMerge/> без val (или val="continue") — продолжение.
    3. Считаем rowspan: для каждой restart-ячейки идём вниз по той же
       колонке, считая continue-ячейки.

    Возвращает CellMerge только для ячеек с rowspan>1 или colspan>1.
    """
    from gostforge.model import CellMerge  # noqa: PLC0415

    merges: list[CellMerge] = []
    # Построим матрицу tc-элементов для удобного поиска.
    tc_grid: list[list[Any]] = []
    for row in rows_raw:
        tc_grid.append([cell._tc for cell in row.cells])

    n_rows = len(tc_grid)
    for r, row_tcs in enumerate(tc_grid):
        for c, tc in enumerate(row_tcs):
            tcPr = tc.find(f"{{{W_NS}}}tcPr")
            colspan = 1
            v_merge_kind: str | None = None
            if tcPr is not None:
                grid_span = tcPr.find(f"{{{W_NS}}}gridSpan")
                if grid_span is not None:
                    try:
                        colspan = int(grid_span.get(f"{{{W_NS}}}val", "1"))
                    except ValueError:
                        colspan = 1
                v_merge = tcPr.find(f"{{{W_NS}}}vMerge")
                if v_merge is not None:
                    v_merge_kind = v_merge.get(f"{{{W_NS}}}val") or "continue"
            # Continue-ячейки пропускаем — они уже учтены в restart выше.
            if v_merge_kind == "continue":
                continue
            # Считаем rowspan: смотрим в ту же колонку (с учётом colspan
            # — python-docx обычно возвращает физические ячейки, и
            # continue-tc лежит в той же позиции в row_tcs).
            rowspan = 1
            if v_merge_kind == "restart":
                for r2 in range(r + 1, n_rows):
                    if c >= len(tc_grid[r2]):
                        break
                    next_tc = tc_grid[r2][c]
                    next_tcPr = next_tc.find(f"{{{W_NS}}}tcPr")
                    if next_tcPr is None:
                        break
                    next_vm = next_tcPr.find(f"{{{W_NS}}}vMerge")
                    if next_vm is None:
                        break
                    val = next_vm.get(f"{{{W_NS}}}val")
                    if val == "restart":
                        # Новое объединение — текущее закончилось.
                        break
                    # val=None или 'continue' — продолжение.
                    rowspan += 1
            if rowspan > 1 or colspan > 1:
                merges.append(CellMerge(row=r, col=c, rowspan=rowspan, colspan=colspan))
    return merges


def _cell_inline(cell: DocxCell) -> list[InlineElement]:
    """Извлечь inline-содержимое первого параграфа ячейки (только текст)."""
    paragraphs = cell.paragraphs
    if not paragraphs:
        return []
    text = paragraphs[0].text or ""
    if not text:
        return []
    return [TextRun(text=text)]


def _glue_captions(items: list[LogicalSection | Block]) -> None:
    """Эвристика склейки подписей.

    Один проход по списку блоков:
    - параграф-подпись СВЕРХУ присоединяется к следующей таблице;
    - параграф-подпись СНИЗУ присоединяется к предыдущему рисунку;
    - параграфы-подписи удаляются из списка.

    Параграф, ставший подписью таблицы, не должен также считаться
    подписью рисунка — поэтому таблица обрабатывается раньше рисунка
    при проверке предыдущего элемента в `result`.
    """
    if not items:
        return

    result: list[LogicalSection | Block] = []
    i = 0
    n = len(items)
    while i < n:
        item = items[i]
        if isinstance(item, Table):
            if result and isinstance(result[-1], Paragraph) and _is_table_caption(result[-1]):
                cap = result.pop()
                assert isinstance(cap, Paragraph)
                item.caption = list(cap.content)
            result.append(item)
            i += 1
            continue

        if isinstance(item, Figure):
            result.append(item)
            j = i + 1
            if j < n:
                nxt = items[j]
                if isinstance(nxt, Paragraph) and _is_figure_caption(nxt):
                    item.caption = list(nxt.content)
                    i = j + 1
                    continue
            i += 1
            continue

        result.append(item)
        i += 1

    items[:] = result


def _is_figure_caption(paragraph: Paragraph) -> bool:
    """Параграф похож на подпись рисунка: стиль Caption или текст «Рисунок N ...»."""
    if _has_caption_style(paragraph):
        return True
    text = _paragraph_plain_text(paragraph).strip()
    return bool(_FIGURE_CAPTION_TEXT_RE.match(text))


def _is_table_caption(paragraph: Paragraph) -> bool:
    """Параграф похож на подпись таблицы: стиль Caption или текст «Таблица N ...»."""
    if _has_caption_style(paragraph):
        return True
    text = _paragraph_plain_text(paragraph).strip()
    return bool(_TABLE_CAPTION_TEXT_RE.match(text))


def _has_caption_style(paragraph: Paragraph) -> bool:
    """Проверить, что style_name похож на «Caption» (с учётом регистра/вариантов)."""
    if paragraph.style_name is None:
        return False
    return paragraph.style_name.strip().lower() in _CAPTION_STYLE_NAMES


def _paragraph_plain_text(paragraph: Paragraph) -> str:
    """Склейка text всех TextRun параграфа."""
    return "".join(el.text for el in paragraph.content if isinstance(el, TextRun))


def _heading_level(p: DocxParagraph) -> int | None:
    """Вернуть уровень заголовка (1..9), если стиль — Heading N, иначе None."""
    style = p.style
    if style is None or style.name is None:
        return None
    m = _HEADING_RE.match(style.name)
    if not m:
        return None
    return int(m.group(1))


def _build_paragraph(p: DocxParagraph, *, idx: int) -> Paragraph:
    """Сконвертировать docx-параграф в модель Paragraph.

    Идём по детям `<w:p>` в порядке появления (document order). Для каждого
    типа ребёнка строим соответствующий InlineElement:
      - `<w:r>` → TextRun (с bold/italic/underline/font/size/color из `<w:rPr>`).
      - `<m:oMath>` без обёртки `<m:oMathPara>` → InlineFormula.
      - `<w:fldSimple w:instr=" REF target_id ...">` → CrossRef.
      - Прочие служебные элементы (`<w:pPr>`, `<w:bookmarkStart>` и т. п.)
        пропускаются.

    После того как content собран, выполняется эвристика «prefix у CrossRef»:
    если непосредственно перед CrossRef стоит TextRun, чей текст оканчивается
    на «(см. », « (» или подобный «открывающий» хвост — этот хвост переносится
    в `CrossRef.prefix`, а TextRun.text укорачивается.
    """
    style_name = p.style.name if p.style is not None else None
    pf = p.paragraph_format
    alignment = _alignment_to_literal(pf.alignment)
    line_spacing = _line_spacing_value(pf.line_spacing)
    indent_cm = _indent_cm(pf.first_line_indent)
    page_break_before = _page_break_before(p)
    space_before_pt = _spacing_pt(pf.space_before)
    space_after_pt = _spacing_pt(pf.space_after)

    content: list[InlineElement] = []
    style_font_name = _style_font_name(p)
    style_font_size_pt = _style_font_size_pt(p)
    style_bold = _style_bold(p)
    style_italic = _style_italic(p)
    style_color_hex = _style_color_hex(p)

    p_xml = p._p
    for child in p_xml.iterchildren():
        tag = etree.QName(child.tag).localname
        ns = etree.QName(child.tag).namespace

        if ns == W_NS and tag == "r":
            # Внутри <w:r> может лежать <m:oMath> — экспортёр пишет inline-формулу
            # именно так. В этом случае run становится InlineFormula (не TextRun).
            nested_omath = child.find(f"{{{M_NS}}}oMath")
            if nested_omath is not None:
                latex = _extract_inline_formula_latex(nested_omath)
                content.append(InlineFormula(latex=latex))
                continue
            # <w:footnoteReference w:id="N"/> — ссылка на сноску.
            fn_ref = child.find(f"{{{W_NS}}}footnoteReference")
            if fn_ref is not None:
                fn_id = fn_ref.get(f"{{{W_NS}}}id", "")
                if fn_id:
                    content.append(FootnoteRef(footnote_id=fn_id))
                continue
            text_run = _text_run_from_w_r(
                child,
                style_font_name=style_font_name,
                style_font_size_pt=style_font_size_pt,
                style_bold=style_bold,
                style_italic=style_italic,
                style_color_hex=style_color_hex,
            )
            if text_run is not None:
                content.append(text_run)
        elif ns == M_NS and tag == "oMath":
            # Inline-формула: <m:oMath> прямой ребёнок <w:p>, не обёрнутый в
            # <m:oMathPara> (блочные формулы обработаны в _block_from_paragraph).
            latex = _extract_inline_formula_latex(child)
            content.append(InlineFormula(latex=latex))
        elif ns == W_NS and tag == "fldSimple":
            cross_ref = _cross_ref_from_fld_simple(child)
            if cross_ref is not None:
                content.append(cross_ref)
            else:
                # fldSimple без распознаваемой REF-инструкции — берём
                # внутренний текст как обычный TextRun, чтобы не потерять
                # содержимое (например, PAGE-поле внутри тела документа).
                inner_text = _fld_simple_inner_text(child)
                if inner_text:
                    content.append(TextRun(text=inner_text))
        elif ns == W_NS and tag == "hyperlink":
            # <w:hyperlink r:id="rIdN" w:anchor="bookmark">.
            # Resolve URL через rels source-документа; anchor — внутренний.
            hyperlink = _hyperlink_from_w_hyperlink(child, p)
            if hyperlink is not None:
                content.append(hyperlink)
        # Прочие элементы (`<w:pPr>`, `<w:bookmarkStart>`, и т. п.) — игнорируем.

    _attach_cross_ref_prefixes(content)

    return Paragraph(
        id=f"p-{idx}",
        content=content,
        style_name=style_name,
        alignment=alignment,
        line_spacing=line_spacing,
        first_line_indent_cm=indent_cm,
        page_break_before=page_break_before,
        space_before_pt=space_before_pt,
        space_after_pt=space_after_pt,
    )


def _text_run_from_w_r(
    w_r_elem: Any,
    *,
    style_font_name: str | None,
    style_font_size_pt: float | None,
    style_bold: bool | None = None,
    style_italic: bool | None = None,
    style_color_hex: str | None = None,
) -> TextRun | None:
    """Построить TextRun из `<w:r>` элемента, читая `<w:rPr>` и `<w:t>`.

    Возвращает None, если у run нет видимого текста — такие пустые run-ы
    Word оставляет после удаления (бывают, например, после bookmark-маркеров).

    Style-cascade: атрибуты font, size, bold, italic читаются сначала из
    `<w:rPr>` самого run-а; если на run-уровне атрибут не задан явно —
    берётся из стиля абзаца (включая linked character-стиль и цепочку
    наследования стилей). Это необходимо, чтобы такие проверки как
    H.01 (формат заголовка 1: TNR, 14 pt, жирный) корректно срабатывали
    на документах, сгенерированных через python-docx `add_heading()` —
    он не пишет явные run-атрибуты, опирается на стиль Heading{N}.
    """
    # Текст: склейка всех <w:t> внутри run-а.
    texts: list[str] = []
    for t in w_r_elem.findall(f"{{{W_NS}}}t"):
        if t.text:
            texts.append(t.text)
    text = "".join(texts)
    if not text:
        return None

    rpr = w_r_elem.find(f"{{{W_NS}}}rPr")

    bold = _rpr_toggle(rpr, "b")
    italic = _rpr_toggle(rpr, "i")
    underline = _rpr_underline(rpr)
    color_hex = _rpr_color(rpr)

    # Style-cascade: если на run-уровне атрибут не задан (None) — наследуем
    # от стиля параграфа. Если на run-уровне явно False — оставляем False
    # (явное «не жирный» поверх жирного стиля = реальное намерение автора).
    if bold is None and style_bold is not None:
        bold = style_bold
    if italic is None and style_italic is not None:
        italic = style_italic
    if color_hex is None and style_color_hex is not None:
        color_hex = style_color_hex

    # font и size: сначала из rPr, иначе из стиля абзаца.
    font_name = _rpr_font(rpr) or style_font_name
    size_pt = _rpr_size_pt(rpr)
    if size_pt is None:
        size_pt = style_font_size_pt

    return TextRun(
        text=text,
        bold=bold,
        italic=italic,
        underline=underline,
        font=font_name,
        size_pt=size_pt,
        color_hex=color_hex,
    )


def _rpr_toggle(rpr_elem: Any, tag: str) -> bool | None:
    """Прочитать toggle-свойство (`<w:b/>`, `<w:i/>`) из `<w:rPr>`.

    OOXML toggle: элемент с val="0"/"false" — отключено; иначе — включено.
    Отсутствие элемента — None (наследуется).
    """
    if rpr_elem is None:
        return None
    el = rpr_elem.find(f"{{{W_NS}}}{tag}")
    if el is None:
        return None
    val = el.get(f"{{{W_NS}}}val")
    return val not in {"0", "false"}


def _rpr_underline(rpr_elem: Any) -> bool | None:
    """Прочитать `<w:u>` из `<w:rPr>`. underline=True если val != "none"."""
    if rpr_elem is None:
        return None
    el = rpr_elem.find(f"{{{W_NS}}}u")
    if el is None:
        return None
    val = el.get(f"{{{W_NS}}}val")
    # Отсутствие val или val != "none" — подчёркивание включено.
    return bool(val != "none")


def _rpr_color(rpr_elem: Any) -> str | None:
    """Прочитать `<w:color>` из `<w:rPr>` и вернуть '#RRGGBB' или None.

    Значения "auto" и пустые — None (цвет наследуется или автоматический).
    """
    if rpr_elem is None:
        return None
    el = rpr_elem.find(f"{{{W_NS}}}color")
    if el is None:
        return None
    val = el.get(f"{{{W_NS}}}val")
    if not val or val.lower() == "auto":
        return None
    # Word хранит цвет как hex без «#», верхний регистр. Нормализуем.
    return f"#{val.upper()}"


def _rpr_font(rpr_elem: Any) -> str | None:
    """Прочитать имя шрифта из `<w:rFonts>` (ascii > hAnsi > cs)."""
    if rpr_elem is None:
        return None
    rfonts = rpr_elem.find(f"{{{W_NS}}}rFonts")
    if rfonts is None:
        return None
    for attr in ("ascii", "hAnsi", "cs"):
        name = rfonts.get(f"{{{W_NS}}}{attr}")
        if name:
            return str(name)
    return None


def _rpr_size_pt(rpr_elem: Any) -> float | None:
    """Прочитать размер шрифта из `<w:sz>`. Значение — половинки пункта."""
    if rpr_elem is None:
        return None
    sz = rpr_elem.find(f"{{{W_NS}}}sz")
    if sz is None:
        return None
    val = sz.get(f"{{{W_NS}}}val")
    if val is None:
        return None
    try:
        return float(val) / 2.0
    except (TypeError, ValueError):
        return None


# --- CrossRef -----------------------------------------------------------------


# Парсер ищет инструкцию вида « REF <bookmark> [\h] ». Толерантно к лишним
# пробелам и регистру (Word нормализует REF, но другие конвертеры — не всегда).
_FLD_REF_INSTR_RE = re.compile(r"^\s*REF\s+([^\s\\]+)", re.IGNORECASE)


def _hyperlink_from_w_hyperlink(hl_elem: Any, paragraph: Any) -> Hyperlink | None:
    """Построить Hyperlink из ``<w:hyperlink r:id="..." w:anchor="...">``.

    Алгоритм:
    1. Собрать текст из всех вложенных ``<w:r>/<w:t>``.
    2. Если есть ``r:id`` — разрешить URL через
       paragraph.part.rels[r:id].target_ref.
    3. Если есть ``w:anchor`` — внутренняя ссылка на bookmark.

    Возвращает None если нет ни URL, ни anchor, или нет видимого текста.
    """
    # Текст: склейка всех <w:t> внутри.
    texts: list[str] = []
    for t in hl_elem.iter(f"{{{W_NS}}}t"):
        if t.text:
            texts.append(t.text)
    text = "".join(texts)
    if not text:
        return None

    r_id = hl_elem.get(f"{{{_R_NS}}}id")
    anchor = hl_elem.get(f"{{{W_NS}}}anchor")
    url = ""
    if r_id:
        try:
            rels = paragraph.part.rels
            if r_id in rels:
                url = rels[r_id].target_ref
        except AttributeError:
            pass
    if not url and not anchor:
        return None
    return Hyperlink(url=url or "", text=text, anchor=anchor)


def _cross_ref_from_fld_simple(fld_elem: Any) -> CrossRef | None:
    """Построить CrossRef из `<w:fldSimple w:instr=" REF target_id ...">`.

    Возвращает None, если инструкция не REF (например, PAGE) — такой
    fldSimple обрабатывается отдельно (его текст-плейсхолдер кладётся как TextRun).
    """
    instr = fld_elem.get(f"{{{W_NS}}}instr") or ""
    match = _FLD_REF_INSTR_RE.match(instr)
    if match is None:
        return None
    target_id = match.group(1)
    return CrossRef(target_id=target_id)


def _fld_simple_inner_text(fld_elem: Any) -> str:
    """Склейка `<w:t>` внутри `<w:fldSimple>` (для не-REF полей)."""
    parts: list[str] = []
    for t in fld_elem.iter(f"{{{W_NS}}}t"):
        if t.text:
            parts.append(t.text)
    return "".join(parts)


# Эвристика «хвост-prefix» CrossRef: список окончаний TextRun-а, которые
# уместно перенести в CrossRef.prefix. Порядок важен — сначала более
# длинные шаблоны, чтобы не сматчить «(см. » раньше « (см. ».
_CROSS_REF_PREFIX_SUFFIXES: tuple[str, ...] = (
    " (см. ",
    "(см. ",
    " (",
)


def _attach_cross_ref_prefixes(content: list[InlineElement]) -> None:
    """Пост-обработка: переносит «хвост» предыдущего TextRun-а в CrossRef.prefix.

    Эвристика консервативная: prefix присваивается только тогда, когда хвост
    однозначно опознан как «открывающий» (входит в `_CROSS_REF_PREFIX_SUFFIXES`).
    Во всех остальных случаях оставляем CrossRef.prefix=None.
    """
    for i, element in enumerate(content):
        if not isinstance(element, CrossRef):
            continue
        if i == 0:
            continue
        prev = content[i - 1]
        if not isinstance(prev, TextRun):
            continue
        prefix, remainder = _split_cross_ref_prefix(prev.text)
        if prefix is None:
            continue
        element.prefix = prefix
        prev.text = remainder


def _split_cross_ref_prefix(text: str) -> tuple[str | None, str]:
    """Если `text` оканчивается на «(см. » / « (» — отщепить этот хвост.

    Возвращает (prefix, remainder). Если ни один из шаблонов не подошёл —
    (None, text) без изменений.
    """
    for suffix in _CROSS_REF_PREFIX_SUFFIXES:
        if text.endswith(suffix):
            return suffix, text[: -len(suffix)]
    return None, text


def _page_break_before(p: DocxParagraph) -> bool | None:
    """Узнать, выставлен ли у параграфа разрыв страницы перед ним.

    Проверяем сам параграф (w:pPr/w:pageBreakBefore) и цепочку стилей
    (Word-стили могут наследовать pageBreakBefore от base_style). Возвращаем:
      - True, если флаг найден явно;
      - None, если не найден ни на параграфе, ни в стиле (наследуется/
        не задан).

    Возможный False в текущей реализации не выдаём — w:pageBreakBefore
    в OOXML по умолчанию означает True; снятие флага через `w:val="0"`
    встречается редко, но мы его учитываем.
    """
    # 1) Прямое свойство параграфа: w:pPr/w:pageBreakBefore
    direct = _paragraph_break_flag(p._p)
    if direct is not None:
        return direct

    # 2) Свойство, унаследованное от стиля абзаца (с обходом base_style)
    style = p.style
    while style is not None:
        style_element = getattr(style, "element", None)
        if style_element is not None:
            inherited = _style_break_flag(style_element)
            if inherited is not None:
                return inherited
        style = getattr(style, "base_style", None)
    return None


def _paragraph_break_flag(p_elem: object) -> bool | None:
    """Найти w:pageBreakBefore в w:pPr заданного <w:p>."""
    p_pr = p_elem.find(f"{{{W_NS}}}pPr")  # type: ignore[attr-defined]
    if p_pr is None:
        return None
    return _read_break_flag(p_pr)


def _style_break_flag(style_elem: object) -> bool | None:
    """Найти w:pageBreakBefore в w:pPr стилей (<w:style>/<w:docDefaults>)."""
    p_pr = style_elem.find(f"{{{W_NS}}}pPr")  # type: ignore[attr-defined]
    if p_pr is None:
        return None
    return _read_break_flag(p_pr)


def _read_break_flag(p_pr_elem: object) -> bool | None:
    """Прочитать <w:pageBreakBefore> внутри <w:pPr>.

    OOXML toggle-семантика: элемент без атрибута w:val или со значением
    "true"/"1"/"on" — True; "false"/"0"/"off" — False; отсутствие
    элемента — None.
    """
    elem = p_pr_elem.find(f"{{{W_NS}}}pageBreakBefore")  # type: ignore[attr-defined]
    if elem is None:
        return None
    val = elem.get(f"{{{W_NS}}}val")
    if val is None:
        return True
    return val.lower() not in {"false", "0", "off"}


def _alignment_to_literal(value: object | None) -> ParagraphAlignment | None:
    """WD_ALIGN_PARAGRAPH → литералка ParagraphAlignment."""
    if value is None:
        return None
    try:
        return _ALIGN_MAP.get(int(cast(Any, value)))
    except (TypeError, ValueError):
        return None


def _spacing_pt(value: object | None) -> float | None:
    """Преобразовать docx-spacing (EMU/Twips object) в pt.

    python-docx возвращает либо docx.shared.Pt-like объект (с .pt
    атрибутом), либо None если атрибут не задан. Возвращаем pt или
    None.
    """
    if value is None:
        return None
    try:
        return float(value.pt)  # type: ignore[attr-defined]
    except (AttributeError, TypeError):
        return None


def _line_spacing_value(value: object | None) -> float | None:
    """Привести line_spacing к float (множитель) или None."""
    if value is None:
        return None
    try:
        return float(cast(Any, value))
    except (TypeError, ValueError):
        return None


def _indent_cm(value: object | None) -> float | None:
    """Перевести first_line_indent в сантиметры."""
    if value is None:
        return None
    cm = getattr(value, "cm", None)
    if cm is None:
        return None
    return round(float(cm), 2)


def _style_font_name(p: DocxParagraph) -> str | None:
    """Имя шрифта, заданное на стиле абзаца (с обходом цепочки наследования стилей)."""
    style = p.style
    while style is not None:
        name = style.font.name if style.font is not None else None
        if name:
            return str(name)
        style = getattr(style, "base_style", None)
    return None


def _style_font_size_pt(p: DocxParagraph) -> float | None:
    """Размер шрифта (в пунктах) из стиля абзаца."""
    style = p.style
    while style is not None:
        size = style.font.size if style.font is not None else None
        if size is not None:
            return float(size.pt)
        style = getattr(style, "base_style", None)
    return None


def _style_bold(p: DocxParagraph) -> bool | None:
    """Жирность, заданная на стиле абзаца, с обходом цепочки наследования.

    Возвращает True/False, если на каком-то уровне цепочки атрибут задан
    явно. None — если ни на одном уровне явно не задан (наследуется
    дальше от Word-defaults). Это позволяет различать «стиль не задал»
    от «стиль явно сказал не жирный».
    """
    style = p.style
    while style is not None:
        bold = style.font.bold if style.font is not None else None
        if bold is not None:
            return bool(bold)
        style = getattr(style, "base_style", None)
    return None


def _style_italic(p: DocxParagraph) -> bool | None:
    """Курсив, заданный на стиле абзаца, с обходом цепочки наследования."""
    style = p.style
    while style is not None:
        italic = style.font.italic if style.font is not None else None
        if italic is not None:
            return bool(italic)
        style = getattr(style, "base_style", None)
    return None


def _style_color_hex(p: DocxParagraph) -> str | None:
    """Цвет шрифта стиля абзаца (с обходом цепочки наследования и
    linked character-стиля).

    Word при рендере run-ов наследует форматирование сразу из двух мест:
    параграф-стиля (Heading1) И его linked char-стиля (Heading1Char).
    Парсер должен проверять оба места — иначе синий Cambria из
    дефолтного шаблона остаётся незамеченным (синий обычно сидит
    в Heading1Char, а не в Heading1).

    Цвет читается через прямой XML — python-docx style.font.color может
    возвращать None при themeColor (атрибут w:val при этом есть и
    содержит реальный hex).
    """

    def _color_in_style_element(st_elem: Any) -> str | None:
        if st_elem is None:
            return None
        rPr = st_elem.find(f"{{{W_NS}}}rPr")
        if rPr is None:
            return None
        color = rPr.find(f"{{{W_NS}}}color")
        if color is None:
            return None
        val = color.get(f"{{{W_NS}}}val")
        if not val:
            return None
        return f"#{val.upper()}" if val.lower() != "auto" else None

    style = p.style
    styles_root = None
    while style is not None:
        st_elem = getattr(style, "element", None)
        if st_elem is not None:
            # 1. Прямой цвет в этом стиле.
            direct = _color_in_style_element(st_elem)
            if direct is not None:
                return direct
            # 2. Цвет в linked character-стиле (w:link).
            link = st_elem.find(f"{{{W_NS}}}link")
            if link is not None:
                char_id = link.get(f"{{{W_NS}}}val")
                if char_id:
                    if styles_root is None:
                        styles_root = st_elem.getparent()
                    if styles_root is not None:
                        for st in styles_root.findall(f"{{{W_NS}}}style"):
                            if (
                                st.get(f"{{{W_NS}}}type") == "character"
                                and st.get(f"{{{W_NS}}}styleId") == char_id
                            ):
                                via_link = _color_in_style_element(st)
                                if via_link is not None:
                                    return via_link
                                break
        style = getattr(style, "base_style", None)
    return None


# --- список литературы --------------------------------------------------------


def _extract_bibliography(page_sections: list[PageSection]) -> list[BibliographyEntry]:
    """Найти раздел со списком литературы и собрать записи в плоский список.

    Алгоритм:
    1. Перебираем все LogicalSection уровня 1 во всех PageSection.
    2. Сравниваем нормализованный текст заголовка со списком известных
       названий раздела (`_BIBLIOGRAPHY_HEADINGS`).
    3. Каждый прямой дочерний Paragraph такой секции, текст которого
       непустой после strip, превращаем в `BibliographyEntry` с id вида
       `ref-{idx}`, где idx — сквозной счётчик по всем найденным записям.
    4. Тип записи определяется эвристически по содержимому текста.
    5. `fields["raw"]` — полный текст параграфа (минимум для Фазы 1).
    6. Через `_parse_bibliography_fields` извлекаются опциональные
       структурные поля (author, year, url, doi, access_date, place,
       language) — используются проверками R.02/R.03/R.08-R.13.

    Пустые параграфы пропускаются. Параграфы из раздела остаются на месте
    в `LogicalSection.children` — модель сознательно дублирует данные
    в `Document.bibliography`, чтобы валидаторы могли работать как по
    плоскому списку, так и по содержанию страницы.
    """
    entries: list[BibliographyEntry] = []
    idx = 0
    for ps in page_sections:
        for section in _iter_logical_sections(ps.content):
            if section.level != 1:
                continue
            heading_text = _normalize_heading(section.heading)
            if heading_text not in _BIBLIOGRAPHY_HEADINGS:
                continue
            for child in section.children:
                if not isinstance(child, Paragraph):
                    continue
                raw = _paragraph_plain_text(child).strip()
                if not raw:
                    continue
                idx += 1
                fields = {"raw": raw}
                fields.update(_parse_bibliography_fields(raw))
                entries.append(
                    BibliographyEntry(
                        id=f"ref-{idx}",
                        type=_detect_bibliography_type(raw),
                        fields=fields,
                    )
                )
    return entries


def _normalize_heading(content: Sequence[InlineElement]) -> str:
    """Привести inline-заголовок к строке без регистра и лишних пробелов."""
    text = "".join(el.text for el in content if isinstance(el, TextRun))
    return " ".join(text.lower().split())


def _detect_bibliography_type(
    raw: str,
) -> Literal["book", "article", "web", "standard", "thesis", "conference", "law"]:
    """Эвристически определить тип библиографической записи по её тексту."""
    if _BIB_URL_RE.search(raw):
        return "web"
    if _BIB_STANDARD_RE.search(raw) or raw.startswith("ГОСТ "):
        return "standard"
    if _BIB_THESIS_RE.search(raw):
        return "thesis"
    if _BIB_CONFERENCE_RE.search(raw):
        return "conference"
    if _BIB_LAW_RE.search(raw):
        return "law"
    if _BIB_ARTICLE_RE.search(raw):
        return "article"
    return "book"


# --- Citation post-processing -------------------------------------------------


# «[N]» либо «[N, с. P]» / «[N, с.P]». N — 1-3 цифры. P — допускаем «12»,
# «12-15», «12, 17-20» (любые цифры, тире, запятые, пробелы). Запрещены
# вложенные скобки.
# Регистр «с»/«С» — нечувствительно.
_CITATION_PATTERN_RE = re.compile(r"\[(\d{1,3})(?:,\s*[сС]\.\s*([\d,\s\-–—]+?))?\]")


def _annotate_citations(
    page_sections: list[PageSection], bibliography: list[BibliographyEntry]
) -> None:
    """Заменить в TextRun-ах подстроки «[N]» / «[N, с. P]» на Citation.

    Срабатывает только когда N валидный 1-based индекс bibliography
    (1 ≤ N ≤ len(bibliography)). Иначе TextRun не модифицируется.
    Применяется ко всем Paragraph во всех PageSection и LogicalSection
    рекурсивно (включая children внутри Figure.caption / Table.caption —
    обработать позже при необходимости; на Фазе 2.5 ограничиваемся
    Paragraph.content).
    """
    if not bibliography:
        return
    for ps in page_sections:
        _annotate_citations_in_items(ps.content, bibliography)


def _annotate_citations_in_items(
    items: list[LogicalSection | Block], bibliography: list[BibliographyEntry]
) -> None:
    """Рекурсивно обойти элементы и заменить citations в Paragraph.content."""
    for item in items:
        if isinstance(item, LogicalSection):
            _annotate_citations_in_items(item.children, bibliography)
        elif isinstance(item, Paragraph):
            item.content = _annotate_citations_in_inline(item.content, bibliography)


def _annotate_citations_in_inline(
    content: list[InlineElement], bibliography: list[BibliographyEntry]
) -> list[InlineElement]:
    """Заменить «[N]»/«[N, с. P]» в TextRun-ах на Citation там, где N валиден."""
    result: list[InlineElement] = []
    for element in content:
        if not isinstance(element, TextRun):
            result.append(element)
            continue
        result.extend(_split_text_run_by_citations(element, bibliography))
    return result


def _split_text_run_by_citations(
    run: TextRun, bibliography: list[BibliographyEntry]
) -> list[InlineElement]:
    """Разрезать один TextRun на (TextRun, Citation, TextRun, ...) по паттерну.

    Атрибуты форматирования (`bold`, `italic`, `font`, ...) копируются
    в каждый получившийся фрагмент TextRun-а — это сохраняет визуальное
    оформление при обратном раунд-трипе.
    """
    text = run.text
    matches = list(_CITATION_PATTERN_RE.finditer(text))
    if not matches:
        return [run]

    pieces: list[InlineElement] = []
    cursor = 0
    changed = False
    for match in matches:
        try:
            n_value = int(match.group(1))
        except (TypeError, ValueError):
            continue
        # Эвристика: N должен быть валидным 1-based индексом bibliography.
        if n_value < 1 or n_value > len(bibliography):
            continue
        source_id = bibliography[n_value - 1].id
        pages_raw = match.group(2)
        pages = _normalize_citation_pages(pages_raw) if pages_raw else None

        # Текст ДО найденного match-а сохраняем как TextRun с тем же оформлением.
        if match.start() > cursor:
            pieces.append(_clone_text_run(run, text[cursor : match.start()]))
        template = "[{n}, с. {pages}]" if pages else "[{n}]"
        pieces.append(Citation(source_id=source_id, pages=pages, template=template))
        cursor = match.end()
        changed = True

    if not changed:
        return [run]

    if cursor < len(text):
        pieces.append(_clone_text_run(run, text[cursor:]))

    # Удалить возможные пустые TextRun-ы (когда citation шёл встык).
    return [p for p in pieces if not (isinstance(p, TextRun) and not p.text)]


def _normalize_citation_pages(raw: str) -> str:
    """Привести «с. 12-15» к каноническому виду «12-15» (убрать лишние пробелы)."""
    # Сжимаем подряд идущие пробелы и убираем краевые.
    return " ".join(raw.split()).strip()


def _clone_text_run(template_run: TextRun, new_text: str) -> TextRun:
    """Скопировать TextRun, заменив только text. Используется при сплите по citations."""
    return TextRun(
        text=new_text,
        bold=template_run.bold,
        italic=template_run.italic,
        underline=template_run.underline,
        superscript=template_run.superscript,
        subscript=template_run.subscript,
        font=template_run.font,
        size_pt=template_run.size_pt,
        color_hex=template_run.color_hex,
    )
