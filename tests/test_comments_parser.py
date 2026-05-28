"""Тесты парсинга комментариев Word из word/comments.xml.

Создаём docx через python-docx, вручную добавляем CommentsPart с
несколькими <w:comment> и вызываем parse_docx. Проверяем что
Document.comments содержит правильно распарсенные элементы.
"""

from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document as DocxDocument

from gostforge.model import Comment
from gostforge.parser import parse_docx


def _make_docx_with_comments(
    tmp_path: Path,
    comments_xml: str,
) -> Path:
    """Создать docx через python-docx и добавить word/comments.xml
    вручную через zip-patch."""
    doc = DocxDocument()
    doc.add_paragraph("Тестовый текст.")
    raw_path = tmp_path / "raw.docx"
    doc.save(raw_path)

    out_path = tmp_path / "with-comments.docx"
    # Перепаковываем zip с добавлением comments.xml и обновлением
    # [Content_Types].xml + word/_rels/document.xml.rels.
    with zipfile.ZipFile(raw_path, "r") as zin:
        names = zin.namelist()
        contents = {n: zin.read(n) for n in names}

    # Добавляем content type.
    ct = contents["[Content_Types].xml"].decode("utf-8")
    if "comments.xml" not in ct:
        ct = ct.replace(
            "</Types>",
            '<Override PartName="/word/comments.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.'
            'wordprocessingml.comments+xml"/></Types>',
        )
    contents["[Content_Types].xml"] = ct.encode("utf-8")

    # Добавляем relationship.
    rels_path = "word/_rels/document.xml.rels"
    rels = contents[rels_path].decode("utf-8")
    if "comments.xml" not in rels:
        rels = rels.replace(
            "</Relationships>",
            '<Relationship Id="rId100" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/'
            'relationships/comments" Target="comments.xml"/></Relationships>',
        )
    contents[rels_path] = rels.encode("utf-8")

    # Сам comments.xml.
    contents["word/comments.xml"] = comments_xml.encode("utf-8")

    with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in contents.items():
            zout.writestr(name, data)
    return out_path


_TWO_COMMENTS_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:comment w:id="0" w:author="Иванов И.И." w:date="2026-05-26T10:00:00Z" w:initials="ИИ">
    <w:p><w:r><w:t>Замечание про стиль.</w:t></w:r></w:p>
  </w:comment>
  <w:comment w:id="1" w:author="Петров П.П." w:date="2026-05-26T11:30:00Z" w:initials="ПП">
    <w:p><w:r><w:t>Перепиши абзац.</w:t></w:r></w:p>
    <w:p><w:r><w:t>Слишком длинный.</w:t></w:r></w:p>
  </w:comment>
</w:comments>
"""


def test_parser_extracts_comments(tmp_path: Path) -> None:
    path = _make_docx_with_comments(tmp_path, _TWO_COMMENTS_XML)
    doc = parse_docx(path)
    assert len(doc.comments) == 2
    c0, c1 = doc.comments
    assert c0.id == "0"
    assert c0.author == "Иванов И.И."
    assert c0.date == "2026-05-26T10:00:00Z"
    assert c0.text == "Замечание про стиль."
    assert c1.author == "Петров П.П."
    assert "длинный" in c1.text


def test_parser_returns_empty_comments_when_no_comments_part(
    tmp_path: Path,
) -> None:
    """Документ без comments.xml → Document.comments == [], без ошибок."""
    doc = DocxDocument()
    doc.add_paragraph("Без комментариев.")
    path = tmp_path / "no-comments.docx"
    doc.save(path)
    parsed = parse_docx(path)
    assert parsed.comments == []


def test_parser_handles_empty_comments_xml(tmp_path: Path) -> None:
    """Валидный, но пустой comments.xml → comments == [], не падает."""
    empty_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:comments xmlns:w="http://schemas.openxmlformats.org/'
        'wordprocessingml/2006/main"></w:comments>'
    )
    path = _make_docx_with_comments(tmp_path, empty_xml)
    parsed = parse_docx(path)
    assert parsed.comments == []


def test_comment_dataclass_defaults() -> None:
    """Comment имеет разумные дефолты для всех полей."""
    c = Comment(id="x")
    assert c.id == "x"
    assert c.author == ""
    assert c.text == ""
    assert c.date == ""
    assert c.section_id is None


def test_comment_text_joins_multiple_paragraphs(tmp_path: Path) -> None:
    """Многопараграфный комментарий объединяется через \\n."""
    path = _make_docx_with_comments(tmp_path, _TWO_COMMENTS_XML)
    doc = parse_docx(path)
    c1 = doc.comments[1]
    # Два параграфа склеены через \n.
    assert "\n" in c1.text
    parts = c1.text.split("\n")
    assert any("Перепиши" in p for p in parts)
    assert any("длинный" in p for p in parts)
