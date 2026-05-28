"""Тесты режима настоящих OOXML-комментариев Word для аннотатора.

Стиль ``comments`` создаёт в .docx-архиве отдельную part ``word/comments.xml``
и вставляет в проблемные параграфы ``<w:commentRangeStart/End>`` плюс
reference-run. Здесь проверяем структуру архива и XML.
"""

from __future__ import annotations

import zipfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from lxml import etree

from gostforge.annotator import annotate_docx
from gostforge.model import Document, PageGeometry, PageSection
from gostforge.profile import load_profile

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NSMAP = {"w": _W_NS}


def _make_minimal_docx(path: Path, text: str = "Тестовый абзац.") -> None:
    """Создать минимальный .docx с одним параграфом."""
    doc = DocxDocument()
    doc.add_paragraph(text)
    doc.save(str(path))


@pytest.fixture()
def bad_margins_docx(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Path:
    """`.docx` плюс пропатченный parse_docx, выдающий заведомо плохие поля."""
    p = tmp_path / "bad.docx"
    _make_minimal_docx(p)

    def fake_parse(_: object) -> Document:
        doc = Document()
        doc.page_sections.append(
            PageSection(
                id="main",
                name="Основная часть",
                type="main",
                page=PageGeometry(margins_mm={"top": 25, "right": 15, "bottom": 20, "left": 30}),
            )
        )
        return doc

    monkeypatch.setattr("gostforge.annotator.docx_annotator.parse_docx", fake_parse)
    return p


def test_annotate_comments_creates_comments_part(bad_margins_docx: Path, tmp_path: Path) -> None:
    """После annotate_docx со style='comments' в .docx-архиве есть word/comments.xml."""
    out = tmp_path / "annotated.docx"
    profile = load_profile("gost-7.32-2017")
    annotate_docx(bad_margins_docx, out, profile, style="comments")

    with zipfile.ZipFile(out, "r") as z:
        names = z.namelist()
    assert "word/comments.xml" in names, f"В аннотированном .docx нет word/comments.xml: {names}"


def test_annotate_comments_content_has_violation_text(
    bad_margins_docx: Path, tmp_path: Path
) -> None:
    """В comments.xml есть <w:comment> с текстом, содержащим код F.01."""
    out = tmp_path / "annotated.docx"
    profile = load_profile("gost-7.32-2017")
    annotate_docx(bad_margins_docx, out, profile, style="comments")

    with zipfile.ZipFile(out, "r") as z:
        comments_xml = z.read("word/comments.xml")

    root = etree.fromstring(comments_xml)
    comments = root.findall(f"{{{_W_NS}}}comment")
    assert len(comments) > 0, "В comments.xml нет ни одного <w:comment>"

    all_text = "".join(t.text or "" for t in root.findall(f".//{{{_W_NS}}}t"))
    assert "F.01" in all_text, f"В тексте комментариев не найден код F.01: {all_text!r}"


def test_annotate_comments_inserts_reference_in_document(
    bad_margins_docx: Path, tmp_path: Path
) -> None:
    """В document.xml для каждого <w:comment> есть commentRangeStart/End + reference."""
    out = tmp_path / "annotated.docx"
    profile = load_profile("gost-7.32-2017")
    annotate_docx(bad_margins_docx, out, profile, style="comments")

    with zipfile.ZipFile(out, "r") as z:
        document_xml = z.read("word/document.xml")

    root = etree.fromstring(document_xml)
    starts = root.findall(f".//{{{_W_NS}}}commentRangeStart")
    ends = root.findall(f".//{{{_W_NS}}}commentRangeEnd")
    refs = root.findall(f".//{{{_W_NS}}}commentReference")

    assert len(starts) >= 1
    assert len(ends) == len(starts), (
        f"commentRangeStart={len(starts)} != commentRangeEnd={len(ends)}"
    )
    assert len(refs) == len(starts), (
        f"commentReference={len(refs)} != commentRangeStart={len(starts)}"
    )


def test_annotate_comments_updates_content_types(bad_margins_docx: Path, tmp_path: Path) -> None:
    """[Content_Types].xml содержит Override для /word/comments.xml."""
    out = tmp_path / "annotated.docx"
    profile = load_profile("gost-7.32-2017")
    annotate_docx(bad_margins_docx, out, profile, style="comments")

    with zipfile.ZipFile(out, "r") as z:
        content_types_xml = z.read("[Content_Types].xml")

    root = etree.fromstring(content_types_xml)
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    overrides = root.findall(f"{{{ct_ns}}}Override")
    part_names = [o.get("PartName") for o in overrides]
    assert "/word/comments.xml" in part_names, (
        f"Override для /word/comments.xml не найден: {part_names}"
    )


def test_annotate_comments_updates_document_rels(bad_margins_docx: Path, tmp_path: Path) -> None:
    """word/_rels/document.xml.rels содержит relationship типа comments."""
    out = tmp_path / "annotated.docx"
    profile = load_profile("gost-7.32-2017")
    annotate_docx(bad_margins_docx, out, profile, style="comments")

    with zipfile.ZipFile(out, "r") as z:
        rels_xml = z.read("word/_rels/document.xml.rels")

    root = etree.fromstring(rels_xml)
    rel_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    comments_rel_type = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
    )
    rels = root.findall(f"{{{rel_ns}}}Relationship")
    types = [r.get("Type") for r in rels]
    assert comments_rel_type in types, f"Relationship на comments.xml не найден: {types}"


def test_annotate_inline_does_not_create_comments_part(
    bad_margins_docx: Path, tmp_path: Path
) -> None:
    """В режиме inline word/comments.xml в архив не добавляется."""
    out = tmp_path / "annotated.docx"
    profile = load_profile("gost-7.32-2017")
    annotate_docx(bad_margins_docx, out, profile, style="inline")

    with zipfile.ZipFile(out, "r") as z:
        names = z.namelist()
    assert "word/comments.xml" not in names


def test_annotate_count_returns_number_of_comments(bad_margins_docx: Path, tmp_path: Path) -> None:
    """annotate_docx возвращает число вставленных комментариев = число <w:comment>."""
    out = tmp_path / "annotated.docx"
    profile = load_profile("gost-7.32-2017")
    n = annotate_docx(bad_margins_docx, out, profile, style="comments")

    with zipfile.ZipFile(out, "r") as z:
        comments_xml = z.read("word/comments.xml")
    root = etree.fromstring(comments_xml)
    comments = root.findall(f"{{{_W_NS}}}comment")
    assert n == len(comments)
    assert n > 0


def test_annotate_comments_default_style(bad_margins_docx: Path, tmp_path: Path) -> None:
    """Без явного style — используется 'comments' (создаётся comments.xml)."""
    out = tmp_path / "annotated.docx"
    profile = load_profile("gost-7.32-2017")
    annotate_docx(bad_margins_docx, out, profile)

    with zipfile.ZipFile(out, "r") as z:
        names = z.namelist()
    assert "word/comments.xml" in names


def test_annotate_comments_zero_violations_copies_input(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Если нарушений нет — annotate возвращает 0 и output = копия input."""
    p = tmp_path / "clean.docx"
    _make_minimal_docx(p)

    # Подменяем parse_docx так, чтобы validate ничего не вернул.
    def fake_parse(_: object) -> Document:
        return Document()

    monkeypatch.setattr("gostforge.annotator.docx_annotator.parse_docx", fake_parse)
    # И validate возвращает пустой список.
    monkeypatch.setattr("gostforge.annotator.docx_annotator.validate", lambda _doc, _prof: [])

    out = tmp_path / "out.docx"
    profile = load_profile("gost-7.32-2017")
    n = annotate_docx(p, out, profile, style="comments")
    assert n == 0
    assert out.exists()
    # word/comments.xml в архиве не должно быть, так как комментариев нет.
    with zipfile.ZipFile(out, "r") as z:
        assert "word/comments.xml" not in z.namelist()


def test_annotate_unknown_style_raises(bad_margins_docx: Path, tmp_path: Path) -> None:
    """Неизвестный style → ValueError."""
    out = tmp_path / "annotated.docx"
    profile = load_profile("gost-7.32-2017")
    with pytest.raises(ValueError, match="Unknown annotation style"):
        annotate_docx(bad_margins_docx, out, profile, style="bogus")  # type: ignore[arg-type]


def test_annotate_comments_reference_inside_paragraph(
    bad_margins_docx: Path, tmp_path: Path
) -> None:
    """commentRangeStart / End должны быть дочерними элементами <w:p>."""
    out = tmp_path / "annotated.docx"
    profile = load_profile("gost-7.32-2017")
    annotate_docx(bad_margins_docx, out, profile, style="comments")

    with zipfile.ZipFile(out, "r") as z:
        document_xml = z.read("word/document.xml")

    root = etree.fromstring(document_xml)
    # Каждый commentRangeStart должен иметь родителя <w:p>.
    starts = root.findall(f".//{{{_W_NS}}}commentRangeStart")
    assert len(starts) > 0
    for start in starts:
        parent = start.getparent()
        assert parent is not None
        assert parent.tag == f"{{{_W_NS}}}p", (
            f"commentRangeStart внутри {parent.tag}, ожидался <w:p>"
        )


def test_annotate_comments_unique_ids(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    """ID комментариев должны быть уникальны, даже если нарушений много."""
    p = tmp_path / "many.docx"
    _make_minimal_docx(p)

    def fake_parse(_: object) -> Document:
        doc = Document()
        # Несколько секций с разными бредовыми полями — даст несколько F.01-нарушений.
        for sid in ("main", "extra1", "extra2"):
            doc.page_sections.append(
                PageSection(
                    id=sid,
                    name=sid,
                    type="main",
                    page=PageGeometry(
                        margins_mm={
                            "top": 5,
                            "right": 5,
                            "bottom": 5,
                            "left": 5,
                        }
                    ),
                )
            )
        return doc

    monkeypatch.setattr("gostforge.annotator.docx_annotator.parse_docx", fake_parse)

    out = tmp_path / "out.docx"
    profile = load_profile("gost-7.32-2017")
    annotate_docx(p, out, profile, style="comments")

    with zipfile.ZipFile(out, "r") as z:
        comments_xml = z.read("word/comments.xml")

    root = etree.fromstring(comments_xml)
    ids = [c.get(f"{{{_W_NS}}}id") for c in root.findall(f"{{{_W_NS}}}comment")]
    assert len(ids) == len(set(ids)), f"ID комментариев не уникальны: {ids}"
