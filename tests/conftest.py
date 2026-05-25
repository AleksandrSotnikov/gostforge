"""Фабрики синтетических .docx-фикстур для тестов парсера и проверок.

Цель: каждый тест строит .docx программно через python-docx + lxml,
получает `pathlib.Path` к файлу во временной директории pytest и
прогоняет его через парсер / валидатор. Никаких бинарных фикстур
в репозитории.
"""

# ruff: noqa: RUF001, RUF002

from __future__ import annotations

import contextlib
from collections.abc import Iterable
from pathlib import Path

import docx
import pytest
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm, Pt
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def make_docx(
    path: Path,
    *,
    margins_mm: dict[str, float] | None = None,
    body_font: str = "Times New Roman",
    body_size: float = 14,
    line_spacing: float = 1.5,
    paragraphs: list[str] | None = None,
    headings: list[tuple[int, str]] | None = None,
    page_number: bool = True,
    page_number_start: int | None = None,
    headings_break_before: bool = False,
    title: str | None = None,
    author: str | None = None,
    tables: list[dict[str, object]] | None = None,
    figures: list[dict[str, object]] | None = None,
    bibliography: list[str] | None = None,
) -> Path:
    """Сгенерировать .docx по заданным параметрам и вернуть путь.

    Параметры:
      path:        куда сохранить файл.
      margins_mm:  поля {top,right,bottom,left} в мм. None → дефолты Word.
      body_font:   шрифт стиля Normal.
      body_size:   кегль стиля Normal (pt).
      line_spacing: межстрочный интервал стиля Normal (множитель).
      paragraphs:  список обычных абзацев.
      headings:    список (level, text) — порождают doc.add_heading.
      page_number: добавить ли поле PAGE в центральный параграф footer.
      page_number_start: если задано, в sectPr пишется <w:pgNumType w:start="N"/>
                          — стартовое значение нумерации страниц.
      headings_break_before: если True, у каждого заголовка ставим
                          paragraph_format.page_break_before = True.
      title/author: записать в docProps.core (если заданы).
      tables:      список таблиц вида
        {"caption": "Таблица 1 — ...", "headers": [...], "rows": [[...]]}.
        Caption (если задан) добавляется отдельным параграфом со стилем
        Caption НАД таблицей.
      figures:     список «рисунков» вида {"caption": "Рисунок 1 — ..."}.
        Имитируется через вставку пустого <w:drawing/> в run, чтобы парсер
        распознал параграф как рисунок. Caption (если задан) — отдельный
        параграф со стилем Caption под рисунком.
      bibliography: список строк. Если задан, в конец документа добавляется
        заголовок уровня 1 «Список использованных источников», а каждая
        строка превращается в отдельный обычный параграф (Фаза 1 — без
        нумерованного списка).
    """
    document = docx.Document()

    # --- метаданные ---
    if title is not None:
        document.core_properties.title = title
    if author is not None:
        document.core_properties.author = author

    # --- стиль Normal: шрифт/кегль/межстрочный интервал ---
    normal = document.styles["Normal"]
    normal.font.name = body_font
    normal.font.size = Pt(body_size)
    normal.paragraph_format.line_spacing = line_spacing

    # --- поля страницы ---
    section = document.sections[0]
    if margins_mm is not None:
        if "top" in margins_mm:
            section.top_margin = Mm(margins_mm["top"])
        if "right" in margins_mm:
            section.right_margin = Mm(margins_mm["right"])
        if "bottom" in margins_mm:
            section.bottom_margin = Mm(margins_mm["bottom"])
        if "left" in margins_mm:
            section.left_margin = Mm(margins_mm["left"])

    # --- заголовки и абзацы ---
    for level, text in headings or []:
        heading = document.add_heading(text, level=level)
        if headings_break_before:
            heading.paragraph_format.page_break_before = True
    for text in paragraphs or []:
        document.add_paragraph(text)

    # --- таблицы (caption сверху, если задан) ---
    for table_spec in tables or []:
        _append_table(document, table_spec)

    # --- рисунки (caption снизу, если задан) ---
    for figure_spec in figures or []:
        _append_figure(document, figure_spec)

    # --- список литературы (heading 1 + по параграфу на запись) ---
    if bibliography is not None:
        document.add_heading("Список использованных источников", level=1)
        for entry in bibliography:
            document.add_paragraph(entry)

    # --- номер страницы в footer (поле PAGE) ---
    if page_number:
        _inject_page_field_in_footer(section)

    # --- стартовая страница нумерации: <w:pgNumType w:start="N"/> в sectPr ---
    if page_number_start is not None:
        _inject_page_number_start(section, page_number_start)

    document.save(str(path))
    return path


def _append_table(document: object, spec: dict[str, object]) -> None:
    """Добавить таблицу в документ, опционально вставив подпись сверху со стилем Caption."""
    caption = spec.get("caption")
    headers_obj = spec.get("headers") or []
    rows_obj = spec.get("rows") or []
    if not isinstance(headers_obj, list) or not isinstance(rows_obj, list):
        raise TypeError("tables[*].headers и tables[*].rows должны быть list")

    if isinstance(caption, str) and caption:
        cap_para = document.add_paragraph(caption)  # type: ignore[attr-defined]
        with contextlib.suppress(KeyError):
            cap_para.style = document.styles["Caption"]  # type: ignore[attr-defined]

    n_cols = len(headers_obj) if headers_obj else (len(rows_obj[0]) if rows_obj else 0)
    n_rows = (1 if headers_obj else 0) + len(rows_obj)
    if n_rows == 0 or n_cols == 0:
        return

    table = document.add_table(rows=n_rows, cols=n_cols)  # type: ignore[attr-defined]
    row_idx = 0
    if headers_obj:
        for c, value in enumerate(headers_obj):
            table.rows[row_idx].cells[c].text = str(value)
        row_idx += 1
    for row in rows_obj:
        if not isinstance(row, list):
            raise TypeError("tables[*].rows[*] должен быть list")
        for c, value in enumerate(row):
            table.rows[row_idx].cells[c].text = str(value)
        row_idx += 1


def _append_figure(document: object, spec: dict[str, object]) -> None:
    """Добавить «рисунок»: параграф с пустым <w:drawing/> + опциональная подпись снизу."""
    para = document.add_paragraph()  # type: ignore[attr-defined]
    run = para.add_run()
    # Минимальный <w:drawing/> — парсеру достаточно факта присутствия тега.
    etree.SubElement(run._r, f"{{{W_NS}}}drawing")

    caption = spec.get("caption")
    if isinstance(caption, str) and caption:
        cap_para = document.add_paragraph(caption)  # type: ignore[attr-defined]
        with contextlib.suppress(KeyError):
            cap_para.style = document.styles["Caption"]  # type: ignore[attr-defined]


def _inject_page_field_in_footer(section: object) -> None:
    """Вставить <w:fldSimple w:instr="PAGE"/> в центральный параграф footer."""
    footer = section.footer  # type: ignore[attr-defined]
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fld = etree.SubElement(fp._p, f"{{{W_NS}}}fldSimple")
    fld.set(f"{{{W_NS}}}instr", "PAGE")
    r = etree.SubElement(fld, f"{{{W_NS}}}r")
    t = etree.SubElement(r, f"{{{W_NS}}}t")
    t.text = "1"


def _inject_page_number_start(section: object, start_value: int) -> None:
    """Записать <w:pgNumType w:start="N"/> в sectPr секции.

    python-docx не даёт прямого API к pgNumType, поэтому работаем через lxml.
    Если элемент уже есть — обновляем атрибут w:start.
    """
    sect_pr = section._sectPr  # type: ignore[attr-defined]
    pg_num_type = sect_pr.find(f"{{{W_NS}}}pgNumType")
    if pg_num_type is None:
        pg_num_type = etree.SubElement(sect_pr, f"{{{W_NS}}}pgNumType")
    pg_num_type.set(f"{{{W_NS}}}start", str(start_value))


# --- готовые фикстуры под типичные кейсы -------------------------------------


GOST_MARGINS: dict[str, float] = {"top": 20, "right": 15, "bottom": 20, "left": 30}

DEFAULT_PARAGRAPHS: list[str] = [
    "Настоящий отчёт описывает результаты выполненного исследования.",
    "Объект исследования — синтетический документ для проверки нормоконтроля.",
]

DEFAULT_HEADINGS: list[tuple[int, str]] = [
    (1, "Введение"),
    (1, "Заключение"),
    (1, "Список использованных источников"),
]


def _build_default(
    target: Path,
    *,
    margins_mm: dict[str, float] | None = None,
    body_font: str = "Times New Roman",
    body_size: float = 14,
    headings: Iterable[tuple[int, str]] | None = None,
    page_number: bool = True,
    page_number_start: int | None = None,
    headings_break_before: bool = False,
) -> Path:
    """Внутренний helper: документ по ГОСТу с разумными дефолтами."""
    return make_docx(
        target,
        margins_mm=margins_mm if margins_mm is not None else dict(GOST_MARGINS),
        body_font=body_font,
        body_size=body_size,
        paragraphs=DEFAULT_PARAGRAPHS,
        headings=list(headings) if headings is not None else list(DEFAULT_HEADINGS),
        page_number=page_number,
        page_number_start=page_number_start,
        headings_break_before=headings_break_before,
    )


@pytest.fixture
def correct_docx(tmp_path: Path) -> Path:
    """Документ полностью по ГОСТ 7.32-2017."""
    return _build_default(tmp_path / "correct.docx")


@pytest.fixture
def wrong_margins_docx(tmp_path: Path) -> Path:
    """Поле top=25 мм вместо 20 — нарушение F.01."""
    margins = dict(GOST_MARGINS)
    margins["top"] = 25
    return _build_default(tmp_path / "wrong_margins.docx", margins_mm=margins)


@pytest.fixture
def wrong_font_docx(tmp_path: Path) -> Path:
    """Шрифт Arial вместо Times New Roman — нарушение T.01."""
    return _build_default(tmp_path / "wrong_font.docx", body_font="Arial")


@pytest.fixture
def wrong_size_docx(tmp_path: Path) -> Path:
    """Кегль 12 вместо 14 — нарушение T.02."""
    return _build_default(tmp_path / "wrong_size.docx", body_size=12)


@pytest.fixture
def missing_intro_docx(tmp_path: Path) -> Path:
    """Документ без раздела «Введение» — нарушение S.01."""
    headings = [(1, "Заключение"), (1, "Список использованных источников")]
    return _build_default(tmp_path / "missing_intro.docx", headings=headings)


@pytest.fixture
def no_page_number_docx(tmp_path: Path) -> Path:
    """Документ без поля PAGE в footer — нарушение F.04."""
    return _build_default(tmp_path / "no_page_number.docx", page_number=False)


@pytest.fixture
def correct_numbering_docx(tmp_path: Path) -> Path:
    """Документ с корректной стартовой страницей нумерации (3) и разрывами перед разделами."""
    return _build_default(
        tmp_path / "correct_numbering.docx",
        page_number=True,
        page_number_start=3,
        headings_break_before=True,
    )


@pytest.fixture
def wrong_numbering_start_docx(tmp_path: Path) -> Path:
    """Старт нумерации 5 вместо ожидаемой 3 — нарушение F.06."""
    return _build_default(
        tmp_path / "wrong_numbering_start.docx",
        page_number=True,
        page_number_start=5,
        headings_break_before=True,
    )


@pytest.fixture
def no_page_break_docx(tmp_path: Path) -> Path:
    """Разделы 1 уровня без page_break_before — потенциальное нарушение S.06."""
    return _build_default(
        tmp_path / "no_page_break.docx",
        page_number=True,
        page_number_start=3,
        headings_break_before=False,
    )


@pytest.fixture
def table_with_caption_docx(tmp_path: Path) -> Path:
    """Таблица с подписью «Таблица 1 — ...» (стиль Caption) — для B.01."""
    return make_docx(
        tmp_path / "table_with_caption.docx",
        margins_mm=dict(GOST_MARGINS),
        paragraphs=["Текст перед таблицей."],
        tables=[
            {
                "caption": "Таблица 1 — Результаты эксперимента",
                "headers": ["Показатель", "Значение"],
                "rows": [["A", "1"], ["B", "2"]],
            }
        ],
        page_number=True,
    )


@pytest.fixture
def table_without_caption_docx(tmp_path: Path) -> Path:
    """Таблица без подписи — нарушение B.01."""
    return make_docx(
        tmp_path / "table_without_caption.docx",
        margins_mm=dict(GOST_MARGINS),
        paragraphs=["Текст перед таблицей."],
        tables=[
            {
                "headers": ["Показатель", "Значение"],
                "rows": [["A", "1"]],
            }
        ],
        page_number=True,
    )


@pytest.fixture
def figure_with_caption_docx(tmp_path: Path) -> Path:
    """Рисунок с подписью «Рисунок 1 — ...» (стиль Caption) — для I.01."""
    return make_docx(
        tmp_path / "figure_with_caption.docx",
        margins_mm=dict(GOST_MARGINS),
        paragraphs=["Текст перед рисунком."],
        figures=[{"caption": "Рисунок 1 — Схема алгоритма"}],
        page_number=True,
    )


@pytest.fixture
def figure_without_caption_docx(tmp_path: Path) -> Path:
    """Рисунок без подписи — нарушение I.01."""
    return make_docx(
        tmp_path / "figure_without_caption.docx",
        margins_mm=dict(GOST_MARGINS),
        paragraphs=["Текст перед рисунком."],
        figures=[{}],
        page_number=True,
    )


# --- список литературы -------------------------------------------------------


BIBLIOGRAPHY_CORRECT: list[str] = [
    "Иванов И. И. Основы программирования : учебник / И. И. Иванов. — Москва : Наука, 2020. — 320 с.",
    "Петров П. П. Анализ данных // Журнал прикладной информатики. — 2021. — № 3. — С. 15-27.",
    "ГОСТ 7.32-2017. Отчёт о научно-исследовательской работе. — Москва : Стандартинформ, 2017. — 32 с.",
    "Сидоров С. С. Электронный ресурс [Электронный ресурс]. — 2022. — URL: https://example.org/page (дата обращения: 01.05.2023).",
]

BIBLIOGRAPHY_BROKEN: list[str] = [
    # Без года издания.
    "Иванов И. И. Без года : монография / И. И. Иванов. — Москва : Наука. — 100 с.",
    # Без точки в конце.
    "Петров П. П. Без точки в конце / П. П. Петров. — Москва : Наука, 2020. — 50 с",
    # Слишком короткая запись.
    "Иванов 2020.",
    # web-запись без URL-маркера и без «дата обращения».
    "Сидоров С. С. Веб-ресурс без маркера : https://example.org/x, 2022. — С. 1.",
]


@pytest.fixture
def bibliography_correct_docx(tmp_path: Path) -> Path:
    """Документ с корректным разделом «Список использованных источников»."""
    return make_docx(
        tmp_path / "bibliography_correct.docx",
        margins_mm=dict(GOST_MARGINS),
        headings=[(1, "Введение")],
        paragraphs=["Текст введения."],
        bibliography=list(BIBLIOGRAPHY_CORRECT),
        page_number=True,
    )


@pytest.fixture
def bibliography_broken_docx(tmp_path: Path) -> Path:
    """Документ с разными нарушениями формата библиографических записей."""
    return make_docx(
        tmp_path / "bibliography_broken.docx",
        margins_mm=dict(GOST_MARGINS),
        headings=[(1, "Введение")],
        paragraphs=["Текст введения."],
        bibliography=list(BIBLIOGRAPHY_BROKEN),
        page_number=True,
    )
