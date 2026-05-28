"""Тесты многоуровневой шапки таблицы (`Table.extra_header_rows`).

Сценарий «двойная шапка» по ГОСТ Р 2.105: верхний ряд — группы колонок
(объединены через `merges` colspan), нижний ряд — подзаголовки.
"""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document as DocxDocument

from gostforge.exporter import export_docx
from gostforge.model import CellMerge, Document, LogicalSection, PageSection, Table, TextRun
from gostforge.profile import load_profile


def test_multi_header_writes_extra_rows_before_main_header(tmp_path: Path) -> None:
    """`extra_header_rows` пишутся НАД основной шапкой, шапка → 3 строки итого."""
    pytest.importorskip("docx")
    # Двойная шапка: «Группа 1 (colspan=2)» | «Группа 2 (colspan=2)»
    # сверху, а под ней — «A | B | C | D».
    table = Table(
        id="t1",
        caption=[TextRun(text="Таблица 1 — Тест")],
        extra_header_rows=[
            [
                [TextRun(text="Группа 1")],
                [TextRun(text="")],  # будет «съедена» colspan-ом
                [TextRun(text="Группа 2")],
                [TextRun(text="")],
            ]
        ],
        headers=[
            [TextRun(text="A")],
            [TextRun(text="B")],
            [TextRun(text="C")],
            [TextRun(text="D")],
        ],
        rows=[
            [
                [TextRun(text="1")],
                [TextRun(text="2")],
                [TextRun(text="3")],
                [TextRun(text="4")],
            ]
        ],
        merges=[
            CellMerge(row=0, col=0, rowspan=1, colspan=2),  # «Группа 1» — extra row 0
            CellMerge(row=0, col=2, rowspan=1, colspan=2),  # «Группа 2»
        ],
        number=1,
    )
    sec = LogicalSection(id="s1", level=1, heading=[TextRun(text="Глава")], children=[table])
    doc = Document(page_sections=[PageSection(id="p", name="main", type="main", content=[sec])])
    out = tmp_path / "multi_header.docx"
    # Отключаем continuation_caption — это поведение тестируется отдельно
    # в test_table_continuation.py. Здесь смотрим только extra_header_rows.
    profile = load_profile("gost-7.32-2017")
    profile.styles.table.continuation_caption = False
    export_docx(doc, profile, out)
    raw = DocxDocument(str(out))
    t = raw.tables[0]
    # 1 extra + 1 main header + 1 data = 3 строки.
    assert len(t.rows) == 3
    # Верхняя строка: «Группа 1» в col 0, «Группа 2» дальше (после colspan).
    row0 = t.rows[0]
    # Из-за colspan=2 ячеек физически меньше; читаем через .text.
    row0_texts = [cell.text for cell in row0.cells]
    # «Группа 1» должна присутствовать в первой объединённой ячейке.
    assert "Группа 1" in "".join(row0_texts)
    assert "Группа 2" in "".join(row0_texts)
    # Средняя строка — основная шапка.
    row1_texts = [cell.text for cell in t.rows[1].cells]
    assert row1_texts[:4] == ["A", "B", "C", "D"]
    # Нижняя — данные.
    row2_texts = [cell.text for cell in t.rows[2].cells]
    assert row2_texts[:4] == ["1", "2", "3", "4"]


def test_auto_merges_from_empty_cells() -> None:
    """В UI пустые ячейки склеиваются с левой соседней через `_auto_merges_from_extra_header_rows`."""
    pytest.importorskip("streamlit")
    from gostforge.web.builder_editor import _auto_merges_from_extra_header_rows

    merges = _auto_merges_from_extra_header_rows([["Группа 1", "", "Группа 2", ""]])
    assert merges == [
        {"row": 0, "col": 0, "rowspan": 1, "colspan": 2},
        {"row": 0, "col": 2, "rowspan": 1, "colspan": 2},
    ]


def test_auto_merges_handles_no_merges_needed() -> None:
    """Если в ряду нет пустых ячеек — merges пустой."""
    pytest.importorskip("streamlit")
    from gostforge.web.builder_editor import _auto_merges_from_extra_header_rows

    assert _auto_merges_from_extra_header_rows([["A", "B", "C"]]) == []


def test_multi_header_roundtrip_preserves_extra_rows(tmp_path: Path) -> None:
    """Round-trip: экспортируем → парсим обратно, multi-header сохраняется.

    Парсер ориентируется на `<w:tblHeader/>`, который ставит экспортёр
    на каждой строке шапки.
    """
    pytest.importorskip("docx")
    from gostforge.parser import parse_docx

    table = Table(
        id="t1",
        caption=[TextRun(text="Таблица 1 — Тест")],
        extra_header_rows=[
            [
                [TextRun(text="Группа 1")],
                [TextRun(text="")],
                [TextRun(text="Группа 2")],
                [TextRun(text="")],
            ]
        ],
        headers=[
            [TextRun(text="A")],
            [TextRun(text="B")],
            [TextRun(text="C")],
            [TextRun(text="D")],
        ],
        rows=[
            [
                [TextRun(text="1")],
                [TextRun(text="2")],
                [TextRun(text="3")],
                [TextRun(text="4")],
            ]
        ],
        merges=[
            CellMerge(row=0, col=0, rowspan=1, colspan=2),
            CellMerge(row=0, col=2, rowspan=1, colspan=2),
        ],
        number=1,
    )
    sec = LogicalSection(id="s1", level=1, heading=[TextRun(text="Глава")], children=[table])
    doc = Document(page_sections=[PageSection(id="p", name="main", type="main", content=[sec])])
    out = tmp_path / "rt.docx"
    # Отключаем continuation_caption — round-trip тест на extra_header_rows,
    # continuation проверяется отдельно.
    profile = load_profile("gost-7.32-2017")
    profile.styles.table.continuation_caption = False
    export_docx(doc, profile, out)

    reparsed = parse_docx(out)
    parsed_tables = [
        b
        for ps in reparsed.page_sections
        for s in ps.content
        if isinstance(s, LogicalSection)
        for b in s.children
        if isinstance(b, Table)
    ]
    assert parsed_tables, "должна остаться одна таблица после round-trip"
    t = parsed_tables[0]
    # extra_header_rows должен содержать одну строку (верхнюю «Группа 1 / Группа 2»).
    assert len(t.extra_header_rows) == 1
    assert len(t.headers) == 4  # «A B C D»
    assert len(t.rows) == 1  # одна строка данных


def test_parser_skips_continuation_table_row(tmp_path: Path) -> None:
    """Строка «Продолжение таблицы N» парсится обратно как маркер и в модель не попадает."""
    pytest.importorskip("docx")
    from gostforge.parser import parse_docx

    table = Table(
        id="t1",
        caption=[TextRun(text="Таблица 1 — Тест")],
        headers=[[TextRun(text="A")], [TextRun(text="B")]],
        rows=[[[TextRun(text="1")], [TextRun(text="2")]]],
        number=1,
    )
    sec = LogicalSection(id="s1", level=1, heading=[TextRun(text="Глава")], children=[table])
    doc = Document(page_sections=[PageSection(id="p", name="main", type="main", content=[sec])])
    out = tmp_path / "cont.docx"
    profile = load_profile("gost-7.32-2017")
    profile.styles.table.continuation_caption = True
    export_docx(doc, profile, out)

    reparsed = parse_docx(out)
    parsed_tables = [
        b
        for ps in reparsed.page_sections
        for s in ps.content
        if isinstance(s, LogicalSection)
        for b in s.children
        if isinstance(b, Table)
    ]
    t = parsed_tables[0]
    # Continuation-строка отфильтрована — в модели только обычная шапка + данные.
    assert t.extra_header_rows == []
    assert len(t.headers) == 2
    assert len(t.rows) == 1


def test_multi_header_back_compat_no_extra_rows(tmp_path: Path) -> None:
    """Таблица БЕЗ `extra_header_rows` ведёт себя как и раньше (одноуровневая шапка)."""
    pytest.importorskip("docx")
    table = Table(
        id="t1",
        caption=[TextRun(text="Таблица 1 — Тест")],
        headers=[[TextRun(text="X")], [TextRun(text="Y")]],
        rows=[[[TextRun(text="1")], [TextRun(text="2")]]],
        number=1,
    )
    sec = LogicalSection(id="s1", level=1, heading=[TextRun(text="Глава")], children=[table])
    doc = Document(page_sections=[PageSection(id="p", name="main", type="main", content=[sec])])
    out = tmp_path / "single_header.docx"
    # Отключаем continuation_caption — тест проверяет одноуровневую шапку.
    profile = load_profile("gost-7.32-2017")
    profile.styles.table.continuation_caption = False
    export_docx(doc, profile, out)
    raw = DocxDocument(str(out))
    t = raw.tables[0]
    # 1 шапка + 1 данные.
    assert len(t.rows) == 2
    assert [c.text for c in t.rows[0].cells] == ["X", "Y"]
