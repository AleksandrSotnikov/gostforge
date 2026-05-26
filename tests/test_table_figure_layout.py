"""Тесты исправлений трёх UX-багов: вёрстка таблиц, ширина рисунков,
неразрывность подписи."""

from __future__ import annotations

import re
import zipfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH

from gostforge.builder import work
from gostforge.exporter import export_docx
from gostforge.profile import load_profile


def _docx_xml(out: Path, part: str) -> str:
    with zipfile.ZipFile(out) as zf:
        return zf.read(part).decode("utf-8")


# --- 1. Форматирование текста в ячейках таблицы ---


def test_table_cell_no_first_line_indent(tmp_path: Path) -> None:
    """Ячейки таблицы НЕ должны наследовать красную строку 1.25 см от
    стиля Normal — в узких колонках это ломает читаемость."""
    b = (
        work("X", year=2026)
        .section("Введение")
        .table(
            headers=["A", "B"],
            rows=[["x", "y"]],
            caption="T",
        )
    )
    out = tmp_path / "tbl.docx"
    export_docx(b.build(), load_profile("gost-7.32-2017"), out)
    raw = DocxDocument(str(out))
    for t in raw.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    indent = p.paragraph_format.first_line_indent
                    # None или 0 — оба ok (не должно быть 1.25 см).
                    if indent is not None:
                        assert indent.cm <= 0.01, f"Красная строка в ячейке: {indent.cm} см"


def test_table_cell_single_line_spacing(tmp_path: Path) -> None:
    """В ячейках single-spacing (1.0), не 1.5 как в Normal."""
    b = work("X", year=2026).section("Введение").table(headers=["A"], rows=[["x"]], caption="T")
    out = tmp_path / "tbl2.docx"
    export_docx(b.build(), load_profile("gost-7.32-2017"), out)
    raw = DocxDocument(str(out))
    for t in raw.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    assert p.paragraph_format.line_spacing == 1.0


def test_table_header_centered_cells_left(tmp_path: Path) -> None:
    """Шапка по центру (header_alignment), ячейки слева (cell_alignment)."""
    b = (
        work("X", year=2026)
        .section("Введение")
        .table(headers=["A", "B"], rows=[["x", "y"]], caption="T")
    )
    out = tmp_path / "tbl3.docx"
    export_docx(b.build(), load_profile("gost-7.32-2017"), out)
    raw = DocxDocument(str(out))
    t = raw.tables[0]
    # Row 0 — header.
    for cell in t.rows[0].cells:
        assert cell.paragraphs[0].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER
    # Row 1 — data.
    for cell in t.rows[1].cells:
        assert cell.paragraphs[0].paragraph_format.alignment == WD_ALIGN_PARAGRAPH.LEFT


def test_table_cell_no_extra_spacing(tmp_path: Path) -> None:
    """Без интервалов перед/после параграфа в ячейках."""
    b = work("X", year=2026).section("Введение").table(headers=["A"], rows=[["x"]], caption="T")
    out = tmp_path / "tbl4.docx"
    export_docx(b.build(), load_profile("gost-7.32-2017"), out)
    raw = DocxDocument(str(out))
    t = raw.tables[0]
    for cell in t.rows[0].cells:
        for p in cell.paragraphs:
            sb = p.paragraph_format.space_before
            sa = p.paragraph_format.space_after
            assert sb is None or sb.pt == 0
            assert sa is None or sa.pt == 0


# --- 2. Ограничение ширины рисунка ---


def test_figure_max_width_default_165cm(tmp_path: Path) -> None:
    """max_width_cm default = 16.5 (ширина текста A4 при полях ГОСТа)."""
    profile = load_profile("gost-7.32-2017")
    assert profile.styles.figure.max_width_cm == 16.5


def test_figure_keep_with_next_default(tmp_path: Path) -> None:
    """keep_with_next=True по дефолту — рисунок не отрывается от подписи."""
    profile = load_profile("gost-7.32-2017")
    assert profile.styles.figure.keep_with_next is True


def _make_test_png(tmp_path: Path, width_px: int = 3000) -> Path:
    """Создать тестовое PNG большой ширины через PIL."""
    pytest.importorskip("PIL")
    from PIL import Image

    img = Image.new("RGB", (width_px, 200), color=(255, 0, 0))
    p = tmp_path / "big.png"
    img.save(p)
    return p


def test_figure_wider_than_max_is_resized(tmp_path: Path) -> None:
    """Картинка шире 16.5 см уменьшается до этой ширины."""
    png_path = _make_test_png(tmp_path, width_px=3000)
    b = (
        work("X", year=2026)
        .section("Введение")
        .figure(image_path=str(png_path), caption="Большой рисунок")
    )
    out = tmp_path / "fig.docx"
    export_docx(b.build(), load_profile("gost-7.32-2017"), out)
    raw = DocxDocument(str(out))
    images = raw.inline_shapes
    assert len(images) >= 1
    # 16.5 см = 5,940,000 EMU (1 cm = 360,000 EMU).
    max_emu = int(16.5 * 360000)
    for img in images:
        # Допуск 1% на округления.
        assert img.width <= max_emu + (max_emu // 100), (
            f"Картинка не уменьшена: width={img.width} EMU, max={max_emu}"
        )


def test_figure_smaller_than_max_stays(tmp_path: Path) -> None:
    """Маленькая картинка остаётся как есть, не растягивается."""
    png_path = _make_test_png(tmp_path, width_px=300)  # ~5 см
    b = (
        work("X", year=2026)
        .section("Введение")
        .figure(image_path=str(png_path), caption="Маленький")
    )
    out = tmp_path / "small.docx"
    export_docx(b.build(), load_profile("gost-7.32-2017"), out)
    raw = DocxDocument(str(out))
    images = raw.inline_shapes
    # 5 см << 16.5 см — должна остаться оригинальной.
    max_emu = int(16.5 * 360000)
    for img in images:
        assert img.width < max_emu


# --- 3. Подпись не отрывается от рисунка/таблицы ---


def test_figure_paragraph_has_keep_with_next(tmp_path: Path) -> None:
    """Параграф с рисунком имеет <w:keepNext/> — подпись не уезжает."""
    png_path = _make_test_png(tmp_path, width_px=500)
    b = work("X", year=2026).section("Введение").figure(image_path=str(png_path), caption="Схема")
    out = tmp_path / "fig-keep.docx"
    export_docx(b.build(), load_profile("gost-7.32-2017"), out)
    doc_xml = _docx_xml(out, "word/document.xml")
    # Найдём параграф с w:drawing (это рисунок), проверим что
    # в его pPr есть <w:keepNext/>.
    fig_p = re.search(
        r"<w:p\b(?:(?!</w:p>).)*<w:drawing>.*?</w:p>",
        doc_xml,
        re.DOTALL,
    )
    assert fig_p is not None, "Параграф с рисунком не найден"
    assert "<w:keepNext/>" in fig_p.group(0)


def test_caption_has_keep_together(tmp_path: Path) -> None:
    """Подпись имеет <w:keepLines/> — длинная подпись на одной странице."""
    b = (
        work("X", year=2026)
        .section("Введение")
        .table(
            headers=["A"],
            rows=[["x"]],
            caption="Длинная подпись таблицы для проверки",
        )
    )
    out = tmp_path / "cap.docx"
    export_docx(b.build(), load_profile("gost-7.32-2017"), out)
    doc_xml = _docx_xml(out, "word/document.xml")
    # Найдём параграф с «Таблица» в тексте.
    cap_p = re.search(
        r"<w:p\b(?:(?!</w:p>).)*Длинная подпись таблицы.*?</w:p>",
        doc_xml,
        re.DOTALL,
    )
    assert cap_p is not None
    assert "<w:keepLines/>" in cap_p.group(0)


def test_table_caption_has_keep_with_next(tmp_path: Path) -> None:
    """Подпись таблицы (position=above) имеет keepNext — таблица не
    уезжает на следующую страницу без подписи."""
    b = work("X", year=2026).section("Введение").table(headers=["A"], rows=[["x"]], caption="T")
    out = tmp_path / "cap-kn.docx"
    export_docx(b.build(), load_profile("gost-7.32-2017"), out)
    doc_xml = _docx_xml(out, "word/document.xml")
    cap_p = re.search(
        r"<w:p\b(?:(?!</w:p>).)*Таблица 1.*?</w:p>",
        doc_xml,
        re.DOTALL,
    )
    assert cap_p is not None
    assert "<w:keepNext/>" in cap_p.group(0)
