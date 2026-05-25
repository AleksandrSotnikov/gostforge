"""Тесты CLI-команды `gostforge new`."""

from __future__ import annotations

from pathlib import Path

from click.testing import CliRunner
from docx import Document as DocxDocument

from gostforge.cli import main


def test_new_creates_docx(tmp_path: Path) -> None:
    runner = CliRunner()
    out = tmp_path / "out.docx"
    result = runner.invoke(
        main,
        ["new", str(out), "--title", "T", "--year", "2026"],
    )
    assert result.exit_code == 0, result.output
    assert out.exists()
    # Документ читается python-docx и метаданные содержат title.
    docx = DocxDocument(out)
    assert docx.core_properties.title == "T"


def test_new_with_template_bachelor_thesis(tmp_path: Path) -> None:
    runner = CliRunner()
    out = tmp_path / "thesis.docx"
    result = runner.invoke(
        main,
        [
            "new",
            str(out),
            "--template",
            "bachelor_thesis",
            "--title",
            "ВКР",
            "--author",
            "Иванов И. И.",
            "--year",
            "2026",
        ],
    )
    assert result.exit_code == 0, result.output
    assert out.exists()
    docx = DocxDocument(out)
    texts = [p.text for p in docx.paragraphs]
    assert any("ГЛАВА 1" in t.upper() for t in texts)


def test_new_invalid_template_fails(tmp_path: Path) -> None:
    runner = CliRunner()
    out = tmp_path / "x.docx"
    result = runner.invoke(
        main,
        ["new", str(out), "--template", "nonsense", "--title", "T"],
    )
    assert result.exit_code != 0
    # click сам выводит сообщение об ошибке: «Invalid value for '--template'».
    assert "nonsense" in result.output or "Invalid" in result.output
