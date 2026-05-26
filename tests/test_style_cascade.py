# ruff: noqa: RUF001, RUF002, RUF003

"""Тесты на style-cascade в парсере и H.01/H.02-цвет.

Style-cascade = когда у run-а нет явных rPr-атрибутов (font, size,
bold, italic, color), парсер должен поднять их из стиля абзаца
(Heading 1, Normal) и его linked character-стиля (Heading1Char).

Это критично для проверок H.01/H.02 на документах, сгенерированных
через python-docx `add_heading()` — он не пишет явные run-атрибуты,
опирается на стиль Heading{N} (который в дефолтном шаблоне Word
синий Cambria).
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.shared import Pt

from gostforge.parser import parse_docx
from gostforge.profile import load_profile
from gostforge.validator import validate
from gostforge.validator.checks.headings import (
    _color_violates_expected,
    _describe_expected_color,
    all_logical_sections,
)


# --- _color_violates_expected — чистая функция ---


@pytest.mark.parametrize(
    "actual,expected,is_violation",
    [
        # Профиль ожидает auto → допустимы None или #000000.
        (None, None, False),
        (None, "auto", False),
        ("#000000", "auto", False),
        ("#000000", None, False),
        ("#365F91", "auto", True),
        ("#FF0000", None, True),
        # Профиль ожидает конкретный hex.
        ("#000000", "000000", False),
        ("#000000", "#000000", False),
        ("#ff0000", "FF0000", False),
        ("#FF0000", "ff0000", False),
        (None, "FF0000", True),
        ("#000000", "FF0000", True),
    ],
)
def test_color_violates_expected(
    actual: str | None, expected: str | None, is_violation: bool
) -> None:
    assert _color_violates_expected(actual, expected) is is_violation


def test_describe_expected_color_auto() -> None:
    assert "чёрный" in _describe_expected_color(None)
    assert "чёрный" in _describe_expected_color("auto")


def test_describe_expected_color_hex() -> None:
    assert "FF0000" in _describe_expected_color("FF0000")
    assert "FF0000" in _describe_expected_color("#FF0000")


# --- Style-cascade в парсере: bold/italic ---


def _make_docx_with_heading_style(
    tmp_path: Path,
    *,
    heading_color: str = "365F91",
    heading_bold: bool = True,
    heading_font: str = "Cambria",
) -> Path:
    """Создать .docx, где стиль Heading 1 имеет указанные параметры,
    а run-ы заголовка — без явных rPr (только наследование от стиля).
    """
    doc = DocxDocument()
    # Изменяем стиль Heading 1 напрямую — через python-docx.
    h1 = doc.styles["Heading 1"]
    h1.font.name = heading_font
    h1.font.size = Pt(14)
    h1.font.bold = heading_bold
    # Добавим явный цвет через XML (python-docx не имеет нормального API).
    from lxml import etree

    rPr = h1.element.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
    if rPr is None:
        rPr = etree.SubElement(
            h1.element,
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr",
        )
    color = rPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color")
    if color is None:
        color = etree.SubElement(
            rPr,
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color",
        )
    color.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
        heading_color,
    )

    # Добавляем содержимое.
    doc.add_heading("Введение", level=1)
    doc.add_paragraph("Текст параграфа.")

    out = tmp_path / "with-styled-heading.docx"
    doc.save(out)
    return out


def test_parser_inherits_color_from_heading_style(tmp_path: Path) -> None:
    """Run заголовка без явного цвета — парсер подтягивает из стиля."""
    path = _make_docx_with_heading_style(tmp_path, heading_color="365F91")
    parsed = parse_docx(path)
    sections = all_logical_sections(parsed)
    assert sections, "Заголовок не распознан"
    heading_runs = [el for el in sections[0].heading if hasattr(el, "color_hex")]
    assert heading_runs
    assert heading_runs[0].color_hex == "#365F91"


def test_parser_inherits_bold_from_heading_style(tmp_path: Path) -> None:
    """Run заголовка без явного bold — парсер подтягивает True из стиля."""
    path = _make_docx_with_heading_style(tmp_path, heading_bold=True)
    parsed = parse_docx(path)
    sections = all_logical_sections(parsed)
    heading_runs = [el for el in sections[0].heading if hasattr(el, "bold")]
    assert heading_runs[0].bold is True


def test_parser_inherits_font_from_heading_style(tmp_path: Path) -> None:
    """Run заголовка без явного font — парсер подтягивает из стиля."""
    path = _make_docx_with_heading_style(tmp_path, heading_font="Cambria")
    parsed = parse_docx(path)
    sections = all_logical_sections(parsed)
    heading_runs = [el for el in sections[0].heading if hasattr(el, "font")]
    assert heading_runs[0].font == "Cambria"


# --- H.01 теперь ловит синие заголовки из дефолтного шаблона ---


def test_h01_detects_blue_heading_from_default_template(tmp_path: Path) -> None:
    """Документ с синим Heading 1 (как дефолтный шаблон Word) даёт H.01."""
    path = _make_docx_with_heading_style(tmp_path, heading_color="365F91", heading_font="Cambria")
    parsed = parse_docx(path)
    parsed.profile_id = "gost-7.32-2017"
    profile = load_profile("gost-7.32-2017")
    violations = validate(parsed, profile)
    h01 = [v for v in violations if v.check_code == "H.01"]
    # Должны быть нарушения: цвет (синий) И шрифт (Cambria вместо TNR).
    color_v = [v for v in h01 if "365F91" in v.message]
    font_v = [v for v in h01 if "Cambria" in v.message]
    assert color_v, f"H.01 не нашёл синий цвет. Все H.01: {[v.message for v in h01]}"
    assert font_v, f"H.01 не нашёл шрифт Cambria. Все H.01: {[v.message for v in h01]}"


def test_h01_passes_when_heading_style_is_correct(tmp_path: Path) -> None:
    """Если Heading 1 чёрный TNR жирный — H.01 не срабатывает."""
    path = _make_docx_with_heading_style(
        tmp_path,
        heading_color="000000",
        heading_font="Times New Roman",
        heading_bold=True,
    )
    parsed = parse_docx(path)
    parsed.profile_id = "gost-7.32-2017"
    profile = load_profile("gost-7.32-2017")
    violations = validate(parsed, profile)
    h01_color = [v for v in violations if v.check_code == "H.01" and "цвет" in v.message.lower()]
    h01_font = [v for v in violations if v.check_code == "H.01" and "шрифт" in v.message.lower()]
    assert not h01_color, f"H.01 ложно сработал на color: {h01_color}"
    assert not h01_font, f"H.01 ложно сработал на font: {h01_font}"


# --- Регрессионный guard: builder→export не даёт цветовых нарушений ---


def test_builder_export_has_no_h01_color_violation(tmp_path: Path) -> None:
    """Сгенерированный конструктором документ не должен давать H.01-color."""
    from gostforge.builder import work
    from gostforge.exporter import export_docx

    b = (
        work("Демо", year=2026)
        .section("Реферат")
        .paragraph("текст реферата")
        .section("Содержание")
        .paragraph("содержание")
        .section("Введение")
        .paragraph("актуальность темы")
        .section("Заключение")
        .paragraph("итог работы")
        .section("Список использованных источников")
        .reference("Кнут Д. — М., 2007.")
    )
    doc = b.build()
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "demo.docx"
    export_docx(doc, profile, out)
    parsed = parse_docx(out)
    parsed.profile_id = "gost-7.32-2017"
    violations = validate(parsed, profile)
    h01_color = [v for v in violations if v.check_code == "H.01" and "цвет" in v.message.lower()]
    assert not h01_color, f"H.01 нашёл цветовое нарушение в сгенерированном docx: {h01_color}"
