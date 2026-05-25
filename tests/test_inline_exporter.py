# ruff: noqa: RUF001, RUF002, RUF003

"""Тесты экспортёра inline-элементов (Фаза 2.5, §4.5).

Проверяем расширение `_write_runs()` на все 4 типа `InlineElement`:
`TextRun`, `CrossRef`, `InlineFormula`, `Citation`. Поскольку парсер ещё
не доработан под inline-формулы и цитаты, проверки идут по сырому OOXML
из распакованного `.docx`-zip.
"""

from __future__ import annotations

import zipfile
from pathlib import Path

import docx as python_docx
from lxml import etree

from gostforge.exporter import export_docx
from gostforge.model import (
    BibliographyEntry,
    Citation,
    CrossRef,
    Document,
    InlineFormula,
    PageSection,
    Paragraph,
    TextRun,
)
from gostforge.profile import load_profile

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _wrap(content: list[object]) -> Document:
    """Завернуть произвольный inline-контент в Document с одним параграфом."""
    doc = Document()
    doc.page_sections.append(
        PageSection(
            id="main",
            name="m",
            type="main",
            content=[Paragraph(id="p1", content=content)],  # type: ignore[arg-type]
        )
    )
    return doc


def _document_xml(path: Path) -> bytes:
    """Вытащить word/document.xml из готового .docx-архива."""
    with zipfile.ZipFile(str(path)) as z:
        return z.read("word/document.xml")


def _first_paragraph_xml(path: Path) -> etree._Element:
    """Найти XML первого body-параграфа в out.docx."""
    root = etree.fromstring(_document_xml(path))
    body = root.find(f"{{{W_NS}}}body")
    assert body is not None
    paragraphs = body.findall(f"{{{W_NS}}}p")
    # Первый параграф в body — наш `Paragraph(id="p1")`.
    return paragraphs[0]


def test_text_run_underline_writes_w_u(tmp_path: Path) -> None:
    """TextRun(underline=True) → у run появляется <w:u w:val="single"/>."""
    doc = _wrap([TextRun(text="подчёркнуто", underline=True)])
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)

    raw = python_docx.Document(str(out))
    runs = raw.paragraphs[0].runs
    assert runs, "В параграфе не найдено ни одного run-а"
    # python-docx exposes underline через атрибут .underline (True/False/None).
    assert runs[0].underline is True


def test_text_run_color_hex_writes_color(tmp_path: Path) -> None:
    """TextRun(color_hex="#FF0000") → <w:color w:val="FF0000"/> (без #)."""
    doc = _wrap([TextRun(text="красный", color_hex="#FF0000")])
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)

    xml = _document_xml(out).decode("utf-8")
    # Ищем атрибут val без префикса '#'.
    assert 'w:val="FF0000"' in xml, f"Не найден ожидаемый w:color: {xml[:500]}"


def test_cross_ref_writes_fld_simple_with_ref_instr(tmp_path: Path) -> None:
    """CrossRef(target_id="fig-1") → <w:fldSimple w:instr=" REF fig-1 \\h "/>."""
    doc = _wrap([CrossRef(target_id="fig-1")])
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)

    p_xml = _first_paragraph_xml(out)
    flds = p_xml.findall(f"{{{W_NS}}}fldSimple")
    assert len(flds) == 1, f"Ожидался один fldSimple, найдено {len(flds)}"
    instr = flds[0].get(f"{{{W_NS}}}instr")
    assert instr is not None
    assert " REF fig-1 " in instr


def test_cross_ref_prefix_written_before_field(tmp_path: Path) -> None:
    """CrossRef(prefix=' (см. ', target_id='fig-1') → run с prefix ИДЁТ ПЕРЕД fldSimple."""
    doc = _wrap([CrossRef(target_id="fig-1", prefix=" (см. ")])
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)

    p_xml = _first_paragraph_xml(out)
    # Соберём детей в исходном порядке, фильтруя только содержательные.
    children = [c for c in p_xml if c.tag in (f"{{{W_NS}}}r", f"{{{W_NS}}}fldSimple")]
    assert len(children) >= 2, f"Ожидали как минимум run+fld, получили {[c.tag for c in children]}"
    # Первый — наш prefix-run с текстом " (см. "
    first = children[0]
    assert first.tag == f"{{{W_NS}}}r"
    texts = first.findall(f"{{{W_NS}}}t")
    assert any((t.text or "") == " (см. " for t in texts)
    # Следом идёт fldSimple
    assert children[1].tag == f"{{{W_NS}}}fldSimple"


def test_inline_formula_writes_omath_in_run(tmp_path: Path) -> None:
    """[TextRun('E='), InlineFormula('h\\nu')] → <m:oMath> ВНУТРИ <w:r>, а не рядом."""
    doc = _wrap([TextRun(text="E="), InlineFormula(latex="h\\nu")])
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)

    p_xml = _first_paragraph_xml(out)
    # Найдём все w:r, поищем тот, у которого внутри лежит m:oMath.
    runs = p_xml.findall(f"{{{W_NS}}}r")
    nested = [r for r in runs if r.find(f"{{{M_NS}}}oMath") is not None]
    assert len(nested) == 1, (
        "Ожидался ровно один <w:r>, содержащий <m:oMath>, "
        f"найдено {len(nested)}; структура: {[c.tag for c in p_xml]}"
    )
    # Содержимое формулы должно сохраниться.
    m_t = nested[0].find(f"{{{M_NS}}}oMath/{{{M_NS}}}r/{{{M_NS}}}t")
    assert m_t is not None
    assert m_t.text == "h\\nu"


def test_citation_renders_bracket_number(tmp_path: Path) -> None:
    """Citation на первый источник из bibliography → текст параграфа содержит '[1]'."""
    doc = _wrap([Citation(source_id="iv-23")])
    doc.bibliography.append(BibliographyEntry(id="iv-23", type="book", fields={"title": "Книга"}))
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)

    raw = python_docx.Document(str(out))
    assert "[1]" in raw.paragraphs[0].text


def test_citation_with_pages_uses_custom_template(tmp_path: Path) -> None:
    """Citation(pages='42', template='[{n}, с. {pages}]') → '[1, с. 42]'."""
    doc = _wrap([Citation(source_id="iv-23", pages="42", template="[{n}, с. {pages}]")])
    doc.bibliography.append(BibliographyEntry(id="iv-23", type="book", fields={"title": "Книга"}))
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)

    raw = python_docx.Document(str(out))
    assert "[1, с. 42]" in raw.paragraphs[0].text


def test_citation_unknown_source_id_falls_back_to_question(tmp_path: Path) -> None:
    """Citation на отсутствующий id при пустой bibliography → '[?]'."""
    doc = _wrap([Citation(source_id="missing")])
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)

    raw = python_docx.Document(str(out))
    assert "[?]" in raw.paragraphs[0].text


def test_multiple_inline_elements_preserve_order(tmp_path: Path) -> None:
    """[TextRun, CrossRef, TextRun, InlineFormula, TextRun] → 5 потомков в правильном порядке."""
    doc = _wrap(
        [
            TextRun(text="a"),
            CrossRef(target_id="x"),
            TextRun(text="b"),
            InlineFormula(latex="y"),
            TextRun(text="c"),
        ]
    )
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)

    p_xml = _first_paragraph_xml(out)
    # Берём только содержательные элементы (runs и fldSimple) в исходном порядке.
    children = [c for c in p_xml if c.tag in (f"{{{W_NS}}}r", f"{{{W_NS}}}fldSimple")]
    # Ожидаем 5 элементов: r("a"), fld(x), r("b"), r(<m:oMath y>), r("c").
    assert len(children) == 5, (
        f"Ожидали 5 элементов, получили {len(children)}: {[c.tag for c in children]}"
    )
    tags = [c.tag for c in children]
    assert tags[0] == f"{{{W_NS}}}r"
    assert tags[1] == f"{{{W_NS}}}fldSimple"
    assert tags[2] == f"{{{W_NS}}}r"
    assert tags[3] == f"{{{W_NS}}}r"  # run, обёртка вокруг m:oMath
    assert tags[4] == f"{{{W_NS}}}r"
    # У четвёртого run-а должна быть формула внутри.
    assert children[3].find(f"{{{M_NS}}}oMath") is not None
    # Тексты в r-нах 0, 2, 4 — "a", "b", "c" соответственно.
    assert children[0].find(f"{{{W_NS}}}t").text == "a"  # type: ignore[union-attr]
    assert children[2].find(f"{{{W_NS}}}t").text == "b"  # type: ignore[union-attr]
    assert children[4].find(f"{{{W_NS}}}t").text == "c"  # type: ignore[union-attr]


def test_bibliography_map_resets_after_export(tmp_path: Path) -> None:
    """После export с непустой bibliography второй export с пустой не утаскивает старые номера."""
    profile = load_profile("gost-7.32-2017")

    # Первый экспорт: bibliography из одного источника, цитата на него.
    doc1 = _wrap([Citation(source_id="iv-23")])
    doc1.bibliography.append(BibliographyEntry(id="iv-23", type="book", fields={"title": "Книга"}))
    out1 = tmp_path / "out1.docx"
    export_docx(doc1, profile, out1)

    # Проверяем, что модуль-level состояние очищено.
    from gostforge.exporter import docx_exporter as exp_mod

    assert exp_mod._current_bibliography_index is None, (
        "Bibliography-индекс не очистился между вызовами export_docx"
    )

    # Второй экспорт: bibliography пустая, та же цитата → должна стать "[?]",
    # а не унаследовать "[1]" из первого экспорта.
    doc2 = _wrap([Citation(source_id="iv-23")])
    out2 = tmp_path / "out2.docx"
    export_docx(doc2, profile, out2)

    raw2 = python_docx.Document(str(out2))
    assert "[?]" in raw2.paragraphs[0].text
    assert "[1]" not in raw2.paragraphs[0].text
    # И снова проверим, что состояние очищено после второго вызова.
    assert exp_mod._current_bibliography_index is None
