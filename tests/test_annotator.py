# ruff: noqa: RUF001, RUF002, RUF003
"""Тесты inline-аннотации .docx."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document as DocxDocument

from gostforge.annotator import annotate_docx
from gostforge.model import Document, PageGeometry, PageSection
from gostforge.profile import load_profile


def _make_minimal_docx(path: Path, text: str = "Тестовый абзац.") -> None:
    """Создать минимальный .docx с одним параграфом."""
    doc = DocxDocument()
    doc.add_paragraph(text)
    doc.save(str(path))


@pytest.fixture()
def clean_docx(tmp_path: Path) -> Path:
    p = tmp_path / "clean.docx"
    _make_minimal_docx(p)
    return p


@pytest.fixture()
def bad_margins_docx(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Path:
    """`.docx` плюс пропатченный parse_docx, выдающий заведомо плохие поля.

    Парсер в Фазе 0 — заглушка, поэтому подменяем его прямо в модуле
    annotator-а, чтобы валидатор увидел нарушение F.01.
    """
    p = tmp_path / "bad.docx"
    _make_minimal_docx(p)

    def fake_parse(_: object) -> Document:
        doc = Document()
        doc.page_sections.append(
            PageSection(
                id="main",
                name="Основная часть",
                type="main",
                page=PageGeometry(margins_mm={"top": 25, "right": 15, "bottom": 20, "left": 30}),
            )
        )
        return doc

    monkeypatch.setattr("gostforge.annotator.docx_annotator.parse_docx", fake_parse)
    return p


def test_annotate_creates_output_file(bad_margins_docx: Path, tmp_path: Path) -> None:
    out = tmp_path / "annotated.docx"
    profile = load_profile("gost-7.32-2017")
    annotate_docx(bad_margins_docx, out, profile, style="inline")
    assert out.exists()


def test_annotate_returns_count_positive(bad_margins_docx: Path, tmp_path: Path) -> None:
    out = tmp_path / "annotated.docx"
    profile = load_profile("gost-7.32-2017")
    n = annotate_docx(bad_margins_docx, out, profile, style="inline")
    assert n > 0


def test_annotate_inserts_marker_text(bad_margins_docx: Path, tmp_path: Path) -> None:
    out = tmp_path / "annotated.docx"
    profile = load_profile("gost-7.32-2017")
    annotate_docx(bad_margins_docx, out, profile, style="inline")

    annotated = DocxDocument(str(out))
    # В каком-то параграфе должен оказаться run с маркером, начинающимся с [F.01:.
    found = False
    for para in annotated.paragraphs:
        for run in para.runs:
            if run.text.startswith("[F.01:"):
                found = True
                # Маркер должен быть курсивом.
                assert run.italic is True
                break
        if found:
            break
    assert found, "В аннотированном документе не найден маркер [F.01:"


def test_annotate_returns_count_matches_violations(clean_docx: Path, tmp_path: Path) -> None:
    """Counts of inserted markers ≤ count of violations (some могут не разрешиться
    в конкретный параграф и попадают в первый параграф как fallback)."""
    from gostforge.parser import parse_docx
    from gostforge.validator import validate

    out = tmp_path / "annotated.docx"
    profile = load_profile("gost-7.32-2017")
    expected = len(validate(parse_docx(clean_docx), profile))
    n = annotate_docx(clean_docx, out, profile, style="inline")
    assert n == expected
    assert out.exists()


def test_annotate_raises_on_missing_input(tmp_path: Path) -> None:
    profile = load_profile("gost-7.32-2017")
    with pytest.raises(FileNotFoundError):
        annotate_docx(tmp_path / "nope.docx", tmp_path / "out.docx", profile, style="inline")
