"""Тесты конструктора работ."""

from __future__ import annotations

from pathlib import Path

import pytest

from gostforge.builder import WorkBuilder, work
from gostforge.model import (
    BibliographyEntry,
    Document,
    Figure,
    LogicalSection,
    Paragraph,
    Table,
    TextRun,
)
from gostforge.parser import parse_docx
from gostforge.profile import load_profile
from gostforge.validator import validate


def _heading_text(section: LogicalSection) -> str:
    return "".join(el.text for el in section.heading if isinstance(el, TextRun))


def _paragraph_text(para: Paragraph) -> str:
    return "".join(el.text for el in para.content if isinstance(el, TextRun))


def _top_level_sections(doc: Document) -> list[LogicalSection]:
    sections: list[LogicalSection] = []
    for ps in doc.page_sections:
        for child in ps.content:
            if isinstance(child, LogicalSection):
                sections.append(child)
    return sections


# --- Базовая сборка ----------------------------------------------------------


def test_work_builder_creates_document() -> None:
    doc = work("Title", author="Иванов И. И.", year=2026).build()
    assert isinstance(doc, Document)
    assert doc.metadata.title == "Title"
    assert doc.metadata.author == "Иванов И. И."
    assert doc.metadata.year == 2026
    assert doc.metadata.work_type == "coursework"


def test_section_adds_logical_section() -> None:
    doc = work("T").section("Введение").build()
    sections = _top_level_sections(doc)
    assert len(sections) == 1
    assert sections[0].level == 1
    assert _heading_text(sections[0]).upper() == "ВВЕДЕНИЕ"


def test_paragraph_adds_content() -> None:
    doc = work("T").section("Введение").paragraph("Текст параграфа").build()
    section = _top_level_sections(doc)[0]
    paragraphs = [c for c in section.children if isinstance(c, Paragraph)]
    assert len(paragraphs) == 1
    assert _paragraph_text(paragraphs[0]) == "Текст параграфа"


def test_chained_paragraphs_stay_in_same_section() -> None:
    doc = (
        work("T")
        .section("Введение")
        .paragraph("Первый")
        .paragraph("Второй")
        .paragraph("Третий")
        .build()
    )
    section = _top_level_sections(doc)[0]
    paragraphs = [c for c in section.children if isinstance(c, Paragraph)]
    assert [_paragraph_text(p) for p in paragraphs] == ["Первый", "Второй", "Третий"]


def test_subsection_nests_under_current_section() -> None:
    doc = (
        work("T")
        .section("Глава 1")
        .subsection("1.1 Подраздел")
        .paragraph("Текст подраздела")
        .build()
    )
    sections = _top_level_sections(doc)
    assert len(sections) == 1
    parent = sections[0]
    assert parent.level == 1
    subs = [c for c in parent.children if isinstance(c, LogicalSection)]
    assert len(subs) == 1
    assert subs[0].level == 2
    assert _heading_text(subs[0]) == "1.1 Подраздел"
    sub_paras = [c for c in subs[0].children if isinstance(c, Paragraph)]
    assert [_paragraph_text(p) for p in sub_paras] == ["Текст подраздела"]


def test_new_section_resets_paragraphs_target() -> None:
    doc = work("T").section("A").paragraph("x").section("B").paragraph("y").build()
    sections = _top_level_sections(doc)
    assert [_heading_text(s) for s in sections] == ["A", "B"]
    a_paras = [c for c in sections[0].children if isinstance(c, Paragraph)]
    b_paras = [c for c in sections[1].children if isinstance(c, Paragraph)]
    assert [_paragraph_text(p) for p in a_paras] == ["x"]
    assert [_paragraph_text(p) for p in b_paras] == ["y"]


# --- Сборка: разрыв страницы и нумерация -------------------------------------


def test_build_does_not_set_page_break_on_paragraphs() -> None:
    """Builder больше НЕ ставит page_break_before на первые параграфы
    разделов — разрыв страницы делается через стиль Heading 1 в
    экспортёре (profile.styles.heading_1.page_break_before).

    Раньше WorkBuilder.build() искал _find_first_paragraph для каждой
    секции и ставил ей page_break_before=True, но это давало баг для
    глав без вступительного текста: если глава начиналась сразу с
    подраздела 1.1 [текст], page-break оседал на первом параграфе
    ПОДРАЗДЕЛА, и текст уезжал на новую страницу после заголовка 1.1.

    Теперь page_break_before — атрибут стиля Heading 1, применяемый
    через _apply_heading_styles при экспорте.
    """
    doc = (
        work("T")
        .section("A")
        .paragraph("a1")
        .section("B")
        .paragraph("b1")
        .section("C")
        .paragraph("c1")
        .build()
    )
    sections = _top_level_sections(doc)
    for sec in sections:
        for child in sec.children:
            if isinstance(child, Paragraph):
                # Параграф не должен иметь явно установленного
                # page_break_before — он наследуется от стиля Heading 1
                # через применение к секции, не к параграфу.
                assert child.page_break_before is None or child.page_break_before is False


def test_build_sets_footer_and_start_value() -> None:
    doc = work("T").section("Введение").paragraph("...").build()
    assert len(doc.page_sections) == 1
    ps = doc.page_sections[0]
    assert ps.type == "main"
    assert ps.footer is not None
    assert ps.footer.default.center is not None
    center = ps.footer.default.center
    assert isinstance(center[0], TextRun)
    assert center[0].text == "{page}"
    assert ps.page_numbering.start_mode == "start_at"
    assert ps.page_numbering.start_value == 3


# --- Валидация ---------------------------------------------------------------


def test_built_document_passes_validation() -> None:
    """Документ конструктора не должен иметь error-нарушений.

    Сейчас активна только F.01 (поля страницы) — наш PageSection ставит
    правильные поля, поэтому ошибок быть не должно.
    """
    doc = (
        work("Курсовая", author="Иванов", year=2026)
        .section("Введение")
        .paragraph("Актуальность темы")
        .section("Заключение")
        .paragraph("Итог")
        .section("Список использованных источников")
        .build()
    )
    profile = load_profile("gost-7.32-2017")
    violations = validate(doc, profile)
    # K.01 на Фазе 1 builder не покрывает (создаёт только PageSection main,
    # а профиль ждёт title/frontmatter/appendix). save() её игнорирует.
    errors = [v for v in violations if v.severity == "error" and v.check_code != "K.01"]
    assert errors == [], f"Найдены ошибки: {[(v.check_code, v.message) for v in errors]}"


# --- Сохранение --------------------------------------------------------------


def test_save_writes_docx(tmp_path: Path) -> None:
    out = tmp_path / "out.docx"
    builder = (
        work("Курсовая по нормоконтролю", author="Иванов И. И.", year=2026)
        .section("Введение")
        .paragraph("Актуальность темы")
        .section("Заключение")
        .paragraph("Итог")
        .section("Список использованных источников")
        .root
    )
    builder.save(out, profile="gost-7.32-2017")
    assert out.exists()
    assert out.stat().st_size > 0
    parsed = parse_docx(out)
    # Парсер сейчас стаб: берёт title из имени файла. Проверяем, что .docx
    # вообще читается без исключений.
    assert parsed.metadata.title


def test_save_raises_on_validation_errors(tmp_path: Path) -> None:
    """Пустой WorkBuilder с неправильными полями страницы должен падать на validate.

    Чтобы это смоделировать на текущем наборе проверок (F.01), мутируем
    PageSection после build() и проверяем save() отдельно через прямой
    вызов экспортёра. Для самого save() — собираем builder и подменяем
    геометрию в build().
    """
    # Создаём корневой WorkBuilder и патчим его build(), чтобы сломать поля
    # страницы и тем самым спровоцировать F.01 на validate().
    root = work("Empty work")
    root.section("Введение").paragraph("...")
    out = tmp_path / "bad.docx"

    original_build = root.build

    def broken_build() -> Document:
        doc = original_build()
        doc.page_sections[0].page.margins_mm["top"] = 10.0  # неверно
        return doc

    root.build = broken_build  # type: ignore[method-assign]

    with pytest.raises(ValueError, match=r"F\.01"):
        root.save(out, profile="gost-7.32-2017")


# --- Таблицы, рисунки, ссылки ------------------------------------------------


def test_table_and_figure_added_to_section() -> None:
    doc = (
        work("T")
        .section("A")
        .table(headers=["h1", "h2"], rows=[["a", "b"], ["c", "d"]], caption="Test")
        .figure(image_path="x.png", caption="Schema")
        .build()
    )
    section = _top_level_sections(doc)[0]
    tables = [c for c in section.children if isinstance(c, Table)]
    figures = [c for c in section.children if isinstance(c, Figure)]
    assert len(tables) == 1
    assert len(figures) == 1
    assert tables[0].number == 1
    assert figures[0].number == 1
    assert isinstance(tables[0].caption[0], TextRun)
    assert tables[0].caption[0].text.startswith("Таблица 1 — ")
    assert isinstance(figures[0].caption[0], TextRun)
    assert figures[0].caption[0].text.startswith("Рисунок 1 — ")


def test_reference_in_bibliography_section() -> None:
    doc = (
        work("T")
        .section("Введение")
        .paragraph("Актуальность")
        .section("Список использованных источников")
        .reference("Иванов И. И. Программирование. — М. : Наука, 2023. — 320 с.")
        .reference("Петров П. П. Алгоритмы. — СПб. : Питер, 2024. — 450 с.", type="book")
        .build()
    )
    assert len(doc.bibliography) == 2
    assert all(isinstance(e, BibliographyEntry) for e in doc.bibliography)
    assert doc.bibliography[0].type == "book"
    assert "Иванов" in doc.bibliography[0].fields["raw"]
    assert "Петров" in doc.bibliography[1].fields["raw"]


def test_reference_outside_bibliography_section_raises() -> None:
    """reference() вне «Списка ...» должен бросать ValueError."""
    builder = work("T").section("Введение")
    with pytest.raises(ValueError, match="reference"):
        builder.reference("Какая-то книга 2024.")


# --- Фабричная функция и тип возвращаемого ----------------------------------


def test_work_factory_returns_work_builder() -> None:
    builder = work("Some")
    assert isinstance(builder, WorkBuilder)


def test_section_image_creates_figure_with_auto_caption() -> None:
    """`image()` создаёт Figure с автонумерованной подписью."""
    from gostforge.model import Figure

    doc = work("Test").section("Глава").image("/tmp/x.png", "Схема архитектуры").build()
    section = doc.page_sections[0].content[0]
    figs = [c for c in section.children if isinstance(c, Figure)]
    assert len(figs) == 1
    assert figs[0].image_path == "/tmp/x.png"
    caption_text = "".join(r.text for r in figs[0].caption)
    assert "Рисунок 1" in caption_text
    assert "Схема архитектуры" in caption_text


def test_section_list_ordered() -> None:
    """`list(items, ordered=True)` создаёт ListBlock с ordered=True."""
    from gostforge.model import ListBlock

    doc = work("Test").section("Глава").list(["один", "два", "три"], ordered=True).build()
    section = doc.page_sections[0].content[0]
    lists = [c for c in section.children if isinstance(c, ListBlock)]
    assert len(lists) == 1
    assert lists[0].ordered is True
    assert len(lists[0].items) == 3


def test_section_list_bulleted() -> None:
    from gostforge.model import ListBlock

    doc = work("Test").section("Глава").list(["A", "B"]).build()
    section = doc.page_sections[0].content[0]
    lists = [c for c in section.children if isinstance(c, ListBlock)]
    assert len(lists) == 1
    assert lists[0].ordered is False


# --- основная надпись (штамп ЕСКД) в конструкторе ---------------------------


def test_builder_title_block_sets_fields() -> None:
    """`.title_block(...)` проставляет штамп на PageSection с нужными графами."""
    from gostforge.model import TitleBlock

    doc = (
        work("Пояснительная записка", author="Иванов И. И.")
        .title_block(
            designation="АБВГ.123456.001 ПЗ",
            organization="Кафедра ИВТ",
            sheets_total="42",
            roles=[("Разраб.", "Иванов", "01.06.26"), ("Пров.", "Петров", "")],
        )
        .section("Введение")
        .paragraph("Текст введения по теме исследования и его актуальности.")
        .build()
    )
    tb = doc.page_sections[0].title_block
    assert isinstance(tb, TitleBlock)
    assert tb.enabled is True
    assert tb.designation == "АБВГ.123456.001 ПЗ"
    assert tb.organization == "Кафедра ИВТ"
    assert tb.sheets_total == "42"
    assert [r.role for r in tb.roles] == ["Разраб.", "Пров."]
    assert tb.roles[0].name == "Иванов"


def test_builder_title_block_default_roles() -> None:
    """Без roles берётся стандартный набор ролей ГОСТ 2.104."""
    doc = work("Работа").title_block(designation="X.001").section("Глава").build()
    tb = doc.page_sections[0].title_block
    assert tb is not None
    assert [r.role for r in tb.roles] == ["Разраб.", "Пров.", "Т.контр.", "Н.контр.", "Утв."]


def test_builder_without_title_block_is_none() -> None:
    """Без вызова .title_block() штамп не задаётся (None)."""
    doc = work("Работа").section("Глава").build()
    assert doc.page_sections[0].title_block is None


def test_builder_title_block_renders_in_export(tmp_path: Path) -> None:
    """Штамп из конструктора попадает в footer экспортированного .docx."""
    import docx as python_docx

    from gostforge.exporter import export_docx

    doc = (
        work("Пояснительная записка")
        .title_block(designation="X.001", organization="Каф.")
        .section("Введение")
        .paragraph("Текст.")
        .build()
    )
    out = tmp_path / "with-stamp.docx"
    export_docx(doc, load_profile("gost-7.32-2017"), out)
    tables = python_docx.Document(str(out)).sections[0].footer.tables
    assert len(tables) == 1
    text = "\n".join(c.text for r in tables[0].rows for c in r.cells)
    assert "X.001" in text


def test_builder_border_sets_page_border() -> None:
    """`.border(...)` включает рамку листа на PageGeometry."""
    from gostforge.model import PageBorder

    doc = work("Работа").border(size_eighth_pt=8).section("Глава").build()
    border = doc.page_sections[0].page.border
    assert isinstance(border, PageBorder)
    assert border.enabled is True
    assert border.size_eighth_pt == 8


def test_builder_without_border_is_none() -> None:
    """Без .border() рамка не задаётся (None) — приземлится профилем."""
    doc = work("Работа").section("Глава").build()
    assert doc.page_sections[0].page.border is None


def test_builder_border_renders_in_export(tmp_path: Path) -> None:
    """Рамка из конструктора попадает в sectPr экспортированного .docx."""
    import docx as python_docx

    from gostforge.exporter import export_docx

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    doc = work("Работа").border().section("Глава").paragraph("Текст.").build()
    out = tmp_path / "framed.docx"
    export_docx(doc, load_profile("gost-7.32-2017"), out)
    sect_pr = python_docx.Document(str(out)).sections[0]._sectPr
    assert sect_pr.find(f"{{{W_NS}}}pgBorders") is not None
