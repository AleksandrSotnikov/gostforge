"""Тесты сессии 3 плана развития: парсер чужих работ.

* Hyperlink — новый inline-элемент;
* FootnoteRef — ссылка на сноску из word/footnotes.xml;
* CellMerge — объединённые ячейки таблицы (<w:vMerge>, <w:gridSpan>).
"""

from __future__ import annotations

import zipfile
from pathlib import Path

from gostforge.exporter import export_docx
from gostforge.model import (
    CellMerge,
    Document,
    DocumentMetadata,
    FootnoteRef,
    Hyperlink,
    LogicalSection,
    PageGeometry,
    PageNumberingConfig,
    PageSection,
    Paragraph,
    Table,
    TextRun,
)
from gostforge.parser import parse_docx
from gostforge.profile import load_profile

# --- Hyperlink ---


def test_hyperlink_round_trip(tmp_path: Path) -> None:
    """Hyperlink в параграфе → экспорт → импорт сохраняет URL и text."""
    doc = Document(metadata=DocumentMetadata(title="X"))
    p = Paragraph(
        id="p1",
        content=[
            TextRun(text="См. документацию "),
            Hyperlink(url="https://gostforge.io/docs", text="gostforge.io"),
            TextRun(text=" — там подробности."),
        ],
        style_name="Normal",
    )
    sec = LogicalSection(id="s", heading=[TextRun(text="Введение")], level=1, children=[p])
    doc.page_sections.append(
        PageSection(
            id="m",
            name="N",
            type="main",
            page=PageGeometry(),
            page_numbering=PageNumberingConfig(),
            content=[sec],
        )
    )
    out = tmp_path / "hl.docx"
    export_docx(doc, load_profile("gost-7.32-2017"), out)
    parsed = parse_docx(out)
    # Найдём Hyperlink в parsed.
    found = []
    for ps in parsed.page_sections:
        for s in ps.content:
            if hasattr(s, "children"):
                for child in s.children:
                    if hasattr(child, "content"):
                        for el in child.content:
                            if isinstance(el, Hyperlink):
                                found.append(el)
    assert len(found) == 1
    assert found[0].url == "https://gostforge.io/docs"
    assert found[0].text == "gostforge.io"


def test_hyperlink_writes_w_hyperlink_in_xml(tmp_path: Path) -> None:
    """В document.xml появляется <w:hyperlink> с r:id."""
    doc = Document(metadata=DocumentMetadata(title="X"))
    p = Paragraph(
        id="p1",
        content=[Hyperlink(url="https://example.com", text="link")],
        style_name="Normal",
    )
    sec = LogicalSection(id="s", heading=[TextRun(text="X")], level=1, children=[p])
    doc.page_sections.append(
        PageSection(
            id="m",
            name="N",
            type="main",
            page=PageGeometry(),
            page_numbering=PageNumberingConfig(),
            content=[sec],
        )
    )
    out = tmp_path / "hl2.docx"
    export_docx(doc, load_profile("gost-7.32-2017"), out)
    with zipfile.ZipFile(out) as zf:
        doc_xml = zf.read("word/document.xml").decode("utf-8")
        rels = zf.read("word/_rels/document.xml.rels").decode("utf-8")
    assert "<w:hyperlink" in doc_xml
    assert "https://example.com" in rels


def test_hyperlink_with_anchor() -> None:
    """Внутренняя ссылка с anchor сохраняет anchor."""
    hl = Hyperlink(url="", text="Глава 1", anchor="ch1")
    # Тест на уровне модели.
    assert hl.anchor == "ch1"
    assert hl.text == "Глава 1"


# --- FootnoteRef ---


def test_footnote_ref_round_trip_with_text(tmp_path: Path) -> None:
    """Создаём минимальный docx с footnote, парсим — FootnoteRef.text заполнен."""
    # Создаём docx через python-docx + ручной патч zip с footnotes.xml.
    from docx import Document as DocxDocument

    d = DocxDocument()
    d.add_paragraph("Какой-то текст.")
    raw = tmp_path / "raw.docx"
    d.save(raw)

    # Patch zip с footnotes.xml и ссылкой на него.
    footnotes_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:footnote w:id="1">
    <w:p><w:r><w:t>Это сноска номер один.</w:t></w:r></w:p>
  </w:footnote>
</w:footnotes>
"""
    out = tmp_path / "with-fn.docx"
    with zipfile.ZipFile(raw) as zin:
        contents = {n: zin.read(n) for n in zin.namelist()}
    # [Content_Types]
    ct = contents["[Content_Types].xml"].decode("utf-8")
    if "footnotes.xml" not in ct:
        ct = ct.replace(
            "</Types>",
            '<Override PartName="/word/footnotes.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.'
            'wordprocessingml.footnotes+xml"/></Types>',
        )
    contents["[Content_Types].xml"] = ct.encode("utf-8")
    # rels
    rels_path = "word/_rels/document.xml.rels"
    rels = contents[rels_path].decode("utf-8")
    if "footnotes.xml" not in rels:
        rels = rels.replace(
            "</Relationships>",
            '<Relationship Id="rId200" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/'
            'relationships/footnotes" Target="footnotes.xml"/></Relationships>',
        )
    contents[rels_path] = rels.encode("utf-8")
    contents["word/footnotes.xml"] = footnotes_xml.encode("utf-8")
    # Также добавим footnoteReference в document.xml.
    doc_xml = contents["word/document.xml"].decode("utf-8")
    # Внутри первого <w:p>:
    doc_xml = doc_xml.replace(
        "<w:r><w:t>Какой-то текст.</w:t></w:r>",
        '<w:r><w:t>Какой-то текст.</w:t></w:r><w:r><w:footnoteReference w:id="1"/></w:r>',
    )
    contents["word/document.xml"] = doc_xml.encode("utf-8")
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in contents.items():
            zout.writestr(name, data)

    parsed = parse_docx(out)
    # Найдём FootnoteRef в параграфах.
    fn_refs = []
    for ps in parsed.page_sections:
        for sec in ps.content:
            if hasattr(sec, "children"):
                for child in sec.children:
                    if hasattr(child, "content"):
                        for el in child.content:
                            if isinstance(el, FootnoteRef):
                                fn_refs.append(el)
            elif hasattr(sec, "content"):
                for el in sec.content:
                    if isinstance(el, FootnoteRef):
                        fn_refs.append(el)
    assert len(fn_refs) == 1
    assert fn_refs[0].footnote_id == "1"
    assert "сноска" in fn_refs[0].text


def test_footnote_ref_empty_text_when_no_part() -> None:
    """FootnoteRef в документе без footnotes.xml имеет пустой text."""
    fn = FootnoteRef(footnote_id="1")
    assert fn.text == ""


# --- CellMerge ---


def test_cell_merge_dataclass_defaults() -> None:
    m = CellMerge(row=0, col=0)
    assert m.rowspan == 1
    assert m.colspan == 1


def test_table_with_horizontal_merge_writes_grid_span(tmp_path: Path) -> None:
    """colspan=2 → <w:gridSpan w:val="2"/> в OOXML."""
    doc = Document(metadata=DocumentMetadata(title="X"))
    t = Table(
        id="t",
        caption=[TextRun(text="Test")],
        headers=[[TextRun(text="A")], [TextRun(text="B")]],
        rows=[[[TextRun(text="1")], [TextRun(text="2")]]],
        merges=[CellMerge(row=0, col=0, rowspan=1, colspan=2)],
    )
    sec = LogicalSection(id="s", heading=[TextRun(text="X")], level=1, children=[t])
    doc.page_sections.append(
        PageSection(
            id="m",
            name="N",
            type="main",
            page=PageGeometry(),
            page_numbering=PageNumberingConfig(),
            content=[sec],
        )
    )
    out = tmp_path / "merge_h.docx"
    export_docx(doc, load_profile("gost-7.32-2017"), out)
    with zipfile.ZipFile(out) as zf:
        doc_xml = zf.read("word/document.xml").decode("utf-8")
    assert '<w:gridSpan w:val="2"/>' in doc_xml


def test_table_with_vertical_merge_writes_v_merge(tmp_path: Path) -> None:
    """rowspan=2 → <w:vMerge w:val="restart"/> + <w:vMerge/> continue."""
    doc = Document(metadata=DocumentMetadata(title="X"))
    t = Table(
        id="t",
        caption=[TextRun(text="Test")],
        headers=[[TextRun(text="A")], [TextRun(text="B")]],
        rows=[
            [[TextRun(text="1")], [TextRun(text="2")]],
            [[TextRun(text="3")], [TextRun(text="4")]],
        ],
        merges=[CellMerge(row=1, col=1, rowspan=2, colspan=1)],
    )
    sec = LogicalSection(id="s", heading=[TextRun(text="X")], level=1, children=[t])
    doc.page_sections.append(
        PageSection(
            id="m",
            name="N",
            type="main",
            page=PageGeometry(),
            page_numbering=PageNumberingConfig(),
            content=[sec],
        )
    )
    out = tmp_path / "merge_v.docx"
    export_docx(doc, load_profile("gost-7.32-2017"), out)
    with zipfile.ZipFile(out) as zf:
        doc_xml = zf.read("word/document.xml").decode("utf-8")
    assert '<w:vMerge w:val="restart"/>' in doc_xml
    assert "<w:vMerge/>" in doc_xml


def test_parser_extracts_grid_span(tmp_path: Path) -> None:
    """Парсер видит gridSpan и создаёт CellMerge."""
    doc = Document(metadata=DocumentMetadata(title="X"))
    t = Table(
        id="t",
        caption=[TextRun(text="Test")],
        headers=[[TextRun(text="A")], [TextRun(text="B")], [TextRun(text="C")]],
        rows=[[[TextRun(text="1")], [TextRun(text="2")], [TextRun(text="3")]]],
        merges=[CellMerge(row=0, col=0, rowspan=1, colspan=2)],
    )
    sec = LogicalSection(id="s", heading=[TextRun(text="X")], level=1, children=[t])
    doc.page_sections.append(
        PageSection(
            id="m",
            name="N",
            type="main",
            page=PageGeometry(),
            page_numbering=PageNumberingConfig(),
            content=[sec],
        )
    )
    out = tmp_path / "rt_merge.docx"
    export_docx(doc, load_profile("gost-7.32-2017"), out)
    parsed = parse_docx(out)
    found_table = None
    for ps in parsed.page_sections:
        for s in ps.content:
            if hasattr(s, "children"):
                for c in s.children:
                    if isinstance(c, Table):
                        found_table = c
    assert found_table is not None
    # Хотя бы один merge с colspan>1 должен быть.
    has_colspan = any(m.colspan > 1 for m in found_table.merges)
    assert has_colspan
