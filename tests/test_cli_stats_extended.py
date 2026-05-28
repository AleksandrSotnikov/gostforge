"""Тесты расширенных опций `gostforge stats` (--by-section, --json)."""

from __future__ import annotations

import json
from pathlib import Path

import pytest
from click.testing import CliRunner

from gostforge.cli import main

pytest.importorskip("docx")


def _make_minimal_docx(tmp_path: Path) -> Path:
    """Создать минимальный .docx через python-docx (без зависимости от builder)."""
    import docx

    out = tmp_path / "x.docx"
    d = docx.Document()
    d.add_heading("Глава 1", level=1)
    d.add_paragraph("Текст первого раздела с тремя словами.")
    d.add_heading("Глава 2", level=1)
    d.add_paragraph("Текст второго раздела чуть длиннее.")
    d.save(out)
    return out


def test_stats_default_output_shows_extended_metrics(tmp_path: Path) -> None:
    """`gostforge stats` без флагов показывает новые поля (avg, формулы и пр.)."""
    docx_path = _make_minimal_docx(tmp_path)
    runner = CliRunner()
    result = runner.invoke(main, ["stats", str(docx_path)])
    assert result.exit_code == 0, result.output
    assert "Параграфов всего" in result.output
    # Новые поля присутствуют.
    assert "Параграфов с inline-формулами" in result.output
    assert "Параграфов с перекр" in result.output
    assert "средняя длина в словах" in result.output


def test_stats_json_outputs_valid_json(tmp_path: Path) -> None:
    """`gostforge stats --json` выдаёт валидный JSON с дублированной структурой."""
    docx_path = _make_minimal_docx(tmp_path)
    runner = CliRunner()
    result = runner.invoke(main, ["stats", str(docx_path), "--json"])
    assert result.exit_code == 0, result.output
    payload = json.loads(result.output)
    assert "x.docx" in payload
    entry = payload["x.docx"]
    assert "total" in entry
    total = entry["total"]
    # Структурные поля.
    for key in (
        "paragraphs",
        "logical_sections_level_1",
        "avg_words_per_paragraph",
        "bibliography_by_type",
    ):
        assert key in total, f"Нет поля {key} в JSON-выводе"


def test_stats_by_section_lists_chapters(tmp_path: Path) -> None:
    """`gostforge stats --by-section` показывает разбивку по разделам."""
    docx_path = _make_minimal_docx(tmp_path)
    runner = CliRunner()
    result = runner.invoke(main, ["stats", str(docx_path), "--by-section"])
    assert result.exit_code == 0, result.output
    assert "По разделам:" in result.output
    assert "Глава 1" in result.output
    assert "Глава 2" in result.output


def test_stats_json_by_section_includes_per_section(tmp_path: Path) -> None:
    """`--json --by-section` — JSON содержит ключ `by_section` со списком разделов."""
    docx_path = _make_minimal_docx(tmp_path)
    runner = CliRunner()
    result = runner.invoke(main, ["stats", str(docx_path), "--json", "--by-section"])
    assert result.exit_code == 0, result.output
    payload = json.loads(result.output)
    entry = payload["x.docx"]
    assert "by_section" in entry
    sections = entry["by_section"]
    assert isinstance(sections, list)
    headings = [s["heading"] for s in sections]
    assert "Глава 1" in headings
    assert "Глава 2" in headings
    # У каждого раздела свой stats-dict.
    for sec in sections:
        assert "stats" in sec
        assert "paragraphs" in sec["stats"]
