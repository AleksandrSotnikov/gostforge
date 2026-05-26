# ruff: noqa: RUF002, RUF003

"""Тесты CLI-команды `gostforge fix`."""

from __future__ import annotations

import zipfile
from pathlib import Path

import docx as python_docx
from click.testing import CliRunner

from gostforge.cli import main
from gostforge.parser import parse_docx
from gostforge.validator import validate
from tests.conftest import make_docx


def _docx_with_double_spaces(path: Path) -> Path:
    """Создать .docx, в котором есть параграф с двойными пробелами."""
    return make_docx(
        path,
        paragraphs=["hello  world"],
        page_number=True,
    )


def test_cli_fix_creates_output(tmp_path: Path) -> None:
    """`gostforge fix` создаёт исправленный .docx без T.08-нарушений."""
    src = _docx_with_double_spaces(tmp_path / "in.docx")
    out = tmp_path / "out.docx"

    runner = CliRunner()
    result = runner.invoke(
        main,
        ["fix", str(src), "-o", str(out)],
        catch_exceptions=False,
    )

    assert result.exit_code == 0, result.output
    assert out.exists()

    # Перепарсим и проверим, что T.08-нарушений больше нет.
    from gostforge.profile import load_profile

    profile = load_profile("gost-7.32-2017")
    document = parse_docx(out)
    t08 = [v for v in validate(document, profile) if v.check_code == "T.08"]
    assert t08 == []


def test_cli_fix_dry_run_does_not_write(tmp_path: Path) -> None:
    """С флагом --dry-run output не создаётся."""
    src = _docx_with_double_spaces(tmp_path / "in.docx")
    out = tmp_path / "out.docx"

    runner = CliRunner()
    result = runner.invoke(
        main,
        ["fix", str(src), "-o", str(out), "--dry-run"],
        catch_exceptions=False,
    )

    assert result.exit_code == 0, result.output
    assert not out.exists()
    # В выводе должна быть метка dry-run
    assert "dry-run" in result.output


def test_fix_preserves_images(tmp_path: Path) -> None:
    """`gostforge fix` сохраняет реальные изображения через source_docx.

    Сценарий: .docx с реальным PNG → fix → output.docx должен содержать
    <w:drawing> и непустой word/media/* (а не [Рисунок: ...] placeholder).
    """
    try:
        from PIL import Image  # type: ignore[import-not-found]
    except ImportError:
        import pytest

        pytest.skip("Pillow не установлен — тест требует генерации PNG")

    img = tmp_path / "pixel.png"
    Image.new("RGB", (12, 12), color="blue").save(img)

    src = tmp_path / "in.docx"
    raw = python_docx.Document()
    raw.add_paragraph("Текст  до картинки.")  # двойной пробел — будет T.08-фикс
    raw.add_picture(str(img))
    raw.add_paragraph("Рисунок 1 — Пиксель")
    raw.save(str(src))

    out = tmp_path / "out.docx"
    runner = CliRunner()
    result = runner.invoke(
        main,
        ["fix", str(src), "-o", str(out)],
        catch_exceptions=False,
    )

    assert result.exit_code == 0, result.output
    assert out.exists()

    # В выходном .docx должен сохраниться рисунок: w:drawing + непустой
    # media-файл (PNG-blob), а не placeholder-текст.
    with zipfile.ZipFile(str(out)) as z:
        document_xml = z.read("word/document.xml").decode("utf-8")
        media_names = [n for n in z.namelist() if n.startswith("word/media/")]
        assert media_names, "Картинка не сохранилась: нет word/media/*"
        for name in media_names:
            assert len(z.read(name)) > 0

    assert "w:drawing" in document_xml
    assert "[Рисунок:" not in document_xml, (
        "В document.xml оказался placeholder — fix не передал source_docx"  # noqa: RUF001
    )


def test_cli_fix_only_applies_filter(tmp_path: Path) -> None:
    """`--only T.08` применяет только T.08, T.10 пропускается."""
    # В параграфе и двойной пробел (T.08), и парные кавычки (T.10).
    src = make_docx(
        tmp_path / "in.docx",
        paragraphs=['"a  b"'],
        page_number=True,
    )
    out = tmp_path / "out.docx"

    runner = CliRunner()
    result = runner.invoke(
        main,
        ["fix", str(src), "-o", str(out), "--only", "T.08"],
        catch_exceptions=False,
    )

    assert result.exit_code == 0, result.output
    assert out.exists()

    from gostforge.profile import load_profile

    profile = load_profile("gost-7.32-2017")
    document = parse_docx(out)
    violations = validate(document, profile)
    # T.08 пофиксен.
    assert not [v for v in violations if v.check_code == "T.08"]
    # T.10 — нет, кавычки должны остаться (нарушение присутствует).
    assert [v for v in violations if v.check_code == "T.10"]
    # И в stdout должна быть запись о T.08, и НЕ должно быть про T.10.
    assert "T.08" in result.output
    assert "T.10" not in result.output
