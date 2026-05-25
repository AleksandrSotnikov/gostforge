"""Тесты экспортёра модели в .docx."""

from pathlib import Path

import docx as python_docx

from gostforge.exporter import export_docx
from gostforge.model import (
    Document,
    Figure,
    LogicalSection,
    PageGeometry,
    PageSection,
    Paragraph,
    Table,
    TextRun,
)
from gostforge.profile import load_profile


def _minimal_doc() -> Document:
    """Документ с одним параграфом и одним заголовком."""
    doc = Document()
    intro = LogicalSection(
        id="intro",
        level=1,
        heading=[TextRun(text="Введение")],
        children=[
            Paragraph(
                id="p1",
                content=[TextRun(text="Это вводный абзац.")],
            )
        ],
    )
    doc.page_sections.append(
        PageSection(
            id="main",
            name="Основная часть",
            type="main",
            page=PageGeometry(),
            content=[intro],
        )
    )
    return doc


def test_export_creates_file(tmp_path: Path) -> None:
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(_minimal_doc(), profile, out)
    assert out.exists()
    assert out.stat().st_size > 0


def test_export_applies_page_margins(tmp_path: Path) -> None:
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(_minimal_doc(), profile, out)

    raw = python_docx.Document(str(out))
    section = raw.sections[0]
    # Margins возвращаются в EMU; .mm даёт float
    assert round(section.top_margin.mm) == int(profile.styles.page.margins_mm["top"])
    assert round(section.right_margin.mm) == int(profile.styles.page.margins_mm["right"])
    assert round(section.bottom_margin.mm) == int(profile.styles.page.margins_mm["bottom"])
    assert round(section.left_margin.mm) == int(profile.styles.page.margins_mm["left"])


def test_export_applies_normal_style(tmp_path: Path) -> None:
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(_minimal_doc(), profile, out)

    raw = python_docx.Document(str(out))
    normal = raw.styles["Normal"]
    assert normal.font.name == profile.styles.body.font
    assert normal.font.size.pt == profile.styles.body.size_pt


def test_export_writes_heading_and_paragraph(tmp_path: Path) -> None:
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(_minimal_doc(), profile, out)

    raw = python_docx.Document(str(out))
    texts = [p.text for p in raw.paragraphs]
    assert "Введение" in texts
    assert "Это вводный абзац." in texts


def test_export_preserves_bold_italic(tmp_path: Path) -> None:
    doc = Document()
    doc.page_sections.append(
        PageSection(
            id="main",
            name="m",
            type="main",
            content=[
                Paragraph(
                    id="p1",
                    content=[
                        TextRun(text="жирно", bold=True),
                        TextRun(text=" и "),
                        TextRun(text="курсивом", italic=True),
                    ],
                )
            ],
        )
    )
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)

    raw = python_docx.Document(str(out))
    runs = raw.paragraphs[0].runs
    assert runs[0].bold is True
    assert runs[2].italic is True


def test_export_applies_paragraph_level_alignment_and_break(tmp_path: Path) -> None:
    """Per-paragraph alignment, line_spacing, indent, page_break_before."""
    doc = Document()
    doc.page_sections.append(
        PageSection(
            id="main",
            name="m",
            type="main",
            content=[
                Paragraph(
                    id="p1",
                    content=[TextRun(text="Текст")],
                    style_name="Normal",
                    alignment="center",
                    line_spacing=2.0,
                    first_line_indent_cm=0.0,
                    page_break_before=True,
                )
            ],
        )
    )
    from gostforge.profile import load_profile
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)

    raw = python_docx.Document(str(out))
    p = raw.paragraphs[0]
    pf = p.paragraph_format
    assert pf.alignment == 1  # WD_ALIGN_PARAGRAPH.CENTER
    assert pf.line_spacing == 2.0
    assert pf.first_line_indent.cm == 0.0
    assert pf.page_break_before is True


def test_export_writes_table_with_caption(tmp_path: Path) -> None:
    doc = Document()
    table = Table(
        id="t1",
        caption=[TextRun(text="Таблица 1 — Результаты")],
        headers=[
            [TextRun(text="A")],
            [TextRun(text="B")],
        ],
        rows=[
            [[TextRun(text="1")], [TextRun(text="2")]],
            [[TextRun(text="3")], [TextRun(text="4")]],
        ],
    )
    doc.page_sections.append(
        PageSection(id="main", name="m", type="main", content=[table])
    )
    from gostforge.profile import load_profile
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)

    raw = python_docx.Document(str(out))
    # Над таблицей должен быть параграф-подпись
    texts = [p.text for p in raw.paragraphs]
    assert any("Таблица 1" in t for t in texts)
    # И сама таблица 3×2 (1 шапка + 2 строки)
    assert len(raw.tables) == 1
    docx_table = raw.tables[0]
    assert len(docx_table.rows) == 3
    assert len(docx_table.columns) == 2
    assert docx_table.rows[0].cells[0].text == "A"
    assert docx_table.rows[2].cells[1].text == "4"


def test_export_writes_figure_placeholder_with_caption(tmp_path: Path) -> None:
    doc = Document()
    fig = Figure(
        id="fig-1",
        caption=[TextRun(text="Рисунок 1 — Схема")],
    )
    doc.page_sections.append(
        PageSection(id="main", name="m", type="main", content=[fig])
    )
    from gostforge.profile import load_profile
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)

    raw = python_docx.Document(str(out))
    texts = [p.text for p in raw.paragraphs]
    # Плейсхолдер рисунка + подпись
    assert any("[Рисунок: fig-1]" in t for t in texts)
    assert any("Рисунок 1" in t for t in texts)


def test_export_writes_footer_with_page_field(tmp_path: Path) -> None:
    """Footer с {page}-плейсхолдером превращается в <w:fldSimple w:instr=PAGE/>."""
    from gostforge.model import ContentTemplate, HeaderConfig, PageNumberingConfig
    from gostforge.profile import load_profile
    doc = Document()
    doc.page_sections.append(
        PageSection(
            id="main",
            name="m",
            type="main",
            page_numbering=PageNumberingConfig(visible=True),
            footer=HeaderConfig(default=ContentTemplate(center=[TextRun(text="{page}")])),
        )
    )
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)

    # Round-trip через парсер: footer с PAGE-полем должен быть восстановлен.
    from gostforge.parser import parse_docx
    reparsed = parse_docx(out)
    section = reparsed.page_sections[0]
    assert section.page_numbering.visible is True
    assert section.footer is not None
    center = section.footer.default.center
    assert center is not None
    assert any(isinstance(r, TextRun) and "{page}" in r.text for r in center)


def test_export_writes_pgnumtype_start(tmp_path: Path) -> None:
    """start_mode=start_at, start_value=3 → <w:pgNumType w:start=3/>."""
    from gostforge.model import PageNumberingConfig
    from gostforge.profile import load_profile
    doc = Document()
    doc.page_sections.append(
        PageSection(
            id="main",
            name="m",
            type="main",
            page_numbering=PageNumberingConfig(
                visible=True, start_mode="start_at", start_value=3
            ),
        )
    )
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)

    from gostforge.parser import parse_docx
    reparsed = parse_docx(out)
    section = reparsed.page_sections[0]
    assert section.page_numbering.start_mode == "start_at"
    assert section.page_numbering.start_value == 3


def test_export_roundtrip_preserves_f04_and_f06(tmp_path: Path) -> None:
    """Полный round-trip: parse → export → parse не теряет F.04/F.06 информацию."""
    import sys
    sys.path.insert(0, str(Path(__file__).parent))
    from conftest import make_docx
    from gostforge.parser import parse_docx
    from gostforge.profile import load_profile
    from gostforge.validator import validate

    src = tmp_path / "in.docx"
    make_docx(
        src,
        margins_mm={"top": 20, "right": 15, "bottom": 20, "left": 30},
        body_font="Times New Roman",
        body_size=14,
        headings=[
            (1, "ВВЕДЕНИЕ"),
            (1, "ЗАКЛЮЧЕНИЕ"),
            (1, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"),
        ],
        paragraphs=["Текст работы."],
        page_number=True,
        page_number_start=3,
    )
    profile = load_profile("gost-7.32-2017")
    doc = parse_docx(src)
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)
    reparsed = parse_docx(out)
    errors = [v for v in validate(reparsed, profile) if v.severity == "error"]
    # F.04 (поле PAGE) и F.06 (start=3) теперь должны переноситься через export.
    codes = {v.check_code for v in errors}
    assert "F.04" not in codes
    assert "F.06" not in codes


def test_export_writes_pgnumtype_format(tmp_path: Path) -> None:
    """Экспортёр пишет w:fmt в pgNumType, если формат отличается от arabic."""
    from gostforge.model import PageNumberingConfig
    from gostforge.parser import parse_docx
    from gostforge.profile import load_profile
    doc = Document()
    doc.page_sections.append(
        PageSection(
            id="main",
            name="m",
            type="main",
            page_numbering=PageNumberingConfig(visible=True, format="roman"),
        )
    )
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)
    reparsed = parse_docx(out)
    assert reparsed.page_sections[0].page_numbering.format == "roman"


def test_export_roundtrip_paper_size_and_orientation(tmp_path: Path) -> None:
    """Round-trip: A3 landscape должен восстанавливаться."""
    from gostforge.model import PageGeometry
    from gostforge.parser import parse_docx
    from gostforge.profile import load_profile
    doc = Document()
    doc.page_sections.append(
        PageSection(
            id="main",
            name="m",
            type="main",
            page=PageGeometry(paper="A3", orientation="landscape"),
        )
    )
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)
    reparsed = parse_docx(out)
    page = reparsed.page_sections[0].page
    assert page.paper == "A3"
    assert page.orientation == "landscape"


def test_export_writes_listblock_paragraphs(tmp_path: Path) -> None:
    """ListBlock пишется как параграфы со стилем List Number или префиксами."""
    from gostforge.model import ListBlock
    from gostforge.profile import load_profile

    doc = Document()
    block = ListBlock(
        id="l1",
        ordered=True,
        items=[[TextRun(text="первый")], [TextRun(text="второй")]],
    )
    doc.page_sections.append(
        PageSection(id="main", name="m", type="main", content=[block])
    )
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)

    raw = python_docx.Document(str(out))
    texts = [p.text for p in raw.paragraphs]
    # либо чистый текст элементов (если стиль есть), либо с префиксом «1. »
    assert any("первый" in t for t in texts)
    assert any("второй" in t for t in texts)


def test_export_inserts_real_picture_when_image_exists(tmp_path: Path) -> None:
    """Если Figure.image_path указывает на существующий файл, вставляется <w:drawing>."""
    from gostforge.model import Figure
    from gostforge.profile import load_profile

    try:
        from PIL import Image  # type: ignore[import-not-found]
    except ImportError:
        import pytest
        pytest.skip("Pillow не установлен — тест требует генерации PNG")
    img = tmp_path / "pixel.png"
    Image.new("RGB", (10, 10), color="red").save(img)

    doc = Document()
    fig = Figure(
        id="fig-1",
        image_path=str(img),
        caption=[TextRun(text="Рисунок 1 — Пиксель")],
    )
    doc.page_sections.append(
        PageSection(id="main", name="m", type="main", content=[fig])
    )
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)

    # В document.xml должен быть w:drawing
    import zipfile
    with zipfile.ZipFile(str(out)) as z:
        document_xml = z.read("word/document.xml").decode("utf-8")
    assert "w:drawing" in document_xml


def test_export_roundtrip_formula(tmp_path: Path) -> None:
    """Round-trip Formula → export → parse сохраняет latex и number."""
    from gostforge.model import Formula
    from gostforge.parser import parse_docx
    from gostforge.profile import load_profile

    doc = Document()
    fml = Formula(id="formula-1", latex="E=mc^2", number=3)
    doc.page_sections.append(
        PageSection(id="main", name="m", type="main", content=[fml])
    )
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)

    reparsed = parse_docx(out)
    formulas = [b for b in reparsed.page_sections[0].content if isinstance(b, Formula)]
    assert len(formulas) == 1
    assert "E=mc^2" in formulas[0].latex
    assert formulas[0].number == 3


def test_export_formula_without_number(tmp_path: Path) -> None:
    """Формула без номера — никаких текстовых хвостов."""
    from gostforge.model import Formula
    from gostforge.profile import load_profile

    doc = Document()
    fml = Formula(id="formula-1", latex="x+y", number=None)
    doc.page_sections.append(
        PageSection(id="main", name="m", type="main", content=[fml])
    )
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)

    raw = python_docx.Document(str(out))
    # Должен быть параграф с центрированием, содержащий формулу,
    # но без хвостовой текстовой части вида "(N)"
    paragraphs = raw.paragraphs
    assert len(paragraphs) >= 1
    # текст параграфа не содержит "(" если нет номера
    assert "(" not in paragraphs[0].text


def test_export_with_source_docx_preserves_image(tmp_path: Path) -> None:
    """Round-trip: создаём .docx с PNG → parse → export(source_docx=src) → в out есть реальный media."""
    from gostforge.parser import parse_docx
    from gostforge.profile import load_profile

    try:
        from PIL import Image  # type: ignore[import-not-found]
    except ImportError:
        import pytest
        pytest.skip("Pillow не установлен — тест требует генерации PNG")

    img = tmp_path / "pixel.png"
    Image.new("RGB", (16, 16), color="green").save(img)

    src = tmp_path / "in.docx"
    raw = python_docx.Document()
    raw.add_paragraph("Перед рисунком.")
    raw.add_picture(str(img))
    raw.add_paragraph("Рисунок 1 — Тест")
    raw.save(str(src))

    document = parse_docx(src)
    figures = [
        item
        for ps in document.page_sections
        for item in ps.content
        if hasattr(item, "image_path")
    ]
    assert figures, "Парсер не нашёл ни одного рисунка"
    assert figures[0].image_path.startswith("embedded:"), (
        f"Ожидали embedded:rIdN, получили {figures[0].image_path!r}"
    )

    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(document, profile, out, source_docx=src)

    import zipfile
    with zipfile.ZipFile(str(out)) as z:
        names = z.namelist()
        document_xml = z.read("word/document.xml").decode("utf-8")
        media_names = [n for n in names if n.startswith("word/media/")]
        assert media_names, f"В выходном .docx нет media: {names}"
        for media_name in media_names:
            assert len(z.read(media_name)) > 0, f"Media {media_name} пустой"

    assert "w:drawing" in document_xml
    assert "[Рисунок:" not in document_xml, (
        "В document.xml оказался placeholder — экспортёр не воспользовался source_docx"
    )


def test_export_with_embedded_path_without_source_docx_uses_placeholder(tmp_path: Path) -> None:
    """Если image_path начинается с embedded:, но source_docx не задан, рисуем placeholder."""
    from gostforge.model import Figure
    from gostforge.profile import load_profile

    doc = Document()
    fig = Figure(
        id="fig-7",
        image_path="embedded:rId99",
        caption=[TextRun(text="Рисунок 7 — Должен стать placeholder-ом")],
    )
    doc.page_sections.append(
        PageSection(id="main", name="m", type="main", content=[fig])
    )
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "out.docx"
    export_docx(doc, profile, out)

    raw = python_docx.Document(str(out))
    texts = [p.text for p in raw.paragraphs]
    assert any("[Рисунок: fig-7]" in t for t in texts)
