"""Тесты парсера .docx → Document."""

# ruff: noqa: RUF001, RUF002, RUF003

from __future__ import annotations

from pathlib import Path

from gostforge.model import Figure, LogicalSection, Paragraph, Table, TextRun
from gostforge.parser.docx_parser import parse_docx

from .conftest import make_docx


def test_parse_extracts_metadata(tmp_path: Path) -> None:
    """Title из docProps попадает в Document.metadata.title."""
    path = make_docx(
        tmp_path / "with_title.docx",
        paragraphs=["Один абзац"],
        title="Моя курсовая работа",
        author="Студент",
    )
    doc = parse_docx(path)
    assert doc.metadata.title == "Моя курсовая работа"
    assert doc.metadata.author == "Студент"
    # year проставляется из core.created — это либо None, либо валидный
    # четырёхзначный год (python-docx по умолчанию ставит 2013).
    assert doc.metadata.year is None or doc.metadata.year >= 2000


def test_parse_extracts_margins(tmp_path: Path) -> None:
    """Поля страницы в мм извлекаются корректно."""
    path = make_docx(
        tmp_path / "margins.docx",
        margins_mm={"top": 20, "right": 15, "bottom": 20, "left": 30},
        paragraphs=["Текст"],
    )
    doc = parse_docx(path)
    assert len(doc.page_sections) == 1
    margins = doc.page_sections[0].page.margins_mm
    assert margins["top"] == 20.0
    assert margins["right"] == 15.0
    assert margins["bottom"] == 20.0
    assert margins["left"] == 30.0


def test_parse_extracts_font_and_size(tmp_path: Path) -> None:
    """Шрифт и кегль из стиля Normal видны в content рунов."""
    path = make_docx(
        tmp_path / "font.docx",
        body_font="Times New Roman",
        body_size=14,
        paragraphs=["Тестовый абзац для проверки шрифта"],
    )
    doc = parse_docx(path)
    runs: list[TextRun] = []
    for item in doc.page_sections[0].content:
        if isinstance(item, Paragraph):
            for elem in item.content:
                if isinstance(elem, TextRun):
                    runs.append(elem)
    assert any(r.font == "Times New Roman" and r.size_pt == 14.0 for r in runs), (
        f"Не найден run с TNR/14, runs={[(r.font, r.size_pt) for r in runs]}"
    )


def test_parse_extracts_headings(tmp_path: Path) -> None:
    """Каждый Heading становится LogicalSection с соответствующим текстом."""
    path = make_docx(
        tmp_path / "headings.docx",
        headings=[(1, "Введение"), (1, "Заключение")],
        paragraphs=[],
    )
    doc = parse_docx(path)
    sections = [item for item in doc.page_sections[0].content if isinstance(item, LogicalSection)]
    assert len(sections) == 2

    def heading_text(s: LogicalSection) -> str:
        return "".join(e.text for e in s.heading if isinstance(e, TextRun))

    titles = [heading_text(s) for s in sections]
    assert "Введение" in titles
    assert "Заключение" in titles
    assert all(s.level == 1 for s in sections)


def test_parse_detects_page_number_field(tmp_path: Path) -> None:
    """page_number=True → footer.default.center содержит TextRun('{page}')."""
    path_with = make_docx(
        tmp_path / "with_page.docx",
        paragraphs=["Хоть какой-то текст"],
        page_number=True,
    )
    doc_with = parse_docx(path_with)
    footer = doc_with.page_sections[0].footer
    assert footer is not None, "Footer должен быть распознан"
    center = footer.default.center
    assert center is not None
    assert any(isinstance(e, TextRun) and e.text == "{page}" for e in center)
    assert doc_with.page_sections[0].page_numbering.visible is True


def test_parse_no_page_number_field(tmp_path: Path) -> None:
    """page_number=False → footer не содержит маркера {page}."""
    path_without = make_docx(
        tmp_path / "no_page.docx",
        paragraphs=["Текст без номера"],
        page_number=False,
    )
    doc_without = parse_docx(path_without)
    footer = doc_without.page_sections[0].footer
    if footer is not None:
        center = footer.default.center or []
        assert not any(isinstance(e, TextRun) and e.text == "{page}" for e in center)


def test_parse_paragraph_after_heading_goes_into_children(tmp_path: Path) -> None:
    """Абзацы после заголовка кладутся в children соответствующей LogicalSection."""
    path = make_docx(
        tmp_path / "structure.docx",
        headings=[(1, "Введение")],
        paragraphs=["Первый абзац введения."],
    )
    doc = parse_docx(path)
    content = doc.page_sections[0].content
    sections = [c for c in content if isinstance(c, LogicalSection)]
    assert len(sections) == 1
    # Один параграф должен попасть внутрь раздела.
    children = sections[0].children
    paragraphs_inside = [c for c in children if isinstance(c, Paragraph)]
    assert len(paragraphs_inside) == 1
    text = "".join(e.text for e in paragraphs_inside[0].content if isinstance(e, TextRun))
    assert "Первый абзац" in text


def test_parse_extracts_page_number_start(tmp_path: Path) -> None:
    """<w:pgNumType w:start="3"/> → start_mode='start_at', start_value=3."""
    path = make_docx(
        tmp_path / "pgnum_start.docx",
        paragraphs=["Текст"],
        page_number=True,
        page_number_start=3,
    )
    doc = parse_docx(path)
    numbering = doc.page_sections[0].page_numbering
    assert numbering.start_mode == "start_at"
    assert numbering.start_value == 3


def test_parse_no_page_number_start_keeps_default(tmp_path: Path) -> None:
    """Без pgNumType — start_mode остаётся 'continue', start_value=None."""
    path = make_docx(
        tmp_path / "pgnum_default.docx",
        paragraphs=["Текст"],
        page_number=True,
    )
    doc = parse_docx(path)
    numbering = doc.page_sections[0].page_numbering
    assert numbering.start_mode == "continue"
    assert numbering.start_value is None


def test_parse_extracts_page_break_before(tmp_path: Path) -> None:
    """headings_break_before=True → у первого Paragraph в разделе page_break_before=True."""
    path = make_docx(
        tmp_path / "page_break.docx",
        headings=[(1, "Введение"), (1, "Заключение")],
        paragraphs=[],
        headings_break_before=True,
    )
    doc = parse_docx(path)
    sections = [
        item for item in doc.page_sections[0].content if isinstance(item, LogicalSection)
    ]
    # В make_docx абзацы paragraphs идут после всех заголовков, поэтому здесь
    # children секций пусты. Проверяем page_break_before на самом заголовке
    # отдельно через прямое чтение — но в текущей модели заголовок хранится
    # как list[InlineElement], а не Paragraph. Поэтому проверка идёт через
    # отдельный параграф под заголовком: добавим параграф после первого
    # заголовка через add_paragraph — но фабрика этого не делает.
    # Достаточно: оба раздела распознаны, и атрибут page_break_before
    # в модели достижим. Проверка фактического выставления — ниже.
    assert len(sections) == 2

    # Альтернативный путь: документ с одним заголовком и абзацем НЕ под ним.
    # Так как DEFAULT-фабрика этого не делает, построим вручную.
    path2 = make_docx(
        tmp_path / "page_break_para.docx",
        headings=[],
        paragraphs=[],
        headings_break_before=True,
    )
    # Открываем повторно и добавляем параграф с pageBreakBefore через python-docx.
    import docx as _docx

    d = _docx.Document(str(path2))
    para = d.add_paragraph("Параграф с разрывом")
    para.paragraph_format.page_break_before = True
    d.save(str(path2))
    doc2 = parse_docx(path2)

    paragraphs = [
        item for item in doc2.page_sections[0].content if isinstance(item, Paragraph)
    ]
    assert len(paragraphs) >= 1
    assert paragraphs[-1].page_break_before is True


def test_parse_paragraph_without_page_break(tmp_path: Path) -> None:
    """Обычный параграф без pageBreakBefore → page_break_before is None."""
    path = make_docx(
        tmp_path / "no_break.docx",
        paragraphs=["Просто абзац без разрыва."],
        headings=[],
    )
    doc = parse_docx(path)
    paragraphs = [
        item for item in doc.page_sections[0].content if isinstance(item, Paragraph)
    ]
    assert len(paragraphs) == 1
    # Атрибут либо не задан явно (None), либо False — главное, что не True.
    assert paragraphs[0].page_break_before is not True


def test_parse_extracts_table_with_caption(tmp_path: Path) -> None:
    """Таблица с подписью «Таблица N — ...»: parser создаёт Table с caption."""
    path = make_docx(
        tmp_path / "table.docx",
        paragraphs=[],
        tables=[
            {
                "caption": "Таблица 1 — Результаты",
                "headers": ["Показатель", "Значение"],
                "rows": [["A", "1"], ["B", "2"]],
            }
        ],
    )
    doc = parse_docx(path)
    content = doc.page_sections[0].content
    tables = [b for b in content if isinstance(b, Table)]
    assert len(tables) == 1
    table = tables[0]

    cap_text = "".join(e.text for e in table.caption if isinstance(e, TextRun))
    assert "Таблица 1" in cap_text

    # Подпись-параграф НЕ должен остаться отдельным элементом.
    paragraphs_left = [b for b in content if isinstance(b, Paragraph)]
    assert all("Таблица 1" not in _para_text(p) for p in paragraphs_left)

    # Заголовки и строки разложены.
    assert len(table.headers) == 2
    header_texts = ["".join(e.text for e in cell if isinstance(e, TextRun)) for cell in table.headers]
    assert header_texts == ["Показатель", "Значение"]
    assert len(table.rows) == 2
    first_row = ["".join(e.text for e in cell if isinstance(e, TextRun)) for cell in table.rows[0]]
    assert first_row == ["A", "1"]


def test_parse_extracts_figure_with_caption(tmp_path: Path) -> None:
    """Рисунок с подписью «Рисунок N — ...»: parser создаёт Figure с caption."""
    path = make_docx(
        tmp_path / "figure.docx",
        paragraphs=[],
        figures=[{"caption": "Рисунок 1 — Схема алгоритма"}],
    )
    doc = parse_docx(path)
    content = doc.page_sections[0].content
    figures = [b for b in content if isinstance(b, Figure)]
    assert len(figures) == 1
    figure = figures[0]
    cap_text = "".join(e.text for e in figure.caption if isinstance(e, TextRun))
    assert "Рисунок 1" in cap_text

    # Подпись-параграф НЕ должен остаться отдельным элементом.
    paragraphs_left = [b for b in content if isinstance(b, Paragraph)]
    assert all("Рисунок 1" not in _para_text(p) for p in paragraphs_left)


def test_parse_figure_without_caption(tmp_path: Path) -> None:
    """Рисунок без подписи — Figure создаётся, caption пуст."""
    path = make_docx(
        tmp_path / "figure_no_caption.docx",
        paragraphs=[],
        figures=[{}],
    )
    doc = parse_docx(path)
    figures = [b for b in doc.page_sections[0].content if isinstance(b, Figure)]
    assert len(figures) == 1
    assert figures[0].caption == []


def test_parse_keeps_order_of_paragraphs_and_tables(tmp_path: Path) -> None:
    """Порядок paragraphs + tables в исходнике сохраняется в content."""
    path = make_docx(
        tmp_path / "order.docx",
        paragraphs=["alpha", "beta"],
        tables=[{"headers": ["X"], "rows": [["1"]]}],
    )
    doc = parse_docx(path)
    content = doc.page_sections[0].content
    # Ожидаем: Paragraph(alpha), Paragraph(beta), Table.
    types = [type(b).__name__ for b in content]
    assert types == ["Paragraph", "Paragraph", "Table"]
    assert _para_text(content[0]) == "alpha"  # type: ignore[arg-type]
    assert _para_text(content[1]) == "beta"  # type: ignore[arg-type]


def _para_text(p: Paragraph) -> str:
    return "".join(e.text for e in p.content if isinstance(e, TextRun))


# --- список литературы -------------------------------------------------------


def test_parse_extracts_bibliography(tmp_path: Path) -> None:
    """Раздел «Список использованных источников» собирается в Document.bibliography."""
    entries = [
        "Иванов И. И. Основы / И. И. Иванов. — Москва : Наука, 2020. — 320 с.",
        "Петров П. П. Анализ // Журнал. — 2021. — № 3. — С. 15-27.",
        "ГОСТ 7.32-2017. Отчёт о НИР. — Москва : Стандартинформ, 2017. — 32 с.",
    ]
    path = make_docx(
        tmp_path / "bib.docx",
        headings=[(1, "Введение")],
        paragraphs=["Текст введения."],
        bibliography=entries,
    )
    doc = parse_docx(path)
    assert len(doc.bibliography) == 3
    raws = [e.fields["raw"] for e in doc.bibliography]
    assert raws == entries
    ids = [e.id for e in doc.bibliography]
    assert ids == ["ref-1", "ref-2", "ref-3"]


def test_parse_detects_bibliography_types(tmp_path: Path) -> None:
    """Тип записи определяется эвристикой: book / web / standard."""
    entries = [
        # book — обычная монография без особенных маркеров.
        "Иванов И. И. Основы / И. И. Иванов. — Москва : Наука, 2020. — 320 с.",
        # web — содержит https://.
        "Сидоров С. С. Ресурс. — 2022. — URL: https://example.org (дата обращения: 01.05.2023).",
        # standard — начинается с ГОСТ.
        "ГОСТ 7.32-2017. Отчёт о НИР. — Москва : Стандартинформ, 2017. — 32 с.",
    ]
    path = make_docx(
        tmp_path / "bib_types.docx",
        headings=[(1, "Введение")],
        paragraphs=["Текст введения."],
        bibliography=entries,
    )
    doc = parse_docx(path)
    types = [e.type for e in doc.bibliography]
    assert types == ["book", "web", "standard"]


def test_parse_bibliography_paragraphs_kept_in_section(tmp_path: Path) -> None:
    """Параграфы раздела «Список источников» НЕ дублируются в content страницы.

    В Document.bibliography они должны быть, и в LogicalSection.children — тоже.
    Но в `page_section.content` отдельных Paragraph-ов с этим текстом быть
    не должно — они лежат внутри LogicalSection.
    """
    entries = [
        "Иванов И. И. Основы / И. И. Иванов. — Москва : Наука, 2020. — 320 с.",
        "Петров П. П. Анализ // Журнал. — 2021. — № 3. — С. 15-27.",
    ]
    path = make_docx(
        tmp_path / "bib_no_dup.docx",
        headings=[(1, "Введение")],
        paragraphs=["Текст введения."],
        bibliography=entries,
    )
    doc = parse_docx(path)
    assert len(doc.bibliography) == 2

    page_content = doc.page_sections[0].content
    # На верхнем уровне content не должно быть Paragraph с библиографическим
    # текстом — они должны лежать внутри LogicalSection «Список ...».
    top_paragraphs = [b for b in page_content if isinstance(b, Paragraph)]
    for p in top_paragraphs:
        text = _para_text(p)
        assert "Иванов И. И. Основы" not in text
        assert "Петров П. П. Анализ" not in text

    # А в children самого раздела «Список ...» — есть.
    bib_sections = [
        item
        for item in page_content
        if isinstance(item, LogicalSection)
        and "источник" in "".join(e.text for e in item.heading if isinstance(e, TextRun)).lower()
    ]
    assert len(bib_sections) == 1
    bib_paragraphs = [c for c in bib_sections[0].children if isinstance(c, Paragraph)]
    bib_texts = [_para_text(p) for p in bib_paragraphs]
    assert any("Иванов И. И. Основы" in t for t in bib_texts)
    assert any("Петров П. П. Анализ" in t for t in bib_texts)


def test_parser_extracts_header_with_page_field(tmp_path: Path) -> None:
    """Если PAGE-поле в header (а не footer) — парсер кладёт {page} в page_section.header."""
    import docx as python_docx
    from docx.oxml.ns import qn  # type: ignore[import-not-found]
    from lxml import etree

    doc = python_docx.Document()
    section = doc.sections[0]
    header = section.header
    p = header.add_paragraph()
    p.paragraph_format.alignment = python_docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    # Вставляем <w:fldSimple w:instr="PAGE"/> прямо в lxml-узел параграфа.
    fld = etree.SubElement(p._p, qn("w:fldSimple"))
    fld.set(qn("w:instr"), "PAGE")
    run = etree.SubElement(fld, qn("w:r"))
    rt = etree.SubElement(run, qn("w:t"))
    rt.text = ""
    out = tmp_path / "h.docx"
    doc.save(str(out))

    from gostforge.parser import parse_docx
    parsed = parse_docx(out)
    page_section = parsed.page_sections[0]
    assert page_section.header is not None
    center = page_section.header.default.center
    assert center is not None
    assert any(getattr(r, "text", "") == "{page}" for r in center)
    assert page_section.page_numbering.visible is True


def test_parser_recognizes_numbered_list(tmp_path: Path) -> None:
    """Параграфы со стилем 'List Number' группируются в один ListBlock(ordered=True)."""
    import docx as python_docx

    from gostforge.model import ListBlock
    from gostforge.parser import parse_docx

    doc = python_docx.Document()
    doc.add_paragraph("Первый элемент", style="List Number")
    doc.add_paragraph("Второй элемент", style="List Number")
    doc.add_paragraph("Третий элемент", style="List Number")
    out = tmp_path / "list.docx"
    doc.save(str(out))

    parsed = parse_docx(out)
    blocks = parsed.page_sections[0].content
    list_blocks = [b for b in blocks if isinstance(b, ListBlock)]
    assert len(list_blocks) == 1
    assert list_blocks[0].ordered is True
    assert len(list_blocks[0].items) == 3


def test_parser_recognizes_bulleted_list(tmp_path: Path) -> None:
    """Параграфы со стилем 'List Bullet' дают ListBlock(ordered=False)."""
    import docx as python_docx

    from gostforge.model import ListBlock
    from gostforge.parser import parse_docx

    doc = python_docx.Document()
    doc.add_paragraph("Пункт A", style="List Bullet")
    doc.add_paragraph("Пункт B", style="List Bullet")
    out = tmp_path / "bullet.docx"
    doc.save(str(out))

    parsed = parse_docx(out)
    blocks = parsed.page_sections[0].content
    list_blocks = [b for b in blocks if isinstance(b, ListBlock)]
    assert len(list_blocks) == 1
    assert list_blocks[0].ordered is False
    assert len(list_blocks[0].items) == 2


def test_parser_splits_lists_by_intervening_paragraph(tmp_path: Path) -> None:
    """Между двумя списками обычный параграф — должны получить 2 ListBlock."""
    import docx as python_docx

    from gostforge.model import ListBlock
    from gostforge.parser import parse_docx

    doc = python_docx.Document()
    doc.add_paragraph("Один", style="List Number")
    doc.add_paragraph("Два", style="List Number")
    doc.add_paragraph("Обычный текст")
    doc.add_paragraph("Три", style="List Number")
    doc.add_paragraph("Четыре", style="List Number")
    out = tmp_path / "split.docx"
    doc.save(str(out))

    parsed = parse_docx(out)
    blocks = parsed.page_sections[0].content
    list_blocks = [b for b in blocks if isinstance(b, ListBlock)]
    assert len(list_blocks) == 2
    assert len(list_blocks[0].items) == 2
    assert len(list_blocks[1].items) == 2


def test_parser_extracts_image_rid_from_drawing(tmp_path: Path) -> None:
    """Парсер сохраняет rId изображения как 'embedded:rIdN' в Figure.image_path."""
    try:
        from PIL import Image  # type: ignore[import-not-found]
    except ImportError:
        import pytest
        pytest.skip("Pillow не установлен")
    import docx as python_docx

    from gostforge.model import Figure
    from gostforge.parser import parse_docx

    img = tmp_path / "test.png"
    Image.new("RGB", (20, 20), color="blue").save(img)

    doc = python_docx.Document()
    p = doc.add_paragraph()
    p.add_run().add_picture(str(img))
    out = tmp_path / "with_image.docx"
    doc.save(str(out))

    parsed = parse_docx(out)
    figures = [b for b in parsed.page_sections[0].content if isinstance(b, Figure)]
    assert len(figures) == 1
    # image_path должен начинаться с 'embedded:rId' — это идентификатор отношения
    assert figures[0].image_path.startswith("embedded:rId"), figures[0].image_path
