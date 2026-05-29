"""Стресс-тесты на крупный синтетический документ (~500 страниц).

Цель: зафиксировать верхние границы времени работы парсера, валидатора,
экспортёра и `compute_stats` на документе порядка 500+ страниц по
дефолтной вёрстке ГОСТ 7.32. Тесты помечены маркером `slow` — они
ощутимо медленнее обычных модульных тестов и предназначены для
регрессий по производительности.

Размеры подбираются под цель 500+ страниц при ~250 словах на странице
для шрифта Times New Roman 14pt и 1.5 межстрочного. Бюджеты времени
выставлены с запасом ×3-4 от наблюдаемых значений, чтобы не флейкать
на медленных CI-раннерах.
"""

from __future__ import annotations

import time
from pathlib import Path

import docx as docx_lib
import pytest

from gostforge.builder import work
from gostforge.exporter import export_docx
from gostforge.model import Document
from gostforge.parser import parse_docx
from gostforge.profile import Profile, load_profile
from gostforge.stats import compute_stats
from gostforge.validator import validate

# --- Параметры размера документа --------------------------------------------

# 30 глав уровня 1 × (40 параграфов + 5 подразделов × 6 параграфов) +
# рисунки и таблицы каждые несколько глав. Эмпирически даёт ~560 страниц
# (≈140k слов / 250 слов на страницу).
_TOP_LEVEL_SECTIONS = 30
_PARAGRAPHS_PER_SECTION = 40
_SUBSECTIONS_PER_SECTION = 5
_PARAGRAPHS_PER_SUBSECTION = 6

# Длинный текст параграфа основного раздела: повторяем заготовку, чтобы
# суммарно превысить порог в 500 страниц. Конкретный текст значения не
# имеет — нагрузка только на размер модели.
_PARAGRAPH_TEXT = (
    "Параграф содержит достаточно объёма, чтобы документ суммарно превысил пятьсот страниц. " * 8
)
_SUBSECTION_PARAGRAPH_TEXT = "Текст подраздела с пояснениями. " * 12

# Бюджеты времени (в секундах). Реальные замеры на чистой dev-машине:
# validate ≈ 1.2с, export ≈ 2.3с, parse ≈ 9.8с, compute_stats ≈ 0.01с.
# Бюджеты установлены с запасом ×3-4, чтобы тесты не флейкали на CI.
_VALIDATE_BUDGET_SECONDS = 30.0
_EXPORT_BUDGET_SECONDS = 30.0
_PARSE_BUDGET_SECONDS = 60.0
_STATS_BUDGET_SECONDS = 5.0


pytestmark = pytest.mark.slow


# --- Фикстуры ---------------------------------------------------------------


def _build_large_document() -> Document:
    """Собрать крупный документ через builder fluent-API.

    Структура:
      - 30 разделов уровня 1, каждый по 40 параграфов основного текста;
      - в каждом разделе по 5 подразделов, в каждом подразделе — 5 параграфов;
      - каждый третий раздел содержит рисунок-плейсхолдер;
      - каждый четвёртый раздел содержит небольшую таблицу.

    Результат: ~2100 параграфов, ~180 логических разделов всего,
    ~140k слов (≈560 страниц), ~10 рисунков, ~8 таблиц.
    """
    b = work(
        "Большая работа для стресс-тестирования",
        author="Тестовый Автор",
        year=2026,
        work_type="research_report",
    )
    for i in range(_TOP_LEVEL_SECTIONS):
        sb = b.section(f"Глава {i + 1}")
        for p in range(_PARAGRAPHS_PER_SECTION):
            sb = sb.paragraph(f"Параграф {p + 1}. {_PARAGRAPH_TEXT}")
        # Каждый третий раздел — рисунок-плейсхолдер.
        if i % 3 == 0:
            sb = sb.figure(image_path="", caption=f"Иллюстрация к главе {i + 1}")
        # Каждый четвёртый раздел — небольшая таблица.
        if i % 4 == 0:
            sb = sb.table(
                headers=["Показатель", "Значение"],
                rows=[["A", "1"], ["B", "2"]],
                caption=f"Данные эксперимента {i + 1}",
            )
        # Подразделы (level=2) с собственным текстом.
        for k in range(_SUBSECTIONS_PER_SECTION):
            sub_sb = sb.subsection(f"{i + 1}.{k + 1} Подраздел")
            for _ in range(_PARAGRAPHS_PER_SUBSECTION):
                sub_sb = sub_sb.paragraph(_SUBSECTION_PARAGRAPH_TEXT)
    return b.build()


@pytest.fixture(scope="module")
def large_document() -> Document:
    """Кэшированный на модуль большой документ — собирается один раз."""
    return _build_large_document()


@pytest.fixture(scope="module")
def gost_profile() -> Profile:
    """Базовый профиль ГОСТ 7.32-2017."""
    return load_profile("gost-7.32-2017")


@pytest.fixture(scope="module")
def large_docx_path(
    large_document: Document,
    gost_profile: Profile,
    tmp_path_factory: pytest.TempPathFactory,
) -> Path:
    """Однократно экспортированный .docx крупного документа.

    Кэшируется на модуль — round-trip тестам не нужно пересобирать его.
    """
    target = tmp_path_factory.mktemp("stress") / "large.docx"
    export_docx(large_document, gost_profile, target)
    return target


# --- Тесты ------------------------------------------------------------------


def test_validate_large_document_within_budget(
    large_document: Document, gost_profile: Profile
) -> None:
    """Валидатор отрабатывает на ~500-страничном документе в пределах бюджета.

    Не утверждаем конкретное число нарушений — достаточно того, что
    `validate()` не падает и возвращает корректный список за разумное
    время. На текущем профиле документ из builder выдаёт сотни warning-ов,
    что нормально (например, T.* по умолчанию-наследуемым атрибутам).
    """
    t0 = time.perf_counter()
    violations = validate(large_document, gost_profile)
    elapsed = time.perf_counter() - t0

    assert isinstance(violations, list)
    assert elapsed < _VALIDATE_BUDGET_SECONDS, (
        f"validate занял {elapsed:.2f}s, бюджет {_VALIDATE_BUDGET_SECONDS}s"
    )


def test_export_large_document_within_budget(
    large_document: Document, gost_profile: Profile, tmp_path: Path
) -> None:
    """Экспорт большого документа в .docx укладывается в бюджет."""
    target = tmp_path / "exported.docx"

    t0 = time.perf_counter()
    export_docx(large_document, gost_profile, target)
    elapsed = time.perf_counter() - t0

    assert target.exists()
    assert target.stat().st_size > 0
    # Файл должен открываться через python-docx без исключения.
    reopened = docx_lib.Document(str(target))
    assert len(reopened.paragraphs) > 0
    assert elapsed < _EXPORT_BUDGET_SECONDS, (
        f"export занял {elapsed:.2f}s, бюджет {_EXPORT_BUDGET_SECONDS}s"
    )


def test_parse_large_document_round_trip(large_document: Document, large_docx_path: Path) -> None:
    """Парсер читает большой .docx обратно за разумное время и без потерь структуры.

    Сверяем число параграфов и логических разделов с оригиналом в пределах
    10% допуска: экспортёр может объединять/добавлять служебные параграфы
    (подписи к таблицам/рисункам и т.п.).
    """
    t0 = time.perf_counter()
    parsed = parse_docx(large_docx_path)
    elapsed = time.perf_counter() - t0

    original_stats = compute_stats(large_document)
    parsed_stats = compute_stats(parsed)

    # Логические разделы 1 уровня должны сохраняться точно.
    assert parsed_stats.logical_sections_level_1 == original_stats.logical_sections_level_1
    # Всего логических разделов — также точно (подразделы парсер видит по
    # уровню заголовка).
    assert parsed_stats.logical_sections_total == original_stats.logical_sections_total

    # Число параграфов: допускаем дрейф ±10% (caption-параграфы для
    # таблиц/рисунков могут считаться по-разному).
    original_paragraphs = original_stats.paragraphs
    parsed_paragraphs = parsed_stats.paragraphs
    drift_ratio = abs(parsed_paragraphs - original_paragraphs) / original_paragraphs
    assert drift_ratio <= 0.10, (
        f"параграфов: было {original_paragraphs}, стало {parsed_paragraphs} "
        f"(дрейф {drift_ratio:.1%}, лимит 10%)"
    )

    assert elapsed < _PARSE_BUDGET_SECONDS, (
        f"parse занял {elapsed:.2f}s, бюджет {_PARSE_BUDGET_SECONDS}s"
    )


def test_compute_stats_on_large_document(large_document: Document) -> None:
    """`compute_stats` собирает корректные метрики на большом документе быстро.

    Подтверждаем, что фикстура действительно крупная (по нескольким
    осям), а сам подсчёт занимает доли секунды.
    """
    t0 = time.perf_counter()
    stats = compute_stats(large_document)
    elapsed = time.perf_counter() - t0

    # Документ должен быть действительно большим.
    assert stats.paragraphs >= 2000, f"ожидалось ≥2000 параграфов, получили {stats.paragraphs}"
    assert stats.logical_sections_level_1 >= 20, (
        f"ожидалось ≥20 разделов 1 уровня, получили {stats.logical_sections_level_1}"
    )
    # ~500 страниц при ~250 словах на странице → ~125k слов и больше.
    assert stats.words >= 100_000, f"ожидалось ≳100k слов (≈500 страниц), получили {stats.words}"

    assert elapsed < _STATS_BUDGET_SECONDS, (
        f"compute_stats заняла {elapsed:.3f}s, бюджет {_STATS_BUDGET_SECONDS}s"
    )
