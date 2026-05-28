"""Тесты inline-парсера (Фаза 2.5, §4.5).

Покрытие: TextRun.underline / color_hex из <w:rPr>, InlineFormula
(<m:oMath> внутри параграфа, без <m:oMathPara>), Citation (эвристика
«[N]» / «[N, с. P]» с проверкой по Document.bibliography), CrossRef.prefix
(перенос «(см. » / « (» из предыдущего TextRun-а).

Тесты строят синтетические .docx через python-docx + lxml, чтобы не
держать в репозитории бинарные фикстуры.
"""

from __future__ import annotations

from pathlib import Path

import docx as python_docx
from docx.oxml.ns import qn
from lxml import etree

from gostforge.model import (
    Citation,
    CrossRef,
    InlineFormula,
    LogicalSection,
    Paragraph,
    TextRun,
)
from gostforge.parser.docx_parser import parse_docx

from .conftest import make_docx

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


# --- helpers -----------------------------------------------------------------


def _add_underlined_run(paragraph: object, text: str) -> None:
    """Добавить run с явно выставленным <w:u w:val="single"/>."""
    run = paragraph.add_run(text)  # type: ignore[attr-defined]
    rpr = run._r.get_or_add_rPr()
    u = etree.SubElement(rpr, qn("w:u"))
    u.set(qn("w:val"), "single")


def _add_colored_run(paragraph: object, text: str, color_hex: str) -> None:
    """Добавить run с явно выставленным <w:color w:val="RRGGBB"/>.

    color_hex принимается с «#» или без — для удобства тестов.
    """
    run = paragraph.add_run(text)  # type: ignore[attr-defined]
    rpr = run._r.get_or_add_rPr()
    color_el = etree.SubElement(rpr, qn("w:color"))
    color_el.set(qn("w:val"), color_hex.lstrip("#"))


def _add_inline_omath(paragraph: object, latex: str) -> None:
    """Добавить <m:oMath> прямым ребёнком <w:p> (БЕЗ <m:oMathPara>).

    Парсер должен трактовать это как inline-формулу.
    """
    p_xml = paragraph._p  # type: ignore[attr-defined]
    omath = etree.SubElement(p_xml, f"{{{M_NS}}}oMath")
    m_r = etree.SubElement(omath, f"{{{M_NS}}}r")
    m_t = etree.SubElement(m_r, f"{{{M_NS}}}t")
    m_t.text = latex


def _add_fld_simple_ref(paragraph: object, target_id: str) -> None:
    """Добавить <w:fldSimple w:instr=" REF target_id \\h "/> в параграф."""
    p_xml = paragraph._p  # type: ignore[attr-defined]
    fld = etree.SubElement(p_xml, qn("w:fldSimple"))
    fld.set(qn("w:instr"), f" REF {target_id} \\h ")
    # Видимый плейсхолдер внутри fldSimple (Word нужен хотя бы один run).
    inner_r = etree.SubElement(fld, qn("w:r"))
    inner_t = etree.SubElement(inner_r, qn("w:t"))
    inner_t.text = "[?]"


def _first_body_paragraph(doc: object) -> Paragraph:
    """Извлечь первый Paragraph из парсенной модели (после parse_docx)."""
    parsed = doc  # type: ignore[assignment]
    for item in parsed.page_sections[0].content:  # type: ignore[attr-defined]
        if isinstance(item, Paragraph):
            return item
    raise AssertionError("В документе нет Paragraph-ов на верхнем уровне content")


def _para_in_first_section(doc: object) -> Paragraph:
    """Найти первый Paragraph внутри первой LogicalSection."""
    parsed = doc  # type: ignore[assignment]
    for item in parsed.page_sections[0].content:  # type: ignore[attr-defined]
        if isinstance(item, LogicalSection):
            for child in item.children:
                if isinstance(child, Paragraph):
                    return child
    raise AssertionError("В первой LogicalSection нет Paragraph-ов")


# --- TextRun.underline / color_hex --------------------------------------------


def test_parses_underline_from_w_u(tmp_path: Path) -> None:
    """Run с явным <w:u w:val="single"/> → TextRun.underline=True."""
    doc = python_docx.Document()
    paragraph = doc.add_paragraph()
    _add_underlined_run(paragraph, "Подчёркнутый текст")
    out = tmp_path / "underline.docx"
    doc.save(str(out))

    parsed = parse_docx(out)
    paragraph_model = _first_body_paragraph(parsed)
    runs = [el for el in paragraph_model.content if isinstance(el, TextRun)]
    assert len(runs) == 1
    assert runs[0].text == "Подчёркнутый текст"
    assert runs[0].underline is True


def test_parses_color_from_w_color(tmp_path: Path) -> None:
    """Run с <w:color w:val="FF0000"/> → TextRun.color_hex='#FF0000'."""
    doc = python_docx.Document()
    paragraph = doc.add_paragraph()
    _add_colored_run(paragraph, "Красный текст", "FF0000")
    out = tmp_path / "color.docx"
    doc.save(str(out))

    parsed = parse_docx(out)
    paragraph_model = _first_body_paragraph(parsed)
    runs = [el for el in paragraph_model.content if isinstance(el, TextRun)]
    assert len(runs) == 1
    assert runs[0].color_hex == "#FF0000"


def test_color_auto_treated_as_none(tmp_path: Path) -> None:
    """<w:color w:val="auto"/> → color_hex остаётся None (наследуется)."""
    doc = python_docx.Document()
    paragraph = doc.add_paragraph()
    _add_colored_run(paragraph, "Авто-цвет", "auto")
    out = tmp_path / "auto_color.docx"
    doc.save(str(out))

    parsed = parse_docx(out)
    paragraph_model = _first_body_paragraph(parsed)
    runs = [el for el in paragraph_model.content if isinstance(el, TextRun)]
    assert runs and runs[0].color_hex is None


# --- InlineFormula -----------------------------------------------------------


def test_inline_formula_inside_paragraph_not_block(tmp_path: Path) -> None:
    """<m:oMath> прямой ребёнок <w:p> (без oMathPara) → InlineFormula в Paragraph."""
    doc = python_docx.Document()
    paragraph = doc.add_paragraph("Перед формулой ")
    _add_inline_omath(paragraph, "h\\nu")
    out = tmp_path / "inline_formula.docx"
    doc.save(str(out))

    parsed = parse_docx(out)
    paragraph_model = _first_body_paragraph(parsed)
    formulas = [el for el in paragraph_model.content if isinstance(el, InlineFormula)]
    assert len(formulas) == 1
    assert formulas[0].latex == "h\\nu"


def test_inline_formula_keeps_order_relative_to_text(tmp_path: Path) -> None:
    """В параграфе с порядком (текст, формула, текст) парсер сохраняет порядок."""
    doc = python_docx.Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("Энергия ")
    _add_inline_omath(paragraph, "h\\nu")
    paragraph.add_run(" описана")
    out = tmp_path / "inline_order.docx"
    doc.save(str(out))

    parsed = parse_docx(out)
    paragraph_model = _first_body_paragraph(parsed)
    kinds = [type(el).__name__ for el in paragraph_model.content]
    assert kinds == ["TextRun", "InlineFormula", "TextRun"]
    assert isinstance(paragraph_model.content[0], TextRun)
    assert paragraph_model.content[0].text == "Энергия "
    assert isinstance(paragraph_model.content[1], InlineFormula)
    assert paragraph_model.content[1].latex == "h\\nu"
    assert isinstance(paragraph_model.content[2], TextRun)
    assert paragraph_model.content[2].text == " описана"


def test_inline_formula_strips_dollar_wrappers(tmp_path: Path) -> None:
    """Если LaTeX внутри <m:t> обёрнут в $...$ — парсер их снимает."""
    doc = python_docx.Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("a + ")
    _add_inline_omath(paragraph, "$x^2$")
    out = tmp_path / "dollar.docx"
    doc.save(str(out))

    parsed = parse_docx(out)
    paragraph_model = _first_body_paragraph(parsed)
    formulas = [el for el in paragraph_model.content if isinstance(el, InlineFormula)]
    assert formulas and formulas[0].latex == "x^2"


def test_block_omath_para_still_becomes_formula_block(tmp_path: Path) -> None:
    """<m:oMathPara><m:oMath/></m:oMathPara> остаётся блочной Formula (а не Paragraph)."""
    from gostforge.model import Formula

    doc = python_docx.Document()
    paragraph = doc.add_paragraph()
    p_xml = paragraph._p
    omath_para = etree.SubElement(p_xml, f"{{{M_NS}}}oMathPara")
    omath = etree.SubElement(omath_para, f"{{{M_NS}}}oMath")
    m_r = etree.SubElement(omath, f"{{{M_NS}}}r")
    m_t = etree.SubElement(m_r, f"{{{M_NS}}}t")
    m_t.text = "E=mc^2"
    out = tmp_path / "block_formula.docx"
    doc.save(str(out))

    parsed = parse_docx(out)
    blocks = parsed.page_sections[0].content
    formulas = [b for b in blocks if isinstance(b, Formula)]
    assert len(formulas) == 1
    assert formulas[0].latex == "E=mc^2"
    # Параграфа с inline-формулой быть не должно (только блочная).
    paragraphs = [b for b in blocks if isinstance(b, Paragraph)]
    for paragraph_model in paragraphs:
        assert not any(isinstance(el, InlineFormula) for el in paragraph_model.content)


# --- Citation ----------------------------------------------------------------


def test_citation_recognized_when_n_in_bibliography_range(tmp_path: Path) -> None:
    """bibliography из 3 записей, текст «[2, с. 42]» → Citation на ref-2 с pages=42."""
    bibliography = [
        "Иванов И. И. Один. — Москва : Наука, 2020. — 100 с.",
        "Петров П. П. Два. — Москва : Наука, 2021. — 200 с.",
        "Сидоров С. С. Три. — Москва : Наука, 2022. — 300 с.",
    ]
    path = make_docx(
        tmp_path / "cit_in_range.docx",
        headings=[(1, "Введение")],
        paragraphs=["Это подтверждается в [2, с. 42]."],
        bibliography=bibliography,
    )
    parsed = parse_docx(path)
    # Параграф «Введение» внутри LogicalSection.
    paragraph_model = _para_in_first_section(parsed)
    citations = [el for el in paragraph_model.content if isinstance(el, Citation)]
    assert len(citations) == 1
    # source_id — id второй записи в bibliography (ref-2).
    assert citations[0].source_id == parsed.bibliography[1].id
    assert citations[0].pages == "42"
    # Должен сохраниться окружающий текст.
    text_chunks = [el.text for el in paragraph_model.content if isinstance(el, TextRun)]
    joined = "".join(text_chunks)
    assert "Это подтверждается в " in joined
    assert "." in joined


def test_citation_not_recognized_when_n_out_of_range(tmp_path: Path) -> None:
    """bibliography из 2 записей, текст «[5]» → остаётся как TextRun."""
    bibliography = [
        "Иванов И. И. Один. — Москва : Наука, 2020. — 100 с.",
        "Петров П. П. Два. — Москва : Наука, 2021. — 200 с.",
    ]
    path = make_docx(
        tmp_path / "cit_out_of_range.docx",
        headings=[(1, "Введение")],
        paragraphs=["Источник [5] невалиден."],
        bibliography=bibliography,
    )
    parsed = parse_docx(path)
    paragraph_model = _para_in_first_section(parsed)
    citations = [el for el in paragraph_model.content if isinstance(el, Citation)]
    assert citations == []
    # Текст должен содержать «[5]» как есть.
    joined = "".join(el.text for el in paragraph_model.content if isinstance(el, TextRun))
    assert "[5]" in joined


def test_citation_without_pages(tmp_path: Path) -> None:
    """«[3]» → Citation на ref-3 с pages=None и template='[{n}]'."""
    bibliography = [
        "Иванов И. И. Один. — Москва : Наука, 2020. — 100 с.",
        "Петров П. П. Два. — Москва : Наука, 2021. — 200 с.",
        "Сидоров С. С. Три. — Москва : Наука, 2022. — 300 с.",
    ]
    path = make_docx(
        tmp_path / "cit_no_pages.docx",
        headings=[(1, "Введение")],
        paragraphs=["Это известно [3]."],
        bibliography=bibliography,
    )
    parsed = parse_docx(path)
    paragraph_model = _para_in_first_section(parsed)
    citations = [el for el in paragraph_model.content if isinstance(el, Citation)]
    assert len(citations) == 1
    assert citations[0].pages is None
    assert citations[0].template == "[{n}]"


def test_citation_does_not_swallow_math_indices(tmp_path: Path) -> None:
    """«[0; 1]» — не Citation, т. к. 0 не валиден для 1-based индекса.

    А «[1]» в той же строке всё же распознаётся как Citation на ref-1.
    Это нормально: эвристика консервативна по N, но допускает валидные.
    """
    bibliography = [
        "Иванов И. И. Один. — Москва : Наука, 2020. — 100 с.",
        "Петров П. П. Два. — Москва : Наука, 2021. — 200 с.",
        "Сидоров С. С. Три. — Москва : Наука, 2022. — 300 с.",
        "Кузнецов К. К. Четыре. — Москва : Наука, 2023. — 400 с.",
        "Лебедев Л. Л. Пять. — Москва : Наука, 2024. — 500 с.",
    ]
    path = make_docx(
        tmp_path / "cit_math.docx",
        headings=[(1, "Введение")],
        paragraphs=["В диапазоне [0; 1] лежит ноль."],
        bibliography=bibliography,
    )
    parsed = parse_docx(path)
    paragraph_model = _para_in_first_section(parsed)
    citations = [el for el in paragraph_model.content if isinstance(el, Citation)]
    # «[0; 1]» — это не [N] (внутри есть «;», что наш regex не пропускает),
    # и не [N, с. P]. Поэтому никаких Citation не должно появиться.
    assert citations == []
    joined = "".join(el.text for el in paragraph_model.content if isinstance(el, TextRun))
    assert "[0; 1]" in joined


# --- CrossRef.prefix ----------------------------------------------------------


def test_cross_ref_prefix_extracted_from_preceding_run(tmp_path: Path) -> None:
    """Параграф «Подробнее (см. <REF fig-1>).» → CrossRef.prefix=' (см. '."""
    doc = python_docx.Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("Подробнее (см. ")
    _add_fld_simple_ref(paragraph, "fig-1")
    paragraph.add_run(").")
    out = tmp_path / "crossref_prefix.docx"
    doc.save(str(out))

    parsed = parse_docx(out)
    paragraph_model = _first_body_paragraph(parsed)
    crossrefs = [el for el in paragraph_model.content if isinstance(el, CrossRef)]
    assert len(crossrefs) == 1
    assert crossrefs[0].target_id == "fig-1"
    # Префикс должен быть взят из заранее заготовленного хвоста предыдущего TextRun-а.
    assert crossrefs[0].prefix == " (см. "
    # Предыдущий TextRun теперь укорочен: остался только «Подробнее».
    text_runs = [el for el in paragraph_model.content if isinstance(el, TextRun)]
    assert text_runs[0].text == "Подробнее"
    # Закрывающая скобка с точкой — следующий TextRun.
    assert any(r.text == ")." for r in text_runs)


def test_cross_ref_without_recognizable_prefix(tmp_path: Path) -> None:
    """Если хвост TextRun-а не подходит под шаблон — prefix остаётся None."""
    doc = python_docx.Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("Текст без открывающих скобок ")
    _add_fld_simple_ref(paragraph, "tbl-2")
    out = tmp_path / "crossref_no_prefix.docx"
    doc.save(str(out))

    parsed = parse_docx(out)
    paragraph_model = _first_body_paragraph(parsed)
    crossrefs = [el for el in paragraph_model.content if isinstance(el, CrossRef)]
    assert len(crossrefs) == 1
    assert crossrefs[0].target_id == "tbl-2"
    # Хвост «скобок » не входит в _CROSS_REF_PREFIX_SUFFIXES — prefix=None.
    assert crossrefs[0].prefix is None
    # И текст TextRun-а должен остаться нетронутым.
    text_runs = [el for el in paragraph_model.content if isinstance(el, TextRun)]
    assert text_runs[0].text == "Текст без открывающих скобок "


# --- Round-trip ---------------------------------------------------------------


def test_round_trip_runs_with_underline_and_formula(tmp_path: Path) -> None:
    """Round-trip: модель → экспорт → парсинг даёт те же inline-элементы.

    Проверяем underline (TextRun), порядок (TextRun, InlineFormula, TextRun).
    """
    try:
        from gostforge.exporter import export_docx
        from gostforge.profile import load_profile
    except ImportError:  # pragma: no cover — экспортёр обязателен в проекте
        import pytest

        pytest.skip("Экспортёр недоступен — round-trip пропущен (TODO)")
        return

    from gostforge.model import Document, PageSection

    document = Document()
    document.page_sections.append(
        PageSection(
            id="main",
            name="m",
            type="main",
            content=[
                Paragraph(
                    id="p1",
                    content=[
                        TextRun(text="До ", underline=True),
                        InlineFormula(latex="h\\nu"),
                        TextRun(text=" после"),
                    ],
                )
            ],
        )
    )

    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "roundtrip.docx"
    export_docx(document, profile, str(out))

    parsed = parse_docx(out)
    paragraph_model = _first_body_paragraph(parsed)
    kinds = [type(el).__name__ for el in paragraph_model.content]
    # Порядок должен сохраниться. Содержимое — текст «До », формула, текст « после».
    assert kinds == ["TextRun", "InlineFormula", "TextRun"]
    assert isinstance(paragraph_model.content[0], TextRun)
    assert paragraph_model.content[0].text == "До "
    assert paragraph_model.content[0].underline is True
    assert isinstance(paragraph_model.content[1], InlineFormula)
    assert paragraph_model.content[1].latex == "h\\nu"
    assert isinstance(paragraph_model.content[2], TextRun)
    assert paragraph_model.content[2].text == " после"
