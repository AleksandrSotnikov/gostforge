"""Тесты CLI-команды `gostforge annotate`."""

from __future__ import annotations

from pathlib import Path

import pytest
from click.testing import CliRunner
from docx import Document as DocxDocument

from gostforge.cli import main
from gostforge.model import Document, PageGeometry, PageSection


def _make_minimal_docx(path: Path) -> None:
    doc = DocxDocument()
    doc.add_paragraph("Тестовый абзац.")
    doc.save(str(path))


@pytest.fixture()
def bad_docx(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> Path:
    p = tmp_path / "bad.docx"
    _make_minimal_docx(p)

    def fake_parse(_: object) -> Document:
        doc = Document()
        doc.page_sections.append(
            PageSection(
                id="main",
                name="Основная часть",
                type="main",
                page=PageGeometry(
                    margins_mm={"top": 25, "right": 15, "bottom": 20, "left": 30}
                ),
            )
        )
        return doc

    monkeypatch.setattr(
        "gostforge.annotator.docx_annotator.parse_docx", fake_parse
    )
    return p


def test_cli_annotate_creates_output(bad_docx: Path, tmp_path: Path) -> None:
    out = tmp_path / "out.docx"
    runner = CliRunner()
    result = runner.invoke(
        main,
        ["annotate", str(bad_docx), "-o", str(out)],
    )
    assert result.exit_code == 0, result.output
    assert out.exists()


def test_cli_annotate_exit_code_zero(bad_docx: Path, tmp_path: Path) -> None:
    out = tmp_path / "out.docx"
    runner = CliRunner()
    result = runner.invoke(
        main,
        ["annotate", str(bad_docx), "-o", str(out)],
    )
    assert result.exit_code == 0


def test_cli_annotate_prints_count_message(bad_docx: Path, tmp_path: Path) -> None:
    out = tmp_path / "out.docx"
    runner = CliRunner()
    result = runner.invoke(
        main,
        ["annotate", str(bad_docx), "-o", str(out)],
    )
    assert "Создано пометок" in result.output


def test_cli_annotate_default_style_is_comments(
    bad_docx: Path, tmp_path: Path
) -> None:
    """Без --style по умолчанию используется comments-режим: в архиве
    появляется word/comments.xml."""
    import zipfile

    out = tmp_path / "out.docx"
    runner = CliRunner()
    result = runner.invoke(main, ["annotate", str(bad_docx), "-o", str(out)])
    assert result.exit_code == 0, result.output
    with zipfile.ZipFile(out, "r") as z:
        assert "word/comments.xml" in z.namelist()


def test_cli_annotate_style_inline_no_comments_part(
    bad_docx: Path, tmp_path: Path
) -> None:
    """--style inline → старое поведение, без word/comments.xml."""
    import zipfile

    out = tmp_path / "out.docx"
    runner = CliRunner()
    result = runner.invoke(
        main, ["annotate", str(bad_docx), "-o", str(out), "--style", "inline"]
    )
    assert result.exit_code == 0, result.output
    with zipfile.ZipFile(out, "r") as z:
        assert "word/comments.xml" not in z.namelist()


def test_cli_annotate_style_comments_creates_comments_part(
    bad_docx: Path, tmp_path: Path
) -> None:
    """Явный --style comments тоже создаёт comments-part."""
    import zipfile

    out = tmp_path / "out.docx"
    runner = CliRunner()
    result = runner.invoke(
        main, ["annotate", str(bad_docx), "-o", str(out), "--style", "comments"]
    )
    assert result.exit_code == 0, result.output
    with zipfile.ZipFile(out, "r") as z:
        assert "word/comments.xml" in z.namelist()


def test_cli_annotate_rejects_unknown_style(bad_docx: Path, tmp_path: Path) -> None:
    """Неизвестное значение --style → click отвечает кодом != 0."""
    out = tmp_path / "out.docx"
    runner = CliRunner()
    result = runner.invoke(
        main, ["annotate", str(bad_docx), "-o", str(out), "--style", "bogus"]
    )
    assert result.exit_code != 0
