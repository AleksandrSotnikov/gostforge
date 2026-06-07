"""Тесты основной надписи (штампа ЕСКД, ГОСТ 2.104) — рендер в footer."""

from __future__ import annotations

from pathlib import Path

import docx as python_docx

from gostforge.exporter import export_docx
from gostforge.fixer import fix, registered_fixers
from gostforge.model import (
    Document,
    PageGeometry,
    PageSection,
    Paragraph,
    TextRun,
    TitleBlock,
    TitleBlockRole,
)
from gostforge.parser import parse_docx
from gostforge.profile import load_profile
from gostforge.validator import validate


def _profile_requiring_title_block():  # type: ignore[no-untyped-def]
    """Профиль ЕСКД с включённым штампом (для F.08)."""
    from gostforge.profile.schema import TitleBlockProfile

    profile = load_profile("gost-7.32-2017")
    profile.styles.page.title_block = TitleBlockProfile(enabled=True, organization="Кафедра X")
    return profile


def _doc_with_title_block(tb: TitleBlock | None) -> Document:
    doc = Document()
    doc.metadata.title = "Пояснительная записка"
    doc.page_sections.append(
        PageSection(
            id="main",
            name="Основная часть",
            type="main",
            page=PageGeometry(),
            content=[Paragraph(id="p1", content=[TextRun(text="Текст")], style_name="Normal")],
            title_block=tb,
        )
    )
    return doc


def _footer_tables(out: Path):  # type: ignore[no-untyped-def]
    docx_doc = python_docx.Document(str(out))
    return docx_doc.sections[0].footer.tables


def _all_cell_text(table) -> str:  # type: ignore[no-untyped-def]
    return "\n".join(c.text for row in table.rows for c in row.cells)


def test_no_title_block_no_footer_table(tmp_path: Path) -> None:
    """Без штампа в footer-е нет таблиц."""
    doc = _doc_with_title_block(None)
    out = tmp_path / "no-tb.docx"
    export_docx(doc, load_profile("gost-7.32-2017"), out)
    assert _footer_tables(out) == []


def test_disabled_title_block_no_table(tmp_path: Path) -> None:
    """enabled=False → штамп не пишется."""
    doc = _doc_with_title_block(TitleBlock(enabled=False))
    out = tmp_path / "disabled-tb.docx"
    export_docx(doc, load_profile("gost-7.32-2017"), out)
    assert _footer_tables(out) == []


def test_form1_renders_with_graphs(tmp_path: Path) -> None:
    """Форма 1: таблица 8 колонок с обозначением/наименованием/организацией/ролями."""
    tb = TitleBlock(
        enabled=True,
        form="form1",
        designation="АБВГ.123456.001 ПЗ",
        title="Анализ алгоритмов",
        organization="Кафедра ИВТ",
        stage="У",
        sheet="1",
        sheets_total="42",
        roles=[
            TitleBlockRole(role="Разраб.", name="Иванов", date="01.06.26"),
            TitleBlockRole(role="Пров.", name="Петров"),
        ],
    )
    doc = _doc_with_title_block(tb)
    out = tmp_path / "form1.docx"
    export_docx(doc, load_profile("gost-7.32-2017"), out)

    tables = _footer_tables(out)
    assert len(tables) == 1
    table = tables[0]
    assert len(table.columns) == 8
    text = _all_cell_text(table)
    assert "АБВГ.123456.001 ПЗ" in text
    assert "Анализ алгоритмов" in text
    assert "Кафедра ИВТ" in text
    assert "Разраб." in text and "Иванов" in text
    # «Листов» — отдельная графа-метка, значение — в соседней ячейке.
    assert "Листов" in text and "42" in text


def test_title_defaults_to_document_title(tmp_path: Path) -> None:
    """Если наименование штампа пустое — берётся заголовок документа."""
    tb = TitleBlock(enabled=True, form="form1", designation="X.001")
    doc = _doc_with_title_block(tb)
    out = tmp_path / "tb-title-default.docx"
    export_docx(doc, load_profile("gost-7.32-2017"), out)
    text = _all_cell_text(_footer_tables(out)[0])
    assert "Пояснительная записка" in text


def test_form2a_renders_compact(tmp_path: Path) -> None:
    """Форма 2а: узкая таблица из 2 колонок (обозначение + лист)."""
    tb = TitleBlock(enabled=True, form="form2a", designation="X.001", sheet="3")
    doc = _doc_with_title_block(tb)
    out = tmp_path / "form2a.docx"
    export_docx(doc, load_profile("gost-7.32-2017"), out)
    tables = _footer_tables(out)
    assert len(tables) == 1
    assert len(tables[0].columns) == 2
    text = _all_cell_text(tables[0])
    assert "X.001" in text
    assert "Лист 3" in text


def test_profile_editor_title_block_roundtrip() -> None:
    """Штамп переживает save-путь редактора профиля (data dict → YAML → Profile)."""
    from gostforge.profile.schema import Profile
    from gostforge.web.profile_editor import build_profile_yaml, profile_to_data

    data = profile_to_data(load_profile("gost-7.32-2017"))
    data["styles"]["page"]["title_block"] = {
        "enabled": True,
        "form": "form1",
        "organization": "Кафедра ИВТ",
        "roles": [{"role": "Разраб.", "name": "", "date": ""}],
    }
    yaml_text = build_profile_yaml(data)  # валидирует через Pydantic
    assert "title_block" in yaml_text

    import yaml as _yaml

    restored = Profile(**_yaml.safe_load(yaml_text))
    assert restored.styles.page.title_block is not None
    assert restored.styles.page.title_block.enabled is True
    assert restored.styles.page.title_block.roles[0].role == "Разраб."


def test_title_block_applied_from_profile(tmp_path: Path) -> None:
    """Штамп из профиля проставляется при title_block=None в модели."""
    profile = load_profile("gost-7.32-2017")
    assert profile.styles.page.title_block is None or not profile.styles.page.title_block.enabled
    # Программно включаем штамп в профиле.
    from gostforge.profile.schema import TitleBlockProfile

    profile.styles.page.title_block = TitleBlockProfile(enabled=True, organization="Кафедра X")
    doc = _doc_with_title_block(None)
    out = tmp_path / "profile-tb.docx"
    export_docx(doc, profile, out)
    tables = _footer_tables(out)
    assert len(tables) == 1
    assert "Кафедра X" in _all_cell_text(tables[0])


# --- парсер: детекция наличия штампа -----------------------------------------


def test_parser_detects_title_block_presence(tmp_path: Path) -> None:
    """Экспортированный штамп определяется парсером как присутствующий."""
    tb = TitleBlock(enabled=True, form="form1", designation="X.001", organization="Каф.")
    doc = _doc_with_title_block(tb)
    out = tmp_path / "rt.docx"
    export_docx(doc, load_profile("gost-7.32-2017"), out)
    reparsed = parse_docx(out)
    parsed_tb = reparsed.page_sections[0].title_block
    assert parsed_tb is not None
    assert parsed_tb.enabled is True


def test_parser_extracts_form1_fields(tmp_path: Path) -> None:
    """Поля формы 1 переживают round-trip export → parse (best-effort)."""
    tb = TitleBlock(
        enabled=True,
        form="form1",
        designation="АБВГ.123456.001 ПЗ",
        title="Анализ алгоритмов",
        organization="Кафедра ИВТ",
        sheets_total="42",
        roles=[TitleBlockRole(role="Разраб.", name="Иванов", date="01.06.26")],
    )
    doc = _doc_with_title_block(tb)
    out = tmp_path / "rt-fields.docx"
    export_docx(doc, load_profile("gost-7.32-2017"), out)
    parsed = parse_docx(out).page_sections[0].title_block
    assert parsed is not None
    assert parsed.form == "form1"
    assert parsed.designation == "АБВГ.123456.001 ПЗ"
    assert parsed.title == "Анализ алгоритмов"
    assert parsed.organization == "Кафедра ИВТ"
    assert parsed.sheets_total == "42"
    assert any(r.role == "Разраб." and r.name == "Иванов" for r in parsed.roles)


def test_parser_extracts_form2a_fields(tmp_path: Path) -> None:
    """Форма 2а: обозначение и номер листа восстанавливаются."""
    tb = TitleBlock(enabled=True, form="form2a", designation="X.001", sheet="3")
    doc = _doc_with_title_block(tb)
    out = tmp_path / "rt-2a.docx"
    export_docx(doc, load_profile("gost-7.32-2017"), out)
    parsed = parse_docx(out).page_sections[0].title_block
    assert parsed is not None
    assert parsed.form == "form2a"
    assert parsed.designation == "X.001"
    assert parsed.sheet == "3"


def test_parser_no_title_block_when_absent(tmp_path: Path) -> None:
    """Без штампа парсер возвращает title_block=None."""
    doc = _doc_with_title_block(None)
    out = tmp_path / "rt-none.docx"
    export_docx(doc, load_profile("gost-7.32-2017"), out)
    reparsed = parse_docx(out)
    assert reparsed.page_sections[0].title_block is None


# --- F.08: нормоконтроль наличия штампа --------------------------------------


def test_f08_silent_when_not_required() -> None:
    """Профиль без штампа (gost-7.32) → F.08 молчит."""
    profile = load_profile("gost-7.32-2017")
    doc = _doc_with_title_block(None)
    assert [v for v in validate(doc, profile) if v.check_code == "F.08"] == []


def test_f08_errors_when_missing() -> None:
    """Профиль требует штамп, а его нет → ошибка F.08."""
    profile = _profile_requiring_title_block()
    doc = _doc_with_title_block(None)
    f08 = [v for v in validate(doc, profile) if v.check_code == "F.08"]
    assert len(f08) == 1
    assert f08[0].severity == "error"


def test_f08_fixer_adds_title_block() -> None:
    """F.08-фиксер добавляет штамп из профиля; после фикса проверка молчит."""
    assert "F.08" in registered_fixers()
    profile = _profile_requiring_title_block()
    doc = _doc_with_title_block(None)

    applied = fix(doc, profile, codes=["F.08"])
    assert len(applied) == 1
    assert applied[0].fixer_code == "F.08"
    tb = doc.page_sections[0].title_block
    assert tb is not None and tb.enabled is True
    assert tb.organization == "Кафедра X"
    assert [v for v in validate(doc, profile) if v.check_code == "F.08"] == []


def test_f08_fixer_noop_when_not_required() -> None:
    """Профиль штамп не требует — фиксер ничего не делает."""
    profile = load_profile("gost-7.32-2017")
    doc = _doc_with_title_block(None)
    assert fix(doc, profile, codes=["F.08"]) == []
    assert doc.page_sections[0].title_block is None
