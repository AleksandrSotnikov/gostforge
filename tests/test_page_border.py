"""Тесты рамки листа (PageBorder, OOXML <w:pgBorders>) — ЕСКД ГОСТ 2.104.

Покрывают сквозной round-trip: model → export → parse, а также
применение рамки из профиля при экспорте.
"""

from __future__ import annotations

from pathlib import Path

import docx as python_docx

from gostforge.exporter import export_docx
from gostforge.fixer import fix, registered_fixers
from gostforge.model import (
    Document,
    PageBorder,
    PageGeometry,
    PageSection,
    Paragraph,
    TextRun,
)
from gostforge.parser import parse_docx
from gostforge.profile import load_profile
from gostforge.validator import validate

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _profile_requiring_border():  # type: ignore[no-untyped-def]
    """Профиль ЕСКД с включённой рамкой (для проверки/фиксера F.07)."""
    profile = load_profile("gost-r-2.105-2019")
    assert profile.styles.page.border is not None
    profile.styles.page.border.enabled = True
    return profile


def _doc_with_border(border: PageBorder | None) -> Document:
    """Документ из одной секции с заданной (или отсутствующей) рамкой."""
    doc = Document()
    doc.page_sections.append(
        PageSection(
            id="main",
            name="Основная часть",
            type="main",
            page=PageGeometry(border=border),
            content=[Paragraph(id="p1", content=[TextRun(text="Текст")], style_name="Normal")],
        )
    )
    return doc


def test_border_none_writes_no_pgborders(tmp_path: Path) -> None:
    """Без рамки в sectPr не должно быть <w:pgBorders>; парсер вернёт None."""
    doc = _doc_with_border(None)
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "no-border.docx"
    export_docx(doc, profile, out)

    docx_doc = python_docx.Document(str(out))
    sect_pr = docx_doc.sections[0]._sectPr
    assert sect_pr.find(f"{{{W_NS}}}pgBorders") is None

    reparsed = parse_docx(out)
    assert reparsed.page_sections[0].page.border is None


def test_border_roundtrip_preserves_params(tmp_path: Path) -> None:
    """Рамка из модели переживает export → parse без потери параметров."""
    border = PageBorder(
        enabled=True,
        style="single",
        size_eighth_pt=8,  # 1.0 pt
        color="auto",
        offset_from="text",
        space_pt=0,
    )
    doc = _doc_with_border(border)
    profile = load_profile("gost-7.32-2017")
    out = tmp_path / "framed.docx"
    export_docx(doc, profile, out)

    # В XML действительно есть рамка на всех сторонах.
    docx_doc = python_docx.Document(str(out))
    pg_borders = docx_doc.sections[0]._sectPr.find(f"{{{W_NS}}}pgBorders")
    assert pg_borders is not None
    assert pg_borders.get(f"{{{W_NS}}}offsetFrom") == "text"
    for side in ("top", "left", "bottom", "right"):
        assert pg_borders.find(f"{{{W_NS}}}{side}") is not None

    reparsed = parse_docx(out)
    parsed_border = reparsed.page_sections[0].page.border
    assert parsed_border is not None
    assert parsed_border.enabled is True
    assert parsed_border.style == "single"
    assert parsed_border.size_eighth_pt == 8
    assert parsed_border.offset_from == "text"


def test_border_applied_from_profile(tmp_path: Path) -> None:
    """Если рамка задана в профиле, экспорт проставит её даже при border=None."""
    profile = load_profile("gost-r-2.105-2019")
    assert profile.styles.page.border is not None
    # Включаем рамку в профиле (в YAML она по умолчанию выключена).
    profile.styles.page.border.enabled = True

    doc = _doc_with_border(None)
    out = tmp_path / "profile-border.docx"
    export_docx(doc, profile, out)

    reparsed = parse_docx(out)
    parsed_border = reparsed.page_sections[0].page.border
    assert parsed_border is not None
    assert parsed_border.enabled is True


def test_border_disabled_in_profile_writes_nothing(tmp_path: Path) -> None:
    """Выключенная в профиле рамка (default ЕСКД) не пишется."""
    profile = load_profile("gost-r-2.105-2019")
    doc = _doc_with_border(None)
    out = tmp_path / "profile-no-border.docx"
    export_docx(doc, profile, out)

    reparsed = parse_docx(out)
    assert reparsed.page_sections[0].page.border is None


# --- F.07: проверка соответствия рамки профилю -------------------------------


def test_f07_silent_when_profile_requires_no_border() -> None:
    """Профиль без рамки (gost-7.32) → F.07 не срабатывает (отсутствие = норма)."""
    profile = load_profile("gost-7.32-2017")
    assert profile.styles.page.border is None
    doc = _doc_with_border(None)
    assert [v for v in validate(doc, profile) if v.check_code == "F.07"] == []


def test_f07_errors_when_border_missing() -> None:
    """Профиль требует рамку, а её нет → ошибка F.07."""
    profile = _profile_requiring_border()
    doc = _doc_with_border(None)
    f07 = [v for v in validate(doc, profile) if v.check_code == "F.07"]
    assert len(f07) == 1
    assert f07[0].severity == "error"


def test_f07_warns_on_style_mismatch() -> None:
    """Рамка есть, но стиль не тот → warning F.07."""
    profile = _profile_requiring_border()
    doc = _doc_with_border(PageBorder(enabled=True, style="double", size_eighth_pt=4))
    f07 = [v for v in validate(doc, profile) if v.check_code == "F.07"]
    assert len(f07) == 1
    assert f07[0].severity == "warning"


# --- F.07: автофиксер --------------------------------------------------------


def test_f07_fixer_registered() -> None:
    assert "F.07" in registered_fixers()


def test_f07_fixer_adds_border_from_profile() -> None:
    """F.07-фиксер добавляет рамку из профиля; после фикса проверка молчит."""
    profile = _profile_requiring_border()
    doc = _doc_with_border(None)

    applied = fix(doc, profile, codes=["F.07"])
    assert len(applied) == 1
    assert applied[0].fixer_code == "F.07"
    border = doc.page_sections[0].page.border
    assert border is not None and border.enabled is True
    assert [v for v in validate(doc, profile) if v.check_code == "F.07"] == []


def test_f07_fixer_noop_when_profile_requires_no_border() -> None:
    """Если профиль рамку не требует — фиксер ничего не делает."""
    profile = load_profile("gost-7.32-2017")
    doc = _doc_with_border(None)
    assert fix(doc, profile, codes=["F.07"]) == []
    assert doc.page_sections[0].page.border is None
